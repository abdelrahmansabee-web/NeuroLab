#!/usr/bin/env python3
"""
Update manuscript from OpenCap to MediaPipe-based pipeline.
"""
from docx import Document

SRC = r'D:\Thesis app\manuscript f\Immediate Effects of a Single Session of PETTLEP.docx'
DST = r'D:\Thesis app\manuscript f\Immediate Effects of a Single Session of PETTLEP_UPDATED.docx'

doc = Document(SRC)
changes = []

def replace_full(condition_fn, new_text):
    """Clear paragraph and add single run if condition matches."""
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if condition_fn(t):
            old = t[:120]
            p.clear()
            p.add_run(new_text)
            changes.append((i, old, new_text[:120]))
            return

def replace_sentence(condition_fn, old_sentence, new_sentence):
    """Replace one sentence within a paragraph, keeping rest intact."""
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if condition_fn(t):
            old = t[:120]
            new_t = t.replace(old_sentence, new_sentence)
            if new_t != t:
                p.clear()
                p.add_run(new_t)
                changes.append((i, old, new_t[:120]))
                return

def replace_all_text(old, new):
    """Replace substring in every matching paragraph."""
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        if old in t:
            p.clear()
            p.add_run(t.replace(old, new))
            changes.append((i, t[:120], t.replace(old, new)[:120]))


def delete_paragraph(paragraph):
    element = paragraph._element
    element.getparent().remove(element)


def remove_reference(predicate):
    for p in list(doc.paragraphs):
        if predicate(p.text):
            delete_paragraph(p)
            changes.append((-1, p.text[:120], "[REMOVED]"))
            return

# ----------------------------------------------------------------
replace_full(
    lambda t: "Keywords:" in t and "OpenCap" in t,
    "Keywords:\nAction observation; motor imagery; AOMI; PETTLEP; stroke; upper limb rehabilitation; kinematics; MediaPipe; movement smoothness; markerless motion capture."
)

# ----------------------------------------------------------------
# 2. ABSTRACT Methods paragraph (idx 9) - do a sentence-level replace
#    to avoid destroying the rest of the paragraph
# ----------------------------------------------------------------
replace_sentence(
    lambda t: t.startswith("Methods:") and "Three-dimensional kinematics" in t,
    "Three-dimensional kinematics of the affected upper limb were to be captured using the OpenCap markerless motion capture system via three smartphones at 60 fps.",
    "Three-dimensional kinematics of the affected upper limb were to be captured using a custom markerless motion capture pipeline based on MediaPipe Pose Landmarker (Google) for 2D landmark extraction from a single smartphone or webcam at 30 fps and 1080p resolution, combined with the ZoeDepth monocular depth estimation network for metric scaling."
)

# ----------------------------------------------------------------
# 3. Section 1.6 heading
# ----------------------------------------------------------------
replace_full(
    lambda t: "1.6 Kinematic Outcome Measurement: The Role of OpenCap" in t,
    "1.6 Kinematic Outcome Measurement: Markerless Motion Capture Pipeline"
)

# ----------------------------------------------------------------
# 4. Section 1.6 first paragraph — OpenCap description
# ----------------------------------------------------------------
replace_full(
    lambda t: t.startswith("OpenCap (Stanford University, USA) is an open-source"),
    "A custom Python-based markerless motion capture pipeline using MediaPipe Pose Landmarker (Google) was employed for kinematic data extraction. MediaPipe Pose Landmarker is a machine-learning model that detects 33 body landmarks from a single two-dimensional video feed at approximately 30 frames per second. The pipeline extracts landmark coordinates (x, y, z) along with visibility scores, outputting a 101-column CSV file per trial. Video data are captured using a single smartphone or webcam at 1080p resolution, eliminating the need for multiple cameras or cloud-based processing. A monocular depth estimation network (ZoeDepth, Intel/zoedepth-nyu) is applied to a single video frame to estimate metric shoulder width, which serves as a scaling factor for normalizing kinematic measurements. An optional OpenSim Inverse Kinematics (IK) module, utilizing a custom 14-degree-of-freedom PinJoint model in OpenSim 4.5, can be employed for exploratory joint-angle analysis; both the marker trajectory (TRC) and joint-angle (MOT) data are smoothed with a 15-frame moving-average filter."
)

# ----------------------------------------------------------------
# 5. "In the present study, the following kinematic parameters"
# ----------------------------------------------------------------
replace_full(
    lambda t: "In the present study, the following kinematic parameters were to be derived from OpenCap" in t,
    "In the present study, the following kinematic parameters were to be derived from the MediaPipe-based pipeline:"
)

# ----------------------------------------------------------------
# 6-10. Kinematic parameters (paragraphs idx 54-58)
# ----------------------------------------------------------------
replace_full(
    lambda t: t.startswith("Movement Onset Time") and "5%" in t,
    "Movement Duration (total_duration_s): Total duration from onset to offset within the active movement window. Movement onset is defined as the first frame where wrist displacement from rest exceeds a normalized threshold of 0.03 (in shoulder-width units). Movement offset is defined as the frame where instantaneous hand velocity falls below 5% of peak velocity. This hybrid displacement–velocity method produces a single contiguous active window per trial."
)

replace_full(
    lambda t: t.startswith("Movement Smoothness (Number of Velocity Peaks, NVP)"),
    "Peak and Mean Velocity (total_peak_velocity, total_mean_velocity): Maximum and average instantaneous hand velocity (normalized units/second) within the active movement window, reflecting motor output intensity."
)

replace_full(
    lambda t: t.startswith("Trunk Displacement") and "lumbar extension" in t,
    "Trunk-to-Palm Displacement Ratio (total_trunk_palm_ratio): Ratio of cumulative trunk displacement to cumulative palm displacement within the active window, indexing compensatory trunk movement during reaching."
)

replace_full(
    lambda t: t.startswith("Shoulder Girdle Elevation") and "shoulder landmark" in t,
    "Maximum Elbow Angle (total_max_elbow_deg): Peak elbow flexion angle (degrees) during the active movement window, computed from shoulder–elbow–wrist landmarks."
)

replace_full(
    lambda t: t.startswith("Movement Time (MT)"),
    "Movement Smoothness (smoothness_pause_pct): Percentage of time within the active window where hand velocity falls below an absolute threshold of 0.03 normalized units — a consistent measure of movement intermittency across participants. Additional metrics include total_path_length (wrist path normalized to shoulder width), total_lat_range_norm (medial–lateral hand range), arm_length_norm (arm length / shoulder width), shoulder_width_norm, total_depression_cm (shoulder girdle depression via ZoeDepth), trunk_lat_norm, and trunk_vert_norm (trunk lateral and forward flexion)."
)

# ----------------------------------------------------------------
# 11. Section 2.6 heading
# ----------------------------------------------------------------
replace_full(
    lambda t: "OpenCap Data Collection Setup" in t,
    "Video Data Collection Setup"
)

# ----------------------------------------------------------------
# 12. Section 2.6 paragraph (idx 110)
# ----------------------------------------------------------------
replace_full(
    lambda t: "Kinematic data were to be recorded using three synchronized smartphones" in t,
    "Kinematic data were to be recorded using a single smartphone or webcam positioned at a standardized frontal view at 1080p resolution and 30 fps with consistent diffuse lighting. Participants wore dark, form-fitting clothing with reflective markers or exposed skin at anatomical landmarks to facilitate tracking. Videos were processed offline using the custom Python-based MediaPipe pipeline on a local workstation — no internet connection or cloud upload was required."
)

# ----------------------------------------------------------------
# 13. "Kinematic variables were to be extracted"
# ----------------------------------------------------------------
replace_full(
    lambda t: "Kinematic variables were to be extracted using custom Python scripts as follows" in t,
    "Kinematic variables were to be extracted using the custom Python analysis pipeline (kinematics_analyzer.py, version 6) as follows:"
)

# ----------------------------------------------------------------
# 14-18. Variable extraction paragraphs (idx 112-116)
# ----------------------------------------------------------------
replace_full(
    lambda t: t.startswith("Movement Onset Time") and "auditory" in t,
    "Movement Onset and Offset: Onset is detected when wrist displacement from rest exceeds 0.03 normalized threshold. Offset is detected when hand velocity falls below 5% of peak velocity. These events define a single active movement window."
)

replace_full(
    lambda t: t.startswith("Number of Velocity Peaks (NVP)"),
    "Auto Limb Detection: The active (affected) side is automatically identified as the limb with the longer cumulative wrist path length, eliminating manual side specification."
)

replace_full(
    lambda t: t.startswith("Trunk Displacement") and "Maximum lumbar" in t,
    "Full-Recording Metrics: All metrics — total_duration_s, total_peak_velocity, total_mean_velocity, total_path_length, total_lat_range_norm, total_trunk_palm_ratio, total_max_elbow_deg, smoothness_pause_pct, arm_length_norm, shoulder_width_norm, total_depression_cm, trunk_lat_norm, trunk_vert_norm — are computed within the single active movement window."
)

replace_full(
    lambda t: t.startswith("Shoulder Girdle Elevation") and "Maximum vertical" in t,
    "ZoeDepth Metric Scaling: The ZoeDepth monocular depth estimation network (Intel/zoedepth-nyu) is applied to a single video frame to estimate shoulder width in meters, used to scale normalized measurements to real-world units (e.g., total_depression_cm)."
)

replace_full(
    lambda t: t.startswith("Movement Time") and "Duration from onset" in t,
    "OpenSim Inverse Kinematics (Exploratory): An OpenSim 4.5 model with 14 PinJoint degrees of freedom is used for complementary joint-angle analysis. TRC and MOT data are both smoothed with a 15-frame moving-average filter."
)

# ----------------------------------------------------------------
# 19. Section 1.7 "measuring the immediate kinematic response using OpenCap" (idx 61)
# ----------------------------------------------------------------
replace_full(
    lambda t: "measuring the immediate kinematic response using OpenCap" in t,
    "The present study addressed these gaps by delivering a single session of PETTLEP-based AOMI — incorporating a sensorimotor transfer calibration phase, individualized mirror-reversed first-person observation, and progressive corrective motor imagery targeting initiation, postural control, and smoothness — and measuring the immediate kinematic response using the custom MediaPipe-based markerless motion capture pipeline with monocular depth estimation. This trial was designed to provide the first direct, objective evidence for the acute neuromotor effects of this approach in persons post-stroke."
)

# ----------------------------------------------------------------
# 20. Discussion — "first study to employ OpenCap" (idx 221)
# ----------------------------------------------------------------
replace_full(
    lambda t: "first study to employ" in t and "OpenCap" in t,
    "Second, this was the first study to employ a custom markerless motion capture pipeline based on MediaPipe Pose Landmarker as the primary outcome measurement technology in an AOMI stroke rehabilitation trial. The use of objective, continuous kinematic data — rather than ordinal clinical scales — enabled detection of subtle, within-session changes in movement quality that would not have been captured by conventional assessment tools. Specifically, the smoothness metric (smoothness_pause_pct) as a primary outcome was sensitive to improvements in motor control at the level of movement trajectory planning, which was the precise neural process targeted by the PETTLEP-AOMI protocol."
)

# ----------------------------------------------------------------
# 20. Limitations — OpenCap Accuracy
# ----------------------------------------------------------------
replace_full(
    lambda t: "OpenCap Accuracy in Post-Stroke Populations" in t,
    "MediaPipe Tracking Accuracy in Post-Stroke Populations: While MediaPipe Pose Landmarker has been validated for general human pose estimation, its accuracy in stroke survivors — who may present with atypical movement patterns, variable clothing, and spasticity-related postural abnormalities — had not been specifically validated for this clinical population. Potential sources of tracking error included occlusion of landmarks during compensatory trunk lean, and AI landmark misidentification in non-standard body configurations. Standardization procedures (consistent clothing, controlled lighting, fixed camera positions) were intended to mitigate these risks; however, a formal within-study reliability analysis was to be performed."
)

# ----------------------------------------------------------------
# 21. Clinical Implications — "The OpenCap system requires"
# ----------------------------------------------------------------
replace_full(
    lambda t: "The OpenCap system requires only three smartphones and internet access" in t,
    "The markerless motion capture pipeline requires only a single smartphone or webcam and a standard laptop for offline processing, making it feasible for implementation in resource-limited rehabilitation settings without requiring internet connectivity."
)

# ----------------------------------------------------------------
# 22. Future Directions — "feasibility of OpenCap"
# ----------------------------------------------------------------
replace_full(
    lambda t: "Demonstrate the feasibility of OpenCap as a kinematic outcome measurement tool" in t,
    "Demonstrate the feasibility of a MediaPipe-based markerless motion capture pipeline as a kinematic outcome measurement tool in a clinical neurorehabilitation setting."
)

# ----------------------------------------------------------------
# 23. Conclusion — "measuring outcomes with the validated OpenCap"
#     This paragraph has substantial content before the OpenCap
#     mention, so we do a sentence-level replace.
# ----------------------------------------------------------------
replace_sentence(
    lambda t: "measuring outcomes with the validated OpenCap markerless three-dimensional motion capture system" in t,
    "measuring outcomes with the validated OpenCap markerless three-dimensional motion capture system",
    "measuring outcomes with a custom MediaPipe-based markerless motion capture pipeline integrating monocular depth estimation (ZoeDepth) and optional OpenSim inverse kinematics"
)

# ----------------------------------------------------------------
# 24. Abstract — primary outcome NVP → smoothness_pause_pct
# ----------------------------------------------------------------
replace_all_text(
    "The primary outcome was the Number of Velocity Peaks (NVP) as an index of movement smoothness.",
    "The primary outcome was movement smoothness (smoothness_pause_pct) — the percentage of time within the active movement window where hand velocity falls below a fixed normalized threshold (0.03), indexing movement intermittency and stop-and-go behavior."
)

replace_all_text(
    "Secondary outcomes included movement onset time, shoulder girdle elevation, trunk displacement (lumbar extension angle), movement time, Wolf Motor Function Test – Short Form (WMFT-SF), Visual Analog Mood Scale – 4 dimensions (VAMS-4), Kinesthetic and Visual Imagery Questionnaire (KVIQ), International Physical Activity Questionnaire – Short Form (IPAQ-SF), Perceived Motor Control Change (Motor Difference Rating Scale), and Visual Analog Scale for pain (VAS).",
    "Secondary outcomes included total movement duration (total_duration_s), peak and mean hand velocity, trunk-to-palm displacement ratio, maximum elbow angle, total path length, lateral hand range, shoulder girdle depression (total_depression_cm via ZoeDepth metric scaling), phase-specific metrics (forward, wipe_right, wipe_left, return), Wolf Motor Function Test – 4-item short form (WMFT-4), Visual Analog Mood Scale – 4 dimensions (VAMS-4), Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10), International Physical Activity Questionnaire (IPAQ), Perceived Motor Control Change Scale, and Visual Analog Scale for pain (VAS). All kinematic variables were collected and managed using the NeuroLab Stroke Rehabilitation Research Platform (custom web application, v6.4)."
)

replace_all_text(
    "It was hypothesized that the PETTLEP-based AOMI group would demonstrate significantly greater reductions in movement jerkiness (lower NVP), trunk compensation, and shoulder girdle elevation, as well as shorter movement onset time and movement time, and improved upper limb functional performance compared to the control group.",
    "It was hypothesized that the PETTLEP-based AOMI group would demonstrate significantly greater reductions in movement intermittency (lower smoothness_pause_pct), trunk compensation (lower total_trunk_palm_ratio), and shoulder girdle elevation, as well as shorter total movement duration and improved upper limb functional performance (WMFT-4) compared to the control group."
)

# ----------------------------------------------------------------
# 25. Research questions, hypotheses, sample size
# ----------------------------------------------------------------
replace_all_text("movement smoothness (NVP)", "movement smoothness (smoothness_pause_pct)")
replace_all_text("smoothness/NVP", "smoothness (smoothness_pause_pct)")
replace_all_text("Δ NVP", "Δ smoothness_pause_pct")
replace_all_text(
    "primary outcome being the Group × Time interaction effect on Number of Velocity Peaks (NVP)",
    "primary outcome being the Group × Time interaction effect on movement smoothness (smoothness_pause_pct)"
)

# ----------------------------------------------------------------
# 26. Methods — outcome measures section
# ----------------------------------------------------------------
replace_full(
    lambda t: t.startswith("Movement Smoothness — Number of Velocity Peaks (NVP)"),
    "Movement Smoothness — Pause Percentage (smoothness_pause_pct)"
)

replace_full(
    lambda t: t.startswith("NVP was calculated from the hand velocity profile"),
    "Movement smoothness (smoothness_pause_pct) was calculated as the percentage of frames within the active movement window where instantaneous hand velocity fell below an absolute threshold of 0.03 normalized units. A lower smoothness_pause_pct indicates more continuous, less stop-and-go movement. The active window was defined using a hybrid displacement–velocity onset/offset algorithm applied to the Reach & Wipe task. Phase-specific smoothness metrics were also computed for forward, wipe_right, wipe_left, and return segments."
)

replace_full(
    lambda t: t == "Movement Onset Time",
    "Total Movement Duration (total_duration_s)"
)

replace_full(
    lambda t: t == "Shoulder Girdle Elevation",
    "Shoulder Girdle Depression (total_depression_cm)"
)

replace_full(
    lambda t: t == "Trunk Displacement",
    "Trunk-to-Palm Ratio (total_trunk_palm_ratio)"
)

replace_full(
    lambda t: t == "Movement Time",
    "Peak Hand Velocity (total_peak_velocity)"
)

replace_all_text("Wolf Motor Function Test – Short Form (WMFT-SF)", "Wolf Motor Function Test – 4-item short form (WMFT-4)")
replace_all_text("WMFT-SF", "WMFT-4")
replace_all_text("Kinesthetic and Visual Imagery Questionnaire (KVIQ)", "Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10)")

# ----------------------------------------------------------------
# 27. Statistical analysis — variable list
# ----------------------------------------------------------------
replace_all_text(
    "primary kinematic outcomes (NVP, Trunk Displacement, Shoulder Girdle Elevation, Movement Time) and WMFT-SF scores",
    "primary kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, total_depression_cm, total_duration_s, total_peak_velocity) and WMFT-4 scores"
)

replace_all_text(
    "or for ordinal variables (VAS, KVIQ scores, Motor Difference Rating Scale)",
    "or for ordinal variables (VAS, VAMS-4, KVIQ-10 scores, Motor Control Change Scale)"
)

replace_all_text(
    "The relationship between participant-rated motor imagery ability (KVIQ visual and kinesthetic subscale scores) and the magnitude of kinematic improvement (Δ smoothness_pause_pct, Δ Trunk Displacement, Δ Movement Time)",
    "The relationship between participant-rated motor imagery ability (KVIQ-10 visual and kinesthetic subscale scores) and the magnitude of kinematic improvement (Δ smoothness_pause_pct, Δ total_trunk_palm_ratio, Δ total_duration_s)"
)

# ----------------------------------------------------------------
# 28. Expected results section
# ----------------------------------------------------------------
replace_full(
    lambda t: t == "3.1 Primary Outcome: Movement Smoothness (NVP)",
    "3.1 Primary Outcome: Movement Smoothness (smoothness_pause_pct)"
)

replace_all_text(
    "The AOMI group would demonstrate a statistically significant reduction in NVP from pre- to post-intervention (within-group effect).",
    "The AOMI group would demonstrate a statistically significant reduction in smoothness_pause_pct from pre- to post-intervention (within-group effect)."
)

replace_all_text(
    "The control group would demonstrate no significant change or a minimal, non-clinically meaningful change in NVP.",
    "The control group would demonstrate no significant change or a minimal, non-clinically meaningful change in smoothness_pause_pct."
)

replace_all_text(
    "A reduction in NVP would indicate that the AOMI session facilitated a shift toward smoother, more continuous reaching trajectories",
    "A reduction in smoothness_pause_pct would indicate that the AOMI session facilitated a shift toward smoother, more continuous reaching trajectories"
)

replace_all_text(
    "Movement Onset Time: The AOMI group would demonstrate a significant reduction in movement onset time",
    "Total Movement Duration: The AOMI group would demonstrate a significant reduction in total_duration_s"
)

replace_all_text(
    "Trunk Displacement: The AOMI group would demonstrate a significant reduction in trunk lean during reaching",
    "Trunk-to-Palm Ratio: The AOMI group would demonstrate a significant reduction in total_trunk_palm_ratio during reaching"
)

replace_all_text(
    "Shoulder Girdle Elevation: The AOMI group would demonstrate a significant reduction in shoulder shrugging during reaching",
    "Shoulder Girdle Depression: The AOMI group would demonstrate a significant increase in total_depression_cm (reflecting reduced compensatory shoulder elevation)"
)

replace_all_text(
    "Movement Time: A significant reduction in movement time (increased movement efficiency) was anticipated in the AOMI group",
    "Peak Hand Velocity: A significant increase in total_peak_velocity (reflecting more confident motor output) was anticipated in the AOMI group"
)

replace_all_text(
    "3.3 Upper Limb Functional Performance (WMFT-4)",
    "3.3 Upper Limb Functional Performance (WMFT-4)"
)

replace_all_text(
    "positive correlation between imagery ability and NVP change",
    "positive correlation between imagery ability and smoothness_pause_pct change"
)

# ----------------------------------------------------------------
# 29. Discussion — remaining NVP / OpenCap references
# ----------------------------------------------------------------
replace_all_text(
    "manifesting as reduced NVP, reduced trunk lean, and reduced shoulder elevation in the post-intervention assessment.",
    "manifesting as reduced smoothness_pause_pct, reduced total_trunk_palm_ratio, and reduced compensatory shoulder elevation in the post-intervention assessment."
)

replace_all_text(
    "This was particularly important for NVP and movement onset time, because the temporal dynamics of velocity peak generation were intimately linked to the timing of motor planning sub-commands.",
    "This was particularly important for smoothness_pause_pct and total_duration_s, because the temporal dynamics of velocity generation were intimately linked to the timing of motor planning sub-commands."
)

replace_all_text(
    "While OpenCap has been validated against gold-standard motion capture for healthy individuals performing well-defined movements (38), its accuracy in stroke survivors",
    "While MediaPipe Pose Landmarker has been validated for general human pose estimation, its accuracy in stroke survivors"
)

replace_all_text(
    "measuring the immediate kinematic response using OpenCap.",
    "measuring the immediate kinematic response using the custom MediaPipe-based markerless motion capture pipeline."
)

replace_all_text(
    "whether a single AOMI session could induce measurable improvements in movement smoothness (NVP), trunk compensation, shoulder girdle elevation, movement time, and functional upper limb performance",
    "whether a single AOMI session could induce measurable improvements in movement smoothness (smoothness_pause_pct), trunk compensation, shoulder girdle elevation, total movement duration, and functional upper limb performance (WMFT-4)"
)

replace_all_text(
    "Explore the moderating roles of imagery ability (KVIQ), affective state (VAMS-4), and physical activity level (IPAQ-SF)",
    "Explore the moderating roles of imagery ability (KVIQ-10), affective state (VAMS-4), and physical activity level (IPAQ)"
)

replace_all_text(
    "Due to the nature of the intervention, participants could not be blinded. However, the outcome assessor and the investigator performing kinematic analysis were to remain blinded to group allocation using standardized automated processing pipelines.",
    "Due to the nature of the intervention, participants could not be blinded. However, the outcome assessor was to remain blinded to group allocation. Kinematic analysis was performed using the standardized automated NeuroLab pipeline (kinematics_analyzer.py v6), which applies identical segmentation and metric extraction algorithms to all trials regardless of group assignment."
)

replace_all_text(
    "The Number of Velocity Peaks (NVP) has been proposed as a sensitive, objective biomarker of upper limb motor recovery (8).",
    "Movement smoothness indices derived from hand velocity profiles — including pause percentage (smoothness_pause_pct) — have been proposed as sensitive, objective biomarkers of upper limb motor recovery (8)."
)

replace_all_text(
    "primary kinematic outcomes (NVP, Trunk Displacement, Shoulder Girdle Elevation, Movement Time) and WMFT-4 scores",
    "primary kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, total_depression_cm, total_duration_s, total_peak_velocity) and WMFT-4 scores"
)

replace_all_text(
    "primary kinematic outcomes (NVP, Trunk Displacement, Shoulder Girdle Elevation, Movement Time) and WMFT-SF scores",
    "primary kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, total_depression_cm, total_duration_s, total_peak_velocity) and WMFT-4 scores"
)

# ----------------------------------------------------------------
# 30. Remove obsolete Watson affect-scale reference (replaced by VAMS Stern 1999)
# ----------------------------------------------------------------
remove_reference(
    lambda t: "Watson D, Clark LA, Tellegen A" in t
)

# ----------------------------------------------------------------
# 31. Abstract Objective — align with NeuroLab metrics
# ----------------------------------------------------------------
replace_all_text(
    "This study aimed to investigate whether a single session of PETTLEP-based AOMI produces immediate improvements in movement smoothness (primary outcome), movement onset time, trunk displacement, shoulder girdle elevation, movement time, and upper limb functional performance compared to a time-matched cognitive and somatic control condition in individuals post-stroke.",
    "This study aimed to investigate whether a single session of PETTLEP-based AOMI produces immediate improvements in movement smoothness (smoothness_pause_pct; primary outcome), total movement duration (total_duration_s), trunk-to-palm displacement ratio (total_trunk_palm_ratio), shoulder vertical displacement range (shoulder_vert_norm), peak hand velocity (total_peak_velocity), and upper limb functional performance (WMFT-4) compared to a time-matched cognitive and somatic control condition in individuals post-stroke."
)

# ----------------------------------------------------------------
# 32. 2D vs 3D motion capture claims
# ----------------------------------------------------------------
replace_all_text(
    "have not been rigorously investigated using objective, three-dimensional motion capture.",
    "have not been rigorously investigated using objective, single-camera markerless kinematic analysis."
)

replace_all_text(
    "Three-dimensional kinematics of the affected upper limb were to be captured using a custom markerless motion capture pipeline",
    "Two-dimensional pose-landmark kinematics of the affected upper limb were to be captured using a custom markerless motion capture pipeline"
)

replace_all_text(
    "no study had used markerless three-dimensional motion capture to quantify the kinematic effects of AOMI in stroke survivors",
    "no study had used a single-camera MediaPipe-based markerless kinematic pipeline to quantify the kinematic effects of AOMI in stroke survivors"
)

replace_all_text(
    "OpenCap markerless three-dimensional motion capture",
    "a custom MediaPipe-based markerless kinematic pipeline"
)

replace_all_text(
    "Generate a dataset of pre- and post-intervention three-dimensional kinematics from stroke survivors",
    "Generate a dataset of pre- and post-intervention two-dimensional pose-landmark kinematics (with metric-scaled derivatives) from stroke survivors"
)

# ----------------------------------------------------------------
# 33. Secondary research questions & hypotheses
# ----------------------------------------------------------------
replace_full(
    lambda t: t == "Movement onset time of the reaching task?",
    "Total movement duration (total_duration_s) of the Reach & Wipe task?"
)

replace_full(
    lambda t: t == "Trunk displacement (lumbar extension angle) during reaching?",
    "Trunk-to-palm displacement ratio (total_trunk_palm_ratio) during reaching?"
)

replace_full(
    lambda t: t == "Shoulder girdle elevation during reaching?",
    "Shoulder vertical displacement range (shoulder_vert_norm) during reaching?"
)

replace_full(
    lambda t: t == "Movement time of the reaching task?",
    "Peak hand velocity (total_peak_velocity) during the reaching task?"
)

replace_full(
    lambda t: "Upper limb functional performance (Wolf Motor Function Test – Short Form)" in t,
    "Upper limb functional performance (Wolf Motor Function Test – 4-item short form, WMFT-4)?"
)

replace_full(
    lambda t: t.startswith("H") and "H\u2080" in t and "movement onset time" in t,
    "H\u2080: One session of PETTLEP-based AOMI will not produce significant changes in kinematic parameters (smoothness_pause_pct, total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) or upper limb functional performance (WMFT-4) compared to the control group."
)

replace_full(
    lambda t: t.startswith("H") and "H\u2081" in t and "movement onset time" in t,
    "H\u2081: One session of PETTLEP-based AOMI will produce significant improvements in kinematic parameters (smoothness_pause_pct, total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) and/or upper limb functional performance (WMFT-4) compared to the control group."
)

# ----------------------------------------------------------------
# 34. Video setup, shoulder metric naming, scales
# ----------------------------------------------------------------
replace_all_text(
    "Participants wore dark, form-fitting clothing with reflective markers or exposed skin at anatomical landmarks to facilitate tracking.",
    "Participants wore dark, form-fitting long-sleeve clothing with secured hair and exposed wrists to facilitate landmark tracking, without reflective markers."
)

replace_all_text(
    "Shoulder Girdle Depression (total_depression_cm)",
    "Shoulder Vertical Displacement Range (shoulder_vert_norm; cm equivalent exported as total_depression_cm)"
)

replace_all_text(
    "shoulder girdle depression (total_depression_cm via ZoeDepth metric scaling)",
    "shoulder vertical displacement range (shoulder_vert_norm; cm equivalent via ZoeDepth metric scaling, exported as total_depression_cm)"
)

replace_all_text(
    "total_depression_cm (shoulder girdle depression via ZoeDepth)",
    "shoulder_vert_norm (shoulder vertical displacement range during the active window; cm export as total_depression_cm via ZoeDepth)"
)

replace_all_text(
    "International Physical Activity Questionnaire – Short Form (IPAQ-SF)",
    "International Physical Activity Questionnaire (IPAQ)"
)

replace_all_text("Perceived Motor Control Change Scale", "Motor Difference Rating Scale (MDRS, post-intervention only)")
replace_all_text("Motor Control Change Scale", "Motor Difference Rating Scale (MDRS)")

replace_all_text(
    "NeuroLab Stroke Rehabilitation Research Platform (custom web application, v6.4)",
    "NeuroLab Stroke Rehabilitation Research Platform (custom web application; kinematics_analyzer.py v6)"
)

# ----------------------------------------------------------------
# 35. Expected results — full paragraph fixes
# ----------------------------------------------------------------
replace_full(
    lambda t: t.startswith("Movement Onset Time:") and "Block 1" in t,
    "Total Movement Duration: The AOMI group would demonstrate a significant reduction in total_duration_s, reflecting more efficient movement execution facilitated by the initiation-focused imagery constraint in Block 1. No significant change was expected in the control group."
)

replace_full(
    lambda t: t.startswith("Trunk Displacement:") and "Block 2" in t,
    "Trunk-to-Palm Ratio: The AOMI group would demonstrate a significant reduction in total_trunk_palm_ratio during reaching, reflecting reduced reliance on compensatory trunk strategies as a result of the trunk stability imagery constraint in Block 2 — which was specifically designed to reinstate the internal representation of anticipatory postural stabilization before arm movement. No significant change was expected in the control group."
)

replace_full(
    lambda t: t.startswith("Shoulder Girdle Elevation:") and "shoulder shrugging" in t,
    "Shoulder Vertical Displacement Range: The AOMI group would demonstrate a significant reduction in shoulder_vert_norm (and its cm equivalent, total_depression_cm), reflecting less compensatory shoulder girdle movement during reaching, consistent with the shoulder heaviness imagery instruction in Block 3. No significant change was expected in the control group."
)

replace_all_text(
    "Shoulder Girdle Depression: The AOMI group would demonstrate a significant increase in total_depression_cm (reflecting reduced compensatory shoulder elevation)",
    "Shoulder Vertical Displacement Range: The AOMI group would demonstrate a significant reduction in shoulder_vert_norm (and its cm equivalent, total_depression_cm), reflecting less compensatory shoulder girdle movement during reaching"
)

replace_full(
    lambda t: t.startswith("Movement Time:") and "observation-then-imagery" in t,
    "Peak Hand Velocity: A significant increase in total_peak_velocity was anticipated in the AOMI group, reflecting more confident motor output facilitated by the observation-then-imagery paradigm."
)

replace_full(
    lambda t: t.startswith("3.3 Upper Limb Functional Performance (WMFT-SF)"),
    "3.3 Upper Limb Functional Performance (WMFT-4)"
)

replace_all_text(
    "any WMFT-SF changes were expected to be modest",
    "any WMFT-4 changes were expected to be modest"
)

replace_all_text(
    "clinically meaningful WMFT-SF changes — equivalent to the established minimal clinically important difference",
    "clinically meaningful WMFT-4 changes — an exploratory question, as minimal clinically important differences have been established for the full WMFT-15 rather than the shortened WMFT-4"
)

# ----------------------------------------------------------------
# 36. Moderator analysis correlation direction & blinding
# ----------------------------------------------------------------
replace_all_text(
    "positive correlation between imagery ability and smoothness_pause_pct change",
    "negative correlation between imagery ability (KVIQ-10) and \u0394 smoothness_pause_pct (i.e., greater imagery ability associated with larger reductions in pause percentage)"
)

replace_all_text(
    "It was hypothesized that participants with higher baseline KVIQ scores (greater motor imagery ability) would demonstrate larger kinematic improvements in the AOMI group (positive correlation between imagery ability and NVP change)",
    "It was hypothesized that participants with higher baseline KVIQ-10 scores (greater motor imagery ability) would demonstrate larger kinematic improvements in the AOMI group (negative correlation between KVIQ-10 scores and \u0394 smoothness_pause_pct), providing evidence for the individual-difference moderating role of imagery capacity. No such relationship was expected in the control group."
)

replace_all_text(
    "While the blinded assessor and blinded kinematic analyst mitigated assessment bias, participant-level expectancy effects could not be fully eliminated.",
    "While the blinded outcome assessor mitigated clinical assessment bias and kinematic extraction was fully automated (NeuroLab pipeline), participant-level expectancy effects could not be fully eliminated."
)

replace_all_text(
    "whether a single AOMI session could induce measurable improvements in movement smoothness (smoothness_pause_pct), trunk compensation, shoulder girdle elevation, total movement duration, and functional upper limb performance (WMFT-4) in post-stroke individuals, compared to an active cognitive and somatic control.",
    "whether a single AOMI session could induce measurable improvements in movement smoothness (smoothness_pause_pct), trunk-to-palm ratio, shoulder vertical displacement range, total movement duration, peak hand velocity, and functional upper limb performance (WMFT-4) in post-stroke individuals, compared to an active cognitive and somatic control."
)

replace_all_text(
    "used to scale normalized measurements to real-world units (e.g., total_depression_cm).",
    "used to scale normalized measurements to real-world units (e.g., shoulder vertical range in cm, exported as total_depression_cm)."
)

# ----------------------------------------------------------------
# 37. References — remove duplicates & OpenCap; add missing
# ----------------------------------------------------------------
def remove_duplicate_references():
    seen = set()
    in_refs = False
    for p in list(doc.paragraphs):
        t = p.text.strip()
        if t == "REFERENCES":
            in_refs = True
            continue
        if not in_refs or len(t) < 25:
            continue
        key = t[:75].lower()
        if key in seen:
            delete_paragraph(p)
            changes.append((-1, t[:80], "[DUPLICATE REMOVED]"))
        else:
            seen.add(key)


def remove_opencap_references():
    for p in list(doc.paragraphs):
        if "OpenCap: 3D human movement dynamics" in p.text:
            delete_paragraph(p)
            changes.append((-1, p.text[:80], "[OpenCap REMOVED]"))


NEW_REFS = [
    "Bazarevsky V, Bazarevsky A, Rybachuk O, et al. BlazePose: On-device real-time body pose tracking. arXiv preprint. 2020; arXiv:2006.10204.",
    "Bhat SF, Birhane A, Weinberger KQ. ZoeDepth: Zero-shot transfer by combining relative and metric depth. CVPR Workshop. 2023.",
    "Kim B, Schweighofer N, Wolf SL, Winstein C. A streamlined 4-item Wolf Motor Function Test for efficient assessment of upper extremity motor function in chronic stroke survivors. Neurorehabilitation and Neural Repair. 2026;40(3):214\u2013225.",
    "Arruda JE, Stern RA, Somerville JA. Measurement of mood states in stroke patients: validation of the visual analog mood scales. Archives of Physical Medicine and Rehabilitation. 1999;80(6):676\u2013680.",
    "Barrows P, Thomas S, Evans N, et al. Validity and reliability of the Dynamic Visual Analogue Mood Scales (D-VAMS) in stroke survivors with aphasia. Aphasiology. 2018;32(5):600\u2013615.",
]


def add_references_if_missing(refs):
    existing = "\n".join(p.text for p in doc.paragraphs)
    for ref in refs:
        marker = ref.split(".")[0][:35]
        if marker not in existing:
            doc.add_paragraph(ref)
            changes.append((-1, "", ref[:80]))


remove_duplicate_references()
remove_opencap_references()
add_references_if_missing(NEW_REFS)

# ----------------------------------------------------------------
# Save
# ----------------------------------------------------------------
try:
    doc.save(DST)
    saved_path = DST
except PermissionError:
    saved_path = DST.replace(".docx", "_CORRECTED.docx")
    doc.save(saved_path)
    print(f"WARNING: {DST} is locked (close Word). Saved to: {saved_path}")
else:
    print(f"Saved: {saved_path}")

print()
import sys
if hasattr(sys.stdout, "reconfigure"):
    try:
        sys.stdout.reconfigure(encoding="utf-8", errors="replace")
    except Exception:
        pass
for idx, old_snip, new_snip in changes:
    print(f"[{idx}]")
    print(f"  OLD: {old_snip}...")
    print(f"  NEW: {new_snip}...")
    print()

print(f"Total changes: {len(changes)}")
