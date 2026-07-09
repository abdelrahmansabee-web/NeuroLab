# -*- coding: utf-8 -*-
"""
Update the edited ethics form to describe a reach-to-grasp task:
  - Rename Reach & Return -> Reach-to-Grasp (ulaşma-kavrama).
  - Cup/object distance determined by the healthy arm's maximal reach.
  - Add "number of stops" (durak sayısı) as a movement-control variable.
Only the newly inserted/altered text is coloured red; existing formatting
(including previous red edits) is preserved.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

src = Path(r"D:\الاخلاقيات\Ethics BKK Last version  (AutoRecovered)_edited.docx")
out = src.with_stem(src.stem + "_reach2grasp")
backup = src.with_stem(src.stem + "_before_reach2grasp")

import shutil
shutil.copy2(src, backup)

doc = Document(str(src))
RED = RGBColor(0xFF, 0x00, 0x00)


def replace_in_runs(p, old_text, new_text):
    """
    Replace the first occurrence of old_text in paragraph p, preserving the
    formatting of surrounding runs. Only the inserted new_text is coloured red.
    Returns True if a replacement was made.
    """
    full = p.text
    idx = full.find(old_text)
    if idx == -1:
        return False

    # Cumulative character positions of runs
    pos = 0
    run_info = []
    for r in p.runs:
        run_info.append((pos, pos + len(r.text), r))
        pos += len(r.text)

    start_run = None
    end_run = None
    for i, (s, e, r) in enumerate(run_info):
        if s <= idx < e and start_run is None:
            start_run = i
        if s < idx + len(old_text) <= e:
            end_run = i
            break
        elif idx + len(old_text) == e:
            end_run = i
            break

    if start_run is None or end_run is None:
        # Fallback: rebuild paragraph keeping text black except new text red
        replace_paragraph_text(p, full.replace(old_text, new_text, 1), red=False)
        # colour only the new text red
        _colour_substring(p, new_text)
        return True

    # Collect prefix (before old_text) and suffix (after old_text) preserving runs
    prefix_runs = []
    suffix_runs = []
    inside_prefix = True
    for i, (s, e, r) in enumerate(run_info):
        if i < start_run:
            prefix_runs.append(r)
        elif i > end_run:
            suffix_runs.append(r)
        elif i == start_run and i == end_run:
            # Old text is inside a single run
            run_start_in_run = idx - s
            run_end_in_run = idx + len(old_text) - s
            prefix_part = r.text[:run_start_in_run]
            suffix_part = r.text[run_end_in_run:]
            prefix_runs.append((r, prefix_part))
            suffix_runs.append((r, suffix_part))
        elif i == start_run:
            run_start_in_run = idx - s
            prefix_part = r.text[:run_start_in_run]
            prefix_runs.append((r, prefix_part))
        elif i == end_run:
            run_end_in_run = idx + len(old_text) - s
            suffix_part = r.text[run_end_in_run:]
            suffix_runs.insert(0, (r, suffix_part))
        # middle runs between start and end are dropped (they contained old_text)

    # Rebuild runs
    new_runs = []
    for item in prefix_runs:
        if isinstance(item, tuple):
            r, txt = item
        else:
            r, txt = item, item.text
        if txt:
            new_runs.append((txt, r.font.color.rgb if r.font.color and r.font.color.rgb else None))
    new_runs.append((new_text, RED))
    for item in suffix_runs:
        if isinstance(item, tuple):
            r, txt = item
        else:
            r, txt = item, item.text
        if txt:
            new_runs.append((txt, r.font.color.rgb if r.font.color and r.font.color.rgb else None))

    # Clear existing runs
    for run in p.runs:
        run._element.getparent().remove(run._element)

    for txt, color in new_runs:
        run = p.add_run(txt)
        if color is not None:
            run.font.color.rgb = color

    return True


def replace_paragraph_text(p, new_text, red=True):
    """Clear paragraph runs and insert new_text (optionally in red)."""
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(new_text)
    if red:
        run.font.color.rgb = RED
    return run


def _colour_substring(p, sub_text):
    """Colour the first occurrence of sub_text in p red."""
    full = p.text
    idx = full.find(sub_text)
    if idx == -1:
        return
    # Simple rebuild: prefix black, sub red, suffix black
    prefix = full[:idx]
    suffix = full[idx + len(sub_text):]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    if prefix:
        p.add_run(prefix)
    r = p.add_run(sub_text)
    r.font.color.rgb = RED
    if suffix:
        p.add_run(suffix)


def insert_after_paragraph(doc, p, text, red=True):
    """Insert a new paragraph after p."""
    parent = p._element.getparent()
    idx = list(parent).index(p._element)
    np = doc.add_paragraph(text)
    if red and np.runs:
        np.runs[0].font.color.rgb = RED
    parent.insert(idx + 1, np._element)
    return np


# ------------------------------------------------------------------
# Edits
# ------------------------------------------------------------------

# 1. Justification / aim paragraph (para 35): add number of stops
old = ("Üst ekstremite hareket kontrolü (NVP, straightness, pause time), gövde kompanzasyonu (trunk ratio), "
       "omuz kuşağı elevasyonu (shoulder elevation), dirsek açısı, hareket süresi ve tepe hız gibi "
       "kinematik parametreler objektif olarak ölçülecektir (8).")
new = ("Üst ekstremite hareket kontrolü (NVP, straightness, pause time, number of stops), gövde kompanzasyonu (trunk ratio), "
       "omuz kuşağı elevasyonu (shoulder elevation), dirsek açısı, hareket süresi ve tepe hız gibi "
       "kinematik parametreler objektif olarak ölçülecektir (8).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 35 (justification)")
        break

# 2. Revision paragraph (para 50): add number of stops to primary outcomes
old = ("Birincil sonuç ölçütü olarak üst ekstremite hareket kontrolünü değerlendiren Hız Tepe Sayısı "
       "(Number of Velocity Peaks, NVP), Yol Doğrusallığı (straightness) ve Duraklama Süresi (pause time) "
       "parametreleri belirlenmiştir.")
new = ("Birincil sonuç ölçütü olarak üst ekstremite hareket kontrolünü değerlendiren Hız Tepe Sayısı "
       "(Number of Velocity Peaks, NVP), Yol Doğrusallığı (straightness), Duraklama Süresi (pause time) "
       "ve Durak Sayısı (number of stops) parametreleri belirlenmiştir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 50 (revision outcomes)")
        break

# 3. Task heading (para 105)
old = "Görevin Yürütülmesi (Reach & Return — Ulaşma–Dönüş):"
new = "Görevin Yürütülmesi (Reach-to-Grasp — Ulaşma-Kavrama):"
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 105 (task heading)")
        break

# 4. Task execution paragraph (para 107): rename + add cup setup
old = ("Ulaşma–Dönüş görevi için sözlü talimat, dikkatte yanlılığı önlemek amacıyla aynı standart metinden "
       "verilecektir. Katılımcıların her değerlendirme zaman noktasında (önce ve sonra) kaydedilmek üzere üç "
       "deneme tamamlamaları gerekmektedir; analiz için üç denemenin ortalaması kullanılacaktır. Yerleşik "
       "hareket davranışlarını yakalamak için \"hızlı\" veya \"yavaş\" gitmeleri yönünde açık bir talimat "
       "verilmeksizin, görevi kendileri için doğal olan bir hızda (rahat hareket hızı) tamamlamaları söylenecektir.")
new = ("Ulaşma-kavrama görevi için sözlü talimat, dikkatte yanlılığı önlemek amacıyla aynı standart metinden "
       "verilecektir. Katılımcıların her değerlendirme zaman noktasında (önce ve sonra) kaydedilmek üzere üç "
       "deneme tamamlamaları gerekmektedir; analiz için üç denemenin ortalaması kullanılacaktır. Yerleşik "
       "hareket davranışlarını yakalamak için \"hızlı\" veya \"yavaş\" gitmeleri yönünde açık bir talimat "
       "verilmeksizin, görevi kendileri için doğal olan bir hızda (rahat hareket hızı) tamamlamaları söylenecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 107 (task execution)")
        break

# 4b. Insert new paragraph about cup placement after task execution
for p in doc.paragraphs:
    if "görevi kendileri için doğal olan bir hızda" in p.text:
        insert_after_paragraph(
            doc, p,
            "Görevde kullanılan hedef nesne (fincan/bardak), katılımcının etkilenmemiş üst ekstremitesiyle "
            "gövde hareketi olmadan ulaşabildiği maksimum noktanın %80–90'ında, omuz hizasında ve ulaşılabilir "
            "masa yüzeyinin ön kenarına yerleştirilecektir. Hedef mesafe, her katılımcı için sağlıklı kolun "
            "aktif hareket aralığına göre bireysel olarak belirlenecektir.",
            red=True
        )
        print("OK: inserted cup placement paragraph")
        break

# 5. Primary outcome description paragraph (para 108): add number of stops
old = ("Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket kontrolündeki değişimdir (NVP, straightness, pause time). "
       "MediaPipe tabanlı pipeline akıllı telefon veya web kamerası videosundan avuç yörüngesi türetir (8). "
       "Hız Tepe Sayısı (NVP): Avuç içi merkezinin teğetsel hız profilindeki yerel maksimum sayısıdır; daha yüksek NVP değerleri hareketin daha fazla alt bölüme ayrıldığını ve daha kesintili olduğunu gösterir. "
       "Yol Doğrusallığı (straightness): Hedefe ulaşma yörüngesinin ideal düz çizgiye olan benzerliğidir; 1'e yakın değerler daha doğrusal, daha düşük değerler daha sapmalı yolları ifade eder. "
       "Duraklama Süresi (pause time): Hareket penceresi içinde hızın belirlenen eşik değerin altında kaldığı toplam süredir; daha uzun duraklama süreleri daha fazla hareket kesintisine işaret eder.")
new = ("Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket kontrolündeki değişimdir (NVP, straightness, pause time, number of stops). "
       "MediaPipe tabanlı pipeline akıllı telefon veya web kamerası videosundan avuç yörüngesi türetir (8). "
       "Hız Tepe Sayısı (NVP): Avuç içi merkezinin teğetsel hız profilindeki yerel maksimum sayısıdır; daha yüksek NVP değerleri hareketin daha fazla alt bölüme ayrıldığını ve daha kesintili olduğunu gösterir. "
       "Yol Doğrusallığı (straightness): Hedefe ulaşma yörüngesinin ideal düz çizgiye olan benzerliğidir; 1'e yakın değerler daha doğrusal, daha düşük değerler daha sapmalı yolları ifade eder. "
       "Duraklama Süresi (pause time): Hareket penceresi içinde hızın belirlenen eşik değerin altında kaldığı toplam süredir; daha uzun duraklama süreleri daha fazla hareket kesintisine işaret eder. "
       "Durak Sayısı (number of stops): Hareket penceresi içinde hızın belirlenen eşik değerin altına düşüp belirli bir süre (ör. 100 ms) altında kaldığı bağımsız durak sayısıdır; daha fazla durak daha kesintili, kontrolsüz bir hareketi yansıtır.")
for p in doc.paragraphs:
    if "Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket kontrolündeki değişimdir" in p.text:
        replace_paragraph_text(p, new, red=True)
        print("OK: para 108 (primary outcome description)")
        break

# 6. Primary outcome intro paragraph (para 111): add number of stops
old = ("Çalışmanın birincil sonuç ölçütü, etkilenen üst ekstremitenin hareket kontrolündeki değişim olacaktır. "
       "Hareket kontrolü, Hız Tepe Sayısı (Number of Velocity Peaks, NVP), Yol Doğrusallığı (straightness) ve "
       "Duraklama Süresi (pause time) parametreleri kullanılarak değerlendirilecektir. "
       "Bu amaçla, MediaPipe tabanlı işaretsiz hareket analizi sistemi aracılığıyla akıllı telefon veya web kamerası "
       "ile kaydedilen videolardan avuç içi merkezinin hareket yörüngesi elde edilecektir (8).")
new = ("Çalışmanın birincil sonuç ölçütü, etkilenen üst ekstremitenin hareket kontrolündeki değişim olacaktır. "
       "Hareket kontrolü, Hız Tepe Sayısı (Number of Velocity Peaks, NVP), Yol Doğrusallığı (straightness), "
       "Duraklama Süresi (pause time) ve Durak Sayısı (number of stops) parametreleri kullanılarak değerlendirilecektir. "
       "Bu amaçla, MediaPipe tabanlı işaretsiz hareket analizi sistemi aracılığıyla akıllı telefon veya web kamerası "
       "ile kaydedilen videolardan avuç içi merkezinin hareket yörüngesi elde edilecektir (8).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 111 (primary outcome intro)")
        break

# 7. Movement Time description (para 129): keep as wrist-based instead of palm if needed
old = ("Ölçüm Yöntemi: Hareket süresi, avuç içi teğetsel hızının önceden belirlenen eşik değerin üzerinde kaldığı "
       "zaman aralığı (başlangıç–bitiş) olarak tanımlanacak ve saniye cinsinden hesaplanacaktır. Daha kısa süreler daha "
       "hızlı hareket performansını gösterecektir.")
new = ("Ölçüm Yöntemi: Hareket süresi, avuç içi merkezinin teğetsel hızının önceden belirlenen eşik değerin üzerinde kaldığı "
       "zaman aralığı (başlangıç–bitiş) olarak tanımlanacak ve saniye cinsinden hesaplanacaktır. Daha kısa süreler daha "
       "hızlı hareket performansını gösterecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 129 (movement time)")
        break

# 8. Pre-assessment paragraph (para 163): rename task, add number of stops
old = ("Sistem kalibrasyonunun ardından katılımcılardan Reach & Return görevini üç kez gerçekleştirmeleri istenecektir. "
       "Hareketler, MediaPipe tabanlı işaretsiz hareket yakalama sistemi ile kaydedilecek ve hareket kontrolü "
       "(NVP, straightness, pause time), gövde kompanzasyonu (trunk ratio) ve diğer kinematik parametreler analiz edilecektir.")
new = ("Sistem kalibrasyonunun ardından katılımcılardan Reach-to-Grasp (ulaşma-kavrama) görevini üç kez gerçekleştirmeleri istenecektir. "
       "Hareketler, MediaPipe tabanlı işaretsiz hareket yakalama sistemi ile kaydedilecek ve hareket kontrolü "
       "(NVP, straightness, pause time, number of stops), gövde kompanzasyonu (trunk ratio) ve diğer kinematik parametreler analiz edilecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 163 (pre-assessment)")
        break

# 9. PETTLEP intervention description (para 171): rename task
old = ("Hazırlık aşamasında katılımcılardan etkilenmemiş üst ekstremiteleri ile bir kez yavaşça uzanma hareketi yapmaları ve "
       "hareket sırasında oluşan hissi fark ederek etkilenen üst ekstremitelerine aktarmaları istenecektir. "
       "Eylem gözlemi aşamasında katılımcılar, etkilenen üst ekstremite ile gerçekleştirilen Reach & Return hareketinin "
       "birinci şahıs bakış açısından kaydedilmiş videosunu izleyecektir. Videolarda gövde stabilitesi, omuz kontrolü, "
       "dirsek hareketi ve ulaşma–geri dönüş sekansı vurgulanacaktır. Motor imgeleme aşamasında katılımcılar gözleri kapalı şekilde, "
       "hareketi birinci şahıs ve kinestetik perspektiften, gerçek zamanlı olarak zihinsel olarak canlandıracaktır. "
       "Son iki blokta ulaşma, kısa bekleme ve geri dönüş fazlarının zamanlamasını desteklemek amacıyla sesli zamanlama ipuçları kullanılacaktır. "
       "Dinlenme aşamasında katılımcılar herhangi bir görev yerine getirmeden rahat şekilde oturacak ve normal solunumlarını sürdürecektir.")
new = ("Hazırlık aşamasında katılımcılardan etkilenmemiş üst ekstremiteleri ile bir kez yavaşça hedefe uzanma ve kavrama hareketi yapmaları ve "
       "hareket sırasında oluşan hissi fark ederek etkilenen üst ekstremitelerine aktarmaları istenecektir. "
       "Eylem gözlemi aşamasında katılımcılar, etkilenen üst ekstremite ile gerçekleştirilen Reach-to-Grasp (ulaşma-kavrama) hareketinin "
       "birinci şahıs bakış açısından kaydedilmiş videosunu izleyecektir. Videolarda gövde stabilitesi, omuz kontrolü, "
       "dirsek hareketi, hedefe ulaşma ve kavrama sekansı vurgulanacaktır. Motor imgeleme aşamasında katılımcılar gözleri kapalı şekilde, "
       "hareketi birinci şahıs ve kinestetik perspektiften, gerçek zamanlı olarak zihinsel olarak canlandıracaktır. "
       "Son iki blokta hedefe ulaşma, kavrama, kısa bekleme ve geri dönüş fazlarının zamanlamasını desteklemek amacıyla sesli zamanlama ipuçları kullanılacaktır. "
       "Dinlenme aşamasında katılımcılar herhangi bir görev yerine getirmeden rahat şekilde oturacak ve normal solunumlarını sürdürecektir.")
for p in doc.paragraphs:
    if "Hazırlık aşamasında katılımcılardan etkilenmemiş" in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 171 (AO/MI description)")
        break

# 10. PETTLEP paragraph (para 180): rename task
old = ("Katılımcılar değerlendirme ile aynı oturma pozisyonunda bulunacak, aynı çevresel koşullarda çalışacak ve "
       "günlük yaşamda sık kullanılan Reach & Return görevini zihinsel olarak canlandıracaktır.")
new = ("Katılımcılar değerlendirme ile aynı oturma pozisyonunda bulunacak, aynı çevresel koşullarda çalışacak ve "
       "günlük yaşamda sık kullanılan Reach-to-Grasp (ulaşma-kavrama) görevini zihinsel olarak canlandıracaktır.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 180 (PETTLEP task)")
        break

# 11. Intervention protocol paragraph (para 160): rename task
old = ("Müdahale sırasında herhangi bir fiziksel uygulama yaptırılmayacak, Reach & Return (uzanma ve geri dönüş) "
       "görevi yalnızca ön ve son değerlendirmelerde gerçekleştirilecektir.")
new = ("Müdahale sırasında herhangi bir fiziksel uygulama yaptırılmayacak, Reach-to-Grasp (ulaşma-kavrama) "
       "görevi yalnızca ön ve son değerlendirmelerde gerçekleştirilecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 160 (intervention protocol)")
        break

# 12. Control group paragraph (para 178): rename task
old = "Katılımcılardan üst ekstremite hareketlerini veya Reach & Return görevini zihinsel olarak canlandırmamaları özellikle talep edilecektir."
new = "Katılımcılardan üst ekstremite hareketlerini veya Reach-to-Grasp (ulaşma-kavrama) görevini zihinsel olarak canlandırmamaları özellikle talep edilecektir."
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 178 (control group)")
        break

# 13. Statistics correlation paragraph (para 214): add ΔNumber of Stops
old = ("KVIQ-10 puanları ile kinematik değişkenlerdeki değişim skorları (ΔNVP, ΔStraightness, ΔPause Time, "
       "ΔTrunk Ratio vb.) arasındaki ilişkiler Pearson veya Spearman korelasyon analizleri kullanılarak değerlendirilecektir.")
new = ("KVIQ-10 puanları ile kinematik değişkenlerdeki değişim skorları (ΔNVP, ΔStraightness, ΔPause Time, "
       "ΔNumber of Stops, ΔTrunk Ratio vb.) arasındaki ilişkiler Pearson veya Spearman korelasyon analizleri kullanılarak değerlendirilecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 214 (correlations)")
        break

# 14. Revision summary paragraph (para 49): rename task
old = ("Çalışmada kullanılan motor görev \"Reach & Return (uzanma ve geri dönüş)\" olarak belirlenmiştir. "
       "Deneysel müdahale, PETTLEP temelli Eylem Gözlemi ve Motor İmgeleme (AOMI) protokolü kapsamında dört adet 3 dakikalık bloktan oluşacak "
       "şekilde (toplam yaklaşık 13 dakika) yapılandırılmıştır. Kontrol koşulu ise süre ve dikkat yükü açısından eşleştirilmiş imgeleme ve "
       "zihinsel arınma protokolü olarak düzenlenmiştir.")
new = ("Çalışmada kullanılan motor görev \"Reach-to-Grasp (ulaşma-kavrama)\" olarak belirlenmiştir. "
       "Deneysel müdahale, PETTLEP temelli Eylem Gözlemi ve Motor İmgeleme (AOMI) protokolü kapsamında dört adet 3 dakikalık bloktan oluşacak "
       "şekilde (toplam yaklaşık 13 dakika) yapılandırılmıştır. Kontrol koşulu ise süre ve dikkat yükü açısından eşleştirilmiş imgeleme ve "
       "zihinsel arınma protokolü olarak düzenlenmiştir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_runs(p, old, new)
        print("OK: para 49 (revision summary)")
        break

# Save
doc.save(str(out))
print(f"\nSaved updated document to: {out}")
print(f"Backup saved to: {backup}")
