from docx import Document

doc = Document(r"D:\Thesis app\phyphox\ethics commitee\Ethics BKK Last version  (AutoRecovered).docx")
with open(r"C:\Users\acer\AppData\Local\Temp\opencode\docx_paragraphs.txt", "w", encoding="utf-8") as f:
    f.write(f"Number of paragraphs: {len(doc.paragraphs)}\n\n")
    for i, p in enumerate(doc.paragraphs):
        txt = p.text
        if txt.strip():
            f.write(f"[{i}]\n{txt}\n\n")
print("Extracted paragraphs")
