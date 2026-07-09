from docx import Document
from pathlib import Path

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_BACKUP.docx")
doc = Document(str(p))
out = open(r"C:\Users\acer\AppData\Local\Temp\opencode\proposal_paragraphs.txt", "w", encoding="utf-8")
out.write(f"paragraphs: {len(doc.paragraphs)}\n\n")
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t:
        out.write(f"[{i}]\n{t}\n\n")
out.close()
print("done")
