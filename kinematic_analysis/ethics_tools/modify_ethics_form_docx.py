# -*- coding: utf-8 -*-
"""
Modify the ethics-committee Word form to assign variables to tools:
  - Phyphox (wrist smartphone accelerometer): SPARC, Movement Time, Peak Velocity
  - MediaPipe (camera): Trunk Ratio, Shoulder Elevation, Elbow Angle
All changed/new text is coloured red.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

original = Path(r"D:\Thesis app\phyphox\ethics commitee\Ethics BKK Last version  (AutoRecovered).docx")
backup = original.with_stem(original.stem + "_backup")
modified = original.with_stem(original.stem + "_modified")

# Backup original
import shutil
shutil.copy2(original, backup)

doc = Document(str(original))

RED = RGBColor(0xFF, 0x00, 0x00)


def replace_paragraph_text(p, new_text, red=True):
    """Clear paragraph runs and insert new_text (optionally in red)."""
    # Clear existing runs
    for run in p.runs:
        run._element.getparent().remove(run._element)
    run = p.add_run(new_text)
    if red:
        run.font.color.rgb = RED
    return run


def replace_in_paragraph(p, old_text, new_text):
    """
    Replace a substring inside a paragraph, colouring only the inserted
    new text red while keeping surrounding text black.
    Falls back to full-paragraph red replacement if the substring is not found.
    """
    full = p.text
    idx = full.find(old_text)
    if idx == -1:
        replace_paragraph_text(p, full.replace(old_text, new_text), red=True)
        return
    prefix = full[:idx]
    suffix = full[idx + len(old_text):]
    # Remove all runs
    for run in p.runs:
        run._element.getparent().remove(run._element)
    if prefix:
        p.add_run(prefix)
    r = p.add_run(new_text)
    r.font.color.rgb = RED
    if suffix:
        p.add_run(suffix)


# ------------------------------------------------------------------
# 1. Justification / aim paragraph (para [42])
# ------------------------------------------------------------------
old = ("Bu çalışmada, PETTLEP temelli AOMI uygulamasının üst ekstremite hareket kalitesi üzerindeki "
       "kısa dönemli etkileri, MediaPipe Pose Landmarker tabanlı işaretsiz hareket analizi sistemi "
       "kullanılarak değerlendirilecektir. Hareket pürüzsüzlüğü (SPARC), gövde kompanzasyonu (trunk ratio), "
       "omuz kuşağı elevasyonu (shoulder elevation), dirsek açısı, hareket süresi ve tepe hız gibi "
       "kinematik parametreler objektif olarak ölçülecektir (8).")
new = ("Bu çalışmada, PETTLEP temelli AOMI uygulamasının üst ekstremite hareket kalitesi üzerindeki "
       "kısa dönemli etkileri, MediaPipe Pose Landmarker tabanlı işaretsiz hareket analizi sistemi ve "
       "Phyphox akıllı telefon akselerometrisi kullanılarak değerlendirilecektir. Hareket pürüzsüzlüğü (SPARC), "
       "hareket süresi ve tepe hız gibi zamana/ivmeye bağlı kinematik parametreler Phyphox ile; "
       "gövde kompanzasyonu (trunk ratio), omuz kuşağı elevasyonu (shoulder elevation) ve dirsek açısı "
       "gibi uzamsal/açısal parametreler MediaPipe ile objektif olarak ölçülecektir (8,19,20).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 1 (justification)")
        break

# ------------------------------------------------------------------
# 2. Revision paragraph listing outcomes (para [57])
# ------------------------------------------------------------------
old = ("Birincil sonuç ölçütü olarak hareket pürüzsüzlüğünü değerlendiren Spectral Arc Length (SPARC) "
       "parametresi belirlenmiştir. İkincil sonuç ölçütleri arasında gövde kompanzasyonu (trunk_ratio), "
       "omuz kuşağı elevasyonu (shoulder_vert_norm; ZoeDepth tabanlı metrik ölçekleme ile), dirsek açısı "
       "(elbow_angle_mean), hareket süresi (movement_time_sec) ve tepe hız (peak_velocity_px_s) yer almaktadır.")
new = ("Birincil sonuç ölçütü olarak hareket pürüzsüzlüğünü değerlendiren Spectral Arc Length (SPARC) "
       "parametresi Phyphox akselerometre kaydı üzerinden belirlenmiştir. İkincil sonuç ölçütleri arasında; "
       "MediaPipe ile ölçülen gövde kompanzasyonu (trunk_ratio), omuz kuşağı elevasyonu (shoulder_vert_norm; "
       "ZoeDepth tabanlı metrik ölçekleme ile) ve dirsek açısı (elbow_angle_mean) ile Phyphox ile ölçülen "
       "hareket süresi (movement_time_sec) ve tepe hız (peak_velocity_px_s) yer almaktadır.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 2 (revision outcomes)")
        break

# ------------------------------------------------------------------
# 3. Data-collection / kinematic parameters paragraph (para [109])
# ------------------------------------------------------------------
old = ("Kinematik parametreler, MediaPipe Pose Landmarker (33 landmark) kullanılarak video kayıtlarından "
       "otomatik olarak çıkarılacaktır. MediaPipe tabanlı işaretsiz hareket analizinin inme sonrası bireylerde "
       "üst ekstremite ulaşma kinematiğini izlemede geçerliliği daha önce gösterilmiştir (8).")
new = ("Kinematik parametreler, iki farklı sistemle eşzamanlı olarak toplanacaktır: Uzamsal ve açısal parametreler "
       "(gövde kompanzasyonu, omuz kuşağı elevasyonu, dirsek açısı) MediaPipe Pose Landmarker (33 landmark) "
       "kullanılarak video kayıtlarından otomatik olarak çıkarılacaktır; zamana bağlı parametreler (SPARC, "
       "hareket süresi, tepe hız) ise etkilenen taraf bileğine sabitlenen akıllı telefonun Phyphox uygulaması "
       "ile kaydedilen akselerometri verisinden hesaplanacaktır. MediaPipe tabanlı işaretsiz hareket analizinin "
       "inme sonrası bireylerde üst ekstremite ulaşma kinematiğini izlemede geçerliliği daha önce gösterilmiştir (8).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 3 (kinematic data collection)")
        break

# ------------------------------------------------------------------
# 4. Task execution paragraph (para [114])
# ------------------------------------------------------------------
old = ("Ulaşma–Dönüş görevi için sözlü talimat, dikkatte yanlılığı önlemek amacıyla aynı standart metinden "
       "verilecektir. Katılımcıların her değerlendirme zaman noktasında (önce ve sonra) kaydedilmek üzere üç "
       "deneme tamamlamaları gerekmektedir; analiz için üç denemenin ortalaması kullanılacaktır. Yerleşik "
       "hareket davranışlarını yakalamak için \"hızlı\" veya \"yavaş\" gitmeleri yönünde açık bir talimat "
       "verilmeksizin, görevi kendileri için doğal olan bir hızda (rahat hareket hızı) tamamlamaları söylenecektir.")
new = ("Ulaşma–Dönüş görevi için sözlü talimat, dikkatte yanlılığı önlemek amacıyla aynı standart metinden "
       "verilecektir. Katılımcıların her değerlendirme zaman noktasında (önce ve sonra) kaydedilmek üzere üç "
       "deneme tamamlamaları gerekmektedir; analiz için üç denemenin ortalaması kullanılacaktır. Her deneme "
       "sırasında etkilenen taraf bileğine Phyphox uygulaması yüklü akıllı telefon sabitlenerek video kaydıyla "
       "eşzamanlı akselerometri verisi toplanacaktır. Yerleşik hareket davranışlarını yakalamak için \"hızlı\" "
       "veya \"yavaş\" gitmeleri yönünde açık bir talimat verilmeksizin, görevi kendileri için doğal olan bir "
       "hızda (rahat hareket hızı) tamamlamaları söylenecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 4 (task execution)")
        break

# ------------------------------------------------------------------
# 5. SPARC description paragraph (para [115])
# ------------------------------------------------------------------
old = ("Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişimdir (SPARC). "
       "MediaPipe tabanlı pipeline akıllı telefon veya web kamerası videosundan avuç yörüngesi türetir (8). "
       "SPARC (Spectral Arc Length — Spektral Ark Uzunluğu): Avuç merkezinin X–Y yörüngesinden teğetsel hız "
       "profili hesaplanır; hız sinyali normalize edilir ve hızlı Fourier dönüşümü (FFT) ile frekans spektrumu "
       "elde edilir (fc ≤ 10 Hz). Spektral eğrinin ark uzunluğu alınır (15,16). İnme sonrası ulaşma görevlerinde "
       "değerlendirilen 32 pürüzsüzlük metriği arasından yalnızca SPARC'nin matematiksel olarak geçerli bulunduğu "
       "gösterilmiştir (17). \nSPARC'nin FM-UE ile uzunlamasına pozitif ilişkisi doğrulanmıştır (18). Daha negatif "
       "SPARC değerleri daha pürüzsüz, kesintisiz hareketi; daha az negatif (veya pozitif) değerler stop-and-go "
       "davranışını yansıtır.")
new = ("Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişimdir (SPARC). "
       "SPARC, Phyphox akıllı telefon akselerometresi ile etkilenen taraf bileğinden kaydedilen üç eksenli ivme "
       "verisinden hesaplanacaktır. Phyphox yaklaşık 100 Hz örnekleme frekansıyla millisaniye hassasiyetinde zaman "
       "damgası sunar ve tepe hız ile hareket süresinin güvenilir hesaplanmasını destekler (19). "
       "SPARC (Spectral Arc Length — Spektral Ark Uzunluğu): Bilek akselerometrisi mutlak hız profili "
       "[√(x²+y²+z²)] üzerinden teğetsel hız hesaplanır; hız sinyali normalize edilir ve hızlı Fourier dönüşümü "
       "(FFT) ile frekans spektrumu elde edilir (fc ≤ 10 Hz). Spektral eğrinin ark uzunluğu alınır (15,16). "
       "İnme sonrası ulaşma görevlerinde değerlendirilen 32 pürüzsüzlük metriği arasından yalnızca SPARC'nin "
       "matematiksel olarak geçerli bulunduğu gösterilmiştir (17).\nSPARC'nin FM-UE ile uzunlamasına pozitif ilişkisi "
       "doğrulanmıştır (18). Daha negatif SPARC değerleri daha pürüzsüz, kesintisiz hareketi; daha az negatif (veya "
       "pozitif) değerler stop-and-go davranışını yansıtır.")
for p in doc.paragraphs:
    if "Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişimdir (SPARC)" in p.text:
        replace_paragraph_text(p, new, red=True)
        print("OK: para 5 (SPARC description)")
        break

# ------------------------------------------------------------------
# 6. Camera/phone setup paragraph (para [116])
# ------------------------------------------------------------------
old = ("Anlık kinematik değişiklikleri değerlendirmek amacıyla, tek eğitim seansının öncesinde ve hemen "
       "sonrasında kayıtlar alınacaktır. Sensör tutarlılığını sağlamak için tüm katılımcılarda aynı akıllı telefon "
       "veya web kamerası modeli kullanılacaktır. Kamera, sabit bir yüksekliğe monte edilen tripod üzerine "
       "yerleştirilecek ve katılımcıya göre yan görünümde (90°) yaklaşık 1,5–2,0 m standart mesafede "
       "konumlandırılacaktır.\nKayıtlar 1080p çözünürlükte ve saniyede 60 kare (fps) yapılacaktır. Yapay zeka "
       "takibini engelleyebilecek arkadan aydınlatma veya gölgeleri önlemek için kayıtlar tutarlı, dağınık "
       "aydınlatmaya sahip bir odada gerçekleştirilecektir. Video verileri MediaPipe Pose Landmarker (33 landmark, "
       "2D) ile işlenecek; omuz elevasyonu için ZoeDepth monoküler derinlik ağı ile omuz genişliği metrik ölçeklemesi "
       "uygulanacaktır.")
new = ("Anlık kinematik değişiklikleri değerlendirmek amacıyla, tek eğitim seansının öncesinde ve hemen "
       "sonrasında kayıtlar alınacaktır. Video kayıtları için aynı akıllı telefon veya web kamerası modeli "
       "kullanılacaktır. Kamera, sabit bir yüksekliğe monte edilen tripod üzerine yerleştirilecek ve katılımcıya "
       "göre yan görünümde (90°) yaklaşık 1,5–2,0 m standart mesafede konumlandırılacaktır. Kayıtlar 1080p "
       "çözünürlükte ve saniyede 60 kare (fps) yapılacaktır. Yapay zeka takibini engelleyebilecek arkadan "
       "aydınlatma veya gölgeleri önlemek için kayıtlar tutarlı, dağınık aydınlatmaya sahip bir odada "
       "gerçekleştirilecektir. Video verileri MediaPipe Pose Landmarker (33 landmark, 2D) ile işlenecek; omuz "
       "elevasyonu için ZoeDepth monoküler derinlik ağı ile omuz genişliği metrik ölçeklemesi uygulanacaktır.\n"
       "Akselerometrik kayıt için etkilenen taraf bileğine standart bir akıllı telefon Phyphox uygulaması "
       "(RWTH Aachen University) aracılığıyla sabitlenecektir. Linear acceleration sensöründen x, y, z eksenlerindeki "
       "ivmeler yaklaşık 100 Hz frekansla kaydedilecek ve CSV formatında dışa aktarılarak Python tabanlı kod ile "
       "SPARC, hareket süresi ve tepe hız hesaplanacaktır.")
for p in doc.paragraphs:
    if "Anlık kinematik değişiklikleri değerlendirmek amacıyla" in p.text:
        replace_paragraph_text(p, new, red=True)
        print("OK: para 6 (camera/phone setup)")
        break

# ------------------------------------------------------------------
# 7. Primary outcome paragraphs ([117]-[120])
# ------------------------------------------------------------------
# 7a
old = ("Çalışmanın birincil sonuç ölçütü, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişim "
       "olacaktır. Hareket pürüzsüzlüğü, Spektral Ark Uzunluğu (Spectral Arc Length, SPARC) yöntemi kullanılarak "
       "değerlendirilecektir. Bu amaçla, MediaPipe tabanlı işaretsiz hareket analizi sistemi aracılığıyla akıllı "
       "telefon veya web kamerası ile kaydedilen videolardan avuç içi merkezinin hareket yörüngesi elde edilecektir (8).")
new = ("Çalışmanın birincil sonuç ölçütü, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişim "
       "olacaktır. Hareket pürüzsüzlüğü, Spektral Ark Uzunluğu (Spectral Arc Length, SPARC) yöntemi kullanılarak "
       "değerlendirilecektir. Bu amaçla, etkilenen taraf bileğine sabitlenen akıllı telefonun Phyphox uygulaması "
       "ile kaydedilen üç eksenli akselerometri verisi kullanılacaktır (19).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 7a (primary outcome intro)")
        break

# 7b
old = ("SPARC hesaplamasında, avuç içi merkezinin iki boyutlu (X–Y) hareket yörüngesinden teğetsel hız profili "
       "oluşturulacaktır. Hız sinyali normalize edildikten sonra hızlı Fourier dönüşümü (Fast Fourier Transform, FFT) "
       "uygulanarak frekans spektrumu elde edilecek ve 10 Hz'e kadar olan frekans bileşenleri analiz edilecektir. "
       "Ardından spektral eğrinin ark uzunluğu hesaplanarak SPARC değeri belirlenecektir (Balasubramanian ve ark., 2012; 2015).")
new = ("SPARC hesaplamasında, bilek akselerometrisi mutlak hız profili [√(x²+y²+z²)] üzerinden teğetsel hız profili "
       "oluşturulacaktır. Hız sinyali normalize edildikten sonra hızlı Fourier dönüşümü (Fast Fourier Transform, FFT) "
       "uygulanarak frekans spektrumu elde edilecek ve 10 Hz'e kadar olan frekans bileşenleri analiz edilecektir. "
       "Ardından spektral eğrinin ark uzunluğu hesaplanarak SPARC değeri belirlenecektir (Balasubramanian ve ark., 2012; 2015).")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 7b (SPARC computation)")
        break

# ------------------------------------------------------------------
# 8. Secondary outcomes: assign tool per variable
# ------------------------------------------------------------------
# 8a Trunk Ratio
old = ("Ölçüm Yöntemi: Gövde kompanzasyonu, hareket sırasında gövde ve el hareketlerinin göreceli katkısını "
       "yansıtan Trunk Ratio parametresi kullanılarak hesaplanacaktır. Daha yüksek değerler, görevin gerçekleştirilmesi "
       "sırasında gövde kompanzasyonunun daha fazla kullanıldığını gösterecektir.")
new = ("Ölçüm Yöntemi: Gövde kompanzasyonu, MediaPipe ile belirlenen gövde ve el hareketlerinin göreceli katkısını "
       "yansıtan Trunk Ratio parametresi kullanılarak hesaplanacaktır. Daha yüksek değerler, görevin gerçekleştirilmesi "
       "sırasında gövde kompanzasyonunun daha fazla kullanıldığını gösterecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 8a (trunk ratio)")
        break

# 8b Shoulder Elevation
old = ("Ölçüm Yöntemi: Omuz elevasyonu, etkilenen omuz landmark’ının dikey (Y ekseni) yer değiştirmesinin omuz "
       "genişliğine normalize edilmesiyle hesaplanacaktır (shoulder_vert_norm). Omuz genişliği için metrik ölçek, "
       "ZoeDepth monoküler derinlik ağı kullanılarak elde edilecektir. Daha yüksek değerler daha fazla telafi edici "
       "omuz elevasyonunu gösterecektir.")
new = ("Ölçüm Yöntemi: Omuz elevasyonu, MediaPipe ile belirlenen etkilenen omuz landmark’ının dikey (Y ekseni) yer "
       "değiştirmesinin omuz genişliğine normalize edilmesiyle hesaplanacaktır (shoulder_vert_norm). Omuz genişliği için "
       "metrik ölçek, ZoeDepth monoküler derinlik ağı kullanılarak elde edilecektir. Daha yüksek değerler daha fazla "
       "telafi edici omuz elevasyonunu gösterecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 8b (shoulder elevation)")
        break

# 8c Elbow Angle
old = ("Ölçüm Yöntemi: Dirsek açısı, omuz–dirsek–el bileği noktaları arasında oluşturulan vektörler arasındaki açının "
       "hareket süresince her karede hesaplanması ve ortalama değerinin alınmasıyla belirlenecektir. Ölçüm özellikle "
       "sagittal düzlemden elde edilen görüntülerde yüksek doğruluk sağlamaktadır.")
new = ("Ölçüm Yöntemi: Dirsek açısı, MediaPipe ile belirlenen omuz–dirsek–el bileği noktaları arasında oluşturulan "
       "vektörler arasındaki açının hareket süresince her karede hesaplanması ve ortalama değerinin alınmasıyla "
       "belirlenecektir. Ölçüm özellikle sagittal düzlemden elde edilen görüntülerde yüksek doğruluk sağlamaktadır.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 8c (elbow angle)")
        break

# 8d Movement Time
old = ("Ölçüm Yöntemi: Hareket süresi, avuç içi teğetsel hızının önceden belirlenen eşik değerin üzerinde kaldığı "
       "zaman aralığı (başlangıç–bitiş) olarak tanımlanacak ve saniye cinsinden hesaplanacaktır. Daha kısa süreler daha "
       "hızlı hareket performansını gösterecektir.")
new = ("Ölçüm Yöntemi: Hareket süresi, Phyphox ile kaydedilen bilek akselerometrisi mutlak hız profilinin önceden "
       "belirlenen eşik değerin üzerinde kaldığı zaman aralığı (başlangıç–bitiş) olarak tanımlanacak ve milisaniye (ms) "
       "hassasiyetinde hesaplanacaktır. Daha kısa süreler daha hızlı hareket performansını gösterecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 8d (movement time)")
        break

# 8e Peak Velocity
old = ("Ölçüm Yöntemi: Tepe hız, hareket penceresi içerisinde avuç içi teğetsel hızının ulaştığı en yüksek değer "
       "olarak hesaplanacak ve piksel/saniye (px/s) cinsinden raporlanacaktır. Daha yüksek değerler daha hızlı hareket "
       "üretimini yansıtmaktadır.")
new = ("Ölçüm Yöntemi: Tepe hız, Phyphox ile kaydedilen bilek akselerometrisi mutlak hız profilinin hareket penceresi "
       "içerisinde ulaştığı en yüksek değer olarak hesaplanacak ve metre/saniye (m/s) cinsinden raporlanacaktır. Daha yüksek "
       "değerler daha hızlı hareket üretimini yansıtmaktadır.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 8e (peak velocity)")
        break

# ------------------------------------------------------------------
# 9. Data collection methods paragraph ([190])
# ------------------------------------------------------------------
old = ("Veriler görüntü kaydı, klinik ölçekler, performans testleri ve yapay zekâ tabanlı hareket analizi sistemi "
       "(MediaPipe) kullanılarak toplanacaktır. Katılımcı bilgileri araştırma kodları ile kodlanacak ve analizler "
       "kimliksizleştirilmiş veriler üzerinden gerçekleştirilecektir.")
new = ("Veriler görüntü kaydı, klinik ölçekler, performans testleri, yapay zekâ tabanlı hareket analizi sistemi "
       "(MediaPipe) ve akıllı telefon akselerometrisi (Phyphox) kullanılarak toplanacaktır. Katılımcı bilgileri "
       "araştırma kodları ile kodlanacak ve analizler kimliksizleştirilmiş veriler üzerinden gerçekleştirilecektir.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 9 (data collection methods)")
        break

# ------------------------------------------------------------------
# 10. Pre-assessment paragraph ([170])
# ------------------------------------------------------------------
old = ("Sistem kalibrasyonunun ardından katılımcılardan Reach \u0026 Return görevini üç kez gerçekleştirmeleri "
       "istenecektir. Hareketler, MediaPipe tabanlı işaretsiz hareket yakalama sistemi ile kaydedilecek ve hareket "
       "pürüzsüzlüğü (SPARC), gövde kompanzasyonu (trunk ratio) ve diğer kinematik parametreler analiz edilecektir.")
new = ("Sistem kalibrasyonunun ardından katılımcılardan Reach \u0026 Return görevini üç kez gerçekleştirmeleri "
       "istenecektir. Hareketler, MediaPipe tabanlı işaretsiz hareket yakalama sistemi ile video olarak kaydedilecek ve "
       "gövde kompanzasyonu (trunk ratio), omuz kuşağı elevasyonu ile dirsek açısı analiz edilecektir. Aynı denemeler "
       "sırasında etkilenen taraf bileğine sabitlenen Phyphox ile akselerometri verisi kaydedilerek SPARC, hareket süresi "
       "ve tepe hız hesaplanacaktır.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 10 (pre-assessment)")
        break

# ------------------------------------------------------------------
# 11. Add references 19-21 after reference 18
# ------------------------------------------------------------------
ref18_old = ("18. Saes M, Refai MIM, van Kordelaar J, Scheltinga BL, van Beijnum B-JF, Bussmann JB, et al. "
             "Smoothness metric during reach-to-grasp after stroke. Part 2. Longitudinal association with motor "
             "impairment. J Neuroeng Rehabil. 2021;18(1):144.")
new_refs = [
    "19. Staacks S, Hütz S, Heinke H, Stampfer C. Advanced tools for smartphone-based experiments: phyphox. Phys Teach. 2021;59(3):214-215.",
    "20. Dobkin BH, Dorsch A. The promise of mHealth: physical activity, fitness, and stroke. Neurorehabil Neural Repair. 2011;25(8):711-715.",
    "21. Dobkin BH. Wearable motion sensors to continuously measure real-world physical activities. Curr Opin Neurol. 2013;26(6):602-608.",
]
for i, p in enumerate(doc.paragraphs):
    if ref18_old in p.text:
        parent = p._element.getparent()
        idx = list(parent).index(p._element)
        for j, txt in enumerate(new_refs):
            np = doc.add_paragraph(txt)
            np.runs[0].font.color.rgb = RED
            parent.insert(idx + 1 + j, np._element)
        print("OK: references 19-21 added")
        break

# Save modified document
doc.save(str(modified))
print(f"Saved modified document to: {modified}")
print(f"Backup of original saved to: {backup}")
