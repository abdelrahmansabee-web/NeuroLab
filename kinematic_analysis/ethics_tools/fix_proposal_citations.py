# -*- coding: utf-8 -*-
"""
Fix in-text citations and add Balasubramanian SPARC references.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc = Document(str(p))
RED = RGBColor(0xFF, 0x00, 0x00)


def replace_in_paragraph(p, old_sub, new_sub):
    full = p.text
    idx = full.find(old_sub)
    if idx == -1:
        return False
    prefix = full[:idx]
    suffix = full[idx + len(old_sub):]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    if prefix:
        p.add_run(prefix)
    r = p.add_run(new_sub)
    r.font.color.rgb = RED
    if suffix:
        p.add_run(suffix)
    return True


# Fix citations
replacements = [
    ("(8,19,20)", "(8,14,15)"),
    ("(8,19).", "(8,14)."),
    ("(8,19).", "(8,14)."),  # permissions
    ("(15,16).", "(17,18)."),  # SPARC methodology
]

for old_sub, new_sub in replacements:
    for para in doc.paragraphs:
        if old_sub in para.text:
            replace_in_paragraph(para, old_sub, new_sub)
            print(f"OK: replaced {old_sub} with {new_sub}")
            break

# Add Balasubramanian references 17-18 after reference 16
ref16_keyword = "Dobkin BH. Wearable motion sensors"
new_refs = [
    "17. Balasubramanian S, Melendez-Calderon A, Burdet E. A robust and sensitive metric for quantifying movement smoothness. IEEE Trans Biomed Eng. 2012;59(8):2126-2136.",
    "18. Balasubramanian S, Melendez-Calderon A, Roby-Brami A, Burdet E. On the analysis of movement smoothness. J NeuroEngineering Rehabil. 2015;12:112.",
]
for i, para in enumerate(doc.paragraphs):
    if ref16_keyword in para.text:
        parent = para._element.getparent()
        idx = list(parent).index(para._element)
        for j, txt in enumerate(new_refs):
            np = doc.add_paragraph(txt)
            np.runs[0].font.color.rgb = RED
            parent.insert(idx + 1 + j, np._element)
        print("OK: references 17-18 added")
        break

doc.save(str(p))
print(f"Saved to: {p}")
