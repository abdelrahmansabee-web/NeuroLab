from docx import Document
from pathlib import Path

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc = Document(str(p))

out = open(r"C:\Users\acer\AppData\Local\Temp\opencode\proposal_final_check.txt", "w", encoding="utf-8")

# Check remaining OpenCap/OpenSim
out.write("=== Remaining OpenCap/OpenSim mentions ===\n")
for i, para in enumerate(doc.paragraphs):
    if 'OpenCap' in para.text or 'OpenSim' in para.text:
        out.write(f"[{i}] {para.text}\n\n")

out.write("\n=== Key updated paragraphs ===\n")
keywords = [
    "MediaPipe",
    "Phyphox",
    "SPARC",
    "Data Collection Setup",
    "Primary outcome",
    "Baseline Assessment",
    "Shoulder Girdle Elevation",
    "Trunk Displacement",
]
for kw in keywords:
    out.write(f"\n--- {kw} ---\n")
    for p in doc.paragraphs:
        if kw in p.text:
            out.write(p.text[:600] + "\n")
            break

out.close()
print("done")
