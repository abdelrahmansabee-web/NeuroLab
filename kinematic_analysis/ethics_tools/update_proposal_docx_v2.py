# -*- coding: utf-8 -*-
"""
Update the English thesis proposal to reflect the new tool split.
Robust version: uses keyword search and full-paragraph replacement.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
import shutil

src = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_BACKUP.docx")
backup = src.with_stem(src.stem + "_before_tool_update")
modified = src.with_stem(src.stem.replace("_BACKUP", "_updated"))

if not backup.exists():
    shutil.copy2(src, backup)

doc = Document(str(src))
RED = RGBColor(0xFF, 0x00, 0x00)


def norm(s):
    """Normalize whitespace for matching."""
    return " ".join(s.split())


def replace_paragraph_text(p, new_text):
    for run in p.runs:
        run._element.getparent().remove(run._element)
    r = p.add_run(new_text)
    r.font.color.rgb = RED


def find_and_replace(doc, keyword, new_text, by_contains=True):
    """Find paragraph containing keyword and replace entire text."""
    for p in doc.paragraphs:
        if keyword in p.text:
            replace_paragraph_text(p, new_text)
            return True
    return False


def find_and_replace_substring(doc, keyword, old_sub, new_sub):
    """Find paragraph containing keyword and replace old_sub with new_sub (red)."""
    for p in doc.paragraphs:
        if keyword in p.text:
            full = p.text
            idx = full.find(old_sub)
            if idx == -1:
                # try normalized
                nfull = norm(full)
                nold = norm(old_sub)
                nidx = nfull.find(nold)
                if nidx == -1:
                    return False
                # reconstruct positions in original (approximate)
                idx = full.find(old_sub.strip().split()[0])
                if idx == -1:
                    return False
            prefix = full[:idx]
            suffix = full[idx + len(old_sub):]
            for run in p.runs:
                run._element.getparent().remove(run._element)
            if prefix:
                p.add_run(prefix)
            r = p.add_run(new_sub)
            r.font.color.rgb = RED
            if suffix:
                p.add_run(suffix)
            return True
    return False


# 1. Subject/Purpose paragraph
find_and_replace(
    doc,
    "For this purpose, the",
    "For this purpose, a dual-sensor kinematic setup combining MediaPipe Pose Landmarker (markerless video-based motion capture) with Phyphox smartphone accelerometry is introduced. MediaPipe will quantify spatial and angular parameters (trunk displacement, shoulder girdle elevation, elbow angle) from side-view video, while Phyphox will capture time- and acceleration-based parameters (movement smoothness, movement time, peak velocity) from a wrist-mounted smartphone accelerometer. Hence, by objectively quantifying the \"how\" of movement, we can examine the mediator between central neural planning and peripheral motor execution (8,19,20). Furthermore, the combination of this kinematic analysis with the Box and Block Test (BBT) allows us to relate the quality of movement to functional dexterity in everyday life (9)."
)
print("OK: para 1 (subject/purpose)")

# 2. Expectation paragraph
find_and_replace(
    doc,
    "Therefore, the expectation of this study",
    "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature. Using MediaPipe and Phyphox technologies in combination with a customized PETTLEP-based AOMI protocol, this study seeks to investigate whether specific mental practice can lead to measurable short-term changes in movement smoothness, compensatory strategies and functional dexterity in persons post-stroke."
)
print("OK: para 2 (expectation)")

# 3. Keywords
find_and_replace_substring(
    doc,
    "Keywords:",
    "stroke, OpenCap, kinematics",
    "stroke, MediaPipe, Phyphox, kinematics"
)
print("OK: para 3 (keywords)")

# 4. Assessments paragraph
find_and_replace(
    doc,
    "To ensure reliability and consistency, the environmental condition",
    "To ensure reliability and consistency, the environmental condition will be standardized for all evaluations, which are going to be performed in the Neurorehabilitation unit of Istinye University Liv Bahçeşehir Hospital. All participants will be tested at 2 time points—once immediately before and once immediately after the session. This pre-post test design enables the identification of short-term modifications in movement smoothness (Phyphox), trunk displacement (MediaPipe), shoulder girdle elevation (MediaPipe), movement time (Phyphox) and peak velocity (Phyphox) following motor imagery and a control task."
)
print("OK: para 4 (assessments)")

# 5. Description of Task
find_and_replace(
    doc,
    "Description of Task:",
    "Description of Task: To prevent attentional bias, the verbal instruction for the Reach-to-Grasp is verbally presented from the same standardized script. Participants are invited to participate in three trials per each point of assessment (pre and post). During each trial, a side-view video will be recorded with a smartphone/webcam for MediaPipe analysis, and a second smartphone running the Phyphox application will be secured on the affected-side wrist to record tri-axial accelerometry. They will be instructed to reach for the object at a pace that feels natural to them (comfortable speed of movement) and will not receive any specific instructions to go either \"fast\" or \"slow\"; this is to catch their habitual movement behavior."
)
print("OK: para 5 (description of task)")

# 6. Data Collection Setup
find_and_replace(
    doc,
    "Data Collection Setup:",
    "Data Collection Setup: To evaluate immediate kinematic modifications, recordings will be taken prior to and immediately following the single training session. For MediaPipe analysis, a smartphone or webcam will be used to record a side-view (sagittal plane) video of the affected upper limb and trunk. The camera will be mounted on a tripod at a fixed height, positioned perpendicular to the reaching plane at a standardized distance of 1.5 m to 2.0 m. For Phyphox analysis, a standard smartphone will be secured on the affected-side wrist using an armband or Velcro strap; the Phyphox application will record linear acceleration from the built-in tri-axial accelerometer at approximately 100 Hz."
)
print("OK: para 6 (data collection setup)")

# 7. Camera config
find_and_replace(
    doc,
    "The cameras will be configured to record",
    "The video camera will be configured to record at 1080p resolution and 60 frames per second (fps). Recordings will be conducted in a room with consistent, diffuse lighting to avoid backlighting or shadows that could interfere with AI tracking. The same MediaPipe Pose Landmarker model and processing parameters will be used for all participant analyses to ensure data comparability. The Phyphox CSV output will be imported into a Python-based pipeline for SPARC, movement-time and peak-velocity computation."
)
print("OK: para 7 (camera config)")

# 8. Video processing / OpenCap 3D model
find_and_replace(
    doc,
    "The video data is processed by uploading",
    "The video data will be processed locally using the MediaPipe Pose Landmarker (33 landmark, 2D) pipeline to extract joint trajectories, and the smartphone accelerometer data will be processed locally using a Python pipeline. The following specific kinematic measures will be computed according to their functional importance (8,19)."
)
print("OK: para 8 (video processing)")

# 9. Primary outcome
find_and_replace(
    doc,
    "The primary outcome of this study is the change in movement smoothness",
    "·     The primary outcome of this study is the change in movement smoothness of the affected upper limb, as measured by the Spectral Arc Length (SPARC) derived from the Phyphox wrist accelerometer recording. The smartphone, secured on the affected-side wrist, records tri-axial linear acceleration at approximately 100 Hz; the resulting CSV file is imported into a Python pipeline to compute the absolute velocity profile, normalize it, apply Fast Fourier Transform (FFT) and calculate the arc length of the spectral curve (15,16). Due to their functional meaning, the following specific kinematic parameters will be derived (8,19)."
)
print("OK: para 9 (primary outcome)")

# 10. NVP -> SPARC
find_and_replace(
    doc,
    "Movement Smoothness (Number of Velocity Peaks - NVP)",
    "·     Movement Smoothness (Spectral Arc Length - SPARC): Calculated from the wrist accelerometer absolute velocity profile [√(x²+y²+z²)]. SPARC quantifies the smoothness of the speed profile by measuring the arc length of its normalized Fourier spectrum; more negative values indicate smoother, less segmented movement. Increase in SPARC (i.e., less negative) reflects loss of motor control and more stop-and-go behavior. Reduction in SPARC (more negative) will be interpreted as increased movement smoothness."
)
print("OK: para 10 (SPARC description)")

# 11a. Shoulder purpose
find_and_replace_substring(
    doc,
    "Shoulder Girdle Elevation (Compensatory Strategy)",
    "the markerless capture system tracks surface anatomical landmarks",
    "MediaPipe tracks the shoulder landmark from side-view video"
)
print("OK: para 11a (shoulder purpose)")

# 11b. Shoulder method
find_and_replace_substring(
    doc,
    "relative vertical (Y axis) displacement",
    "Calculated by measuring the relative vertical (Y axis) displacement of the Shoulder marker with respect to the Neck marker during movement.",
    "Calculated by measuring the relative vertical (Y axis) displacement of the affected-side Shoulder marker with respect to its resting position during movement. The side-view camera orientation maximizes the accuracy of pure vertical elevation."
)
print("OK: para 11b (shoulder method)")

# 12a. Trunk purpose
find_and_replace_substring(
    doc,
    "Trunk Displacement (Compensatory Strategy)",
    "This variable is expressed not as an absolute biomechanical value, but as an indicator of compensatory strategy quality.",
    "MediaPipe provides the trunk and hand landmarks needed to compute this compensatory strategy index."
)
print("OK: para 12a (trunk purpose)")

# 12b. Trunk method
find_and_replace_substring(
    doc,
    "Lumbar Extension angle",
    "Calculated by measuring the change in the Lumbar Extension angle (lower back extension) in the sagittal plane during the reaching phase, obtained from the OpenSim kinematic model.",
    "Calculated from MediaPipe trunk and hand landmark trajectories as the ratio of trunk displacement to hand displacement during the reaching phase (Trunk Ratio). Higher values indicate greater reliance on trunk compensation."
)
print("OK: para 12b (trunk method)")

# 13. Permissions
find_and_replace(
    doc,
    "The OpenCap system has been made available",
    "MediaPipe Pose Landmarker is an open-source machine-learning pipeline developed by Google, and Phyphox is an open-source smartphone physics-lab application developed by RWTH Aachen University. Both tools are freely available for scientific research; therefore, specific author permission is not required for their use in this study [8,19]."
)
print("OK: para 13 (permissions)")

# 14. Baseline Assessment
find_and_replace(
    doc,
    "1. Baseline Assessment:",
    "1. Baseline Assessment: After the system calibration, each subject will perform 3 runs of a routine Reach-to-Grasp task. Movement kinematics will be recorded simultaneously: MediaPipe will capture side-view video for trunk displacement, shoulder girdle elevation and elbow angle; Phyphox will record wrist accelerometry for movement smoothness (SPARC), movement time and peak velocity (pre-intervention movement quality)."
)
print("OK: para 14 (baseline assessment)")

# 15. Correlations
find_and_replace(
    doc,
    "Correlations Subjective vividness of imagery",
    "Correlations Subjective vividness of imagery (MIQ-3 scores) and the amount of kinematic facilitation (Δ scores of SPARC, trunk displacement, shoulder girdle elevation, elbow angle, movement time and peak velocity) will be correlated with each other using Pearson’s (normal distribution) or Spearman’s (non-normal distribution) to verify whether higher imagery ability can predict the best motor outcome performance."
)
print("OK: para 15 (correlations)")

# 16a. Transport phase
find_and_replace(
    doc,
    "· The Transport Phase",
    "· The Transport Phase (The Reach & The Trunk): Shoulder flexion and elbow extension are distance-related components. Forward leaning trunk compensation and shoulder girdle elevation are in fact very common compensatory mechanisms in this phase of post-stroke patients. To treat this phase, the \"Glued Back\" constraint is introduced as it attempts to recover real limb range-of-motion (ROM) rather than compensatory motion; these motion patterns can be numerically expressed through MediaPipe-derived 'Trunk Displacement' and 'Shoulder Girdle Elevation' variables."
)
print("OK: para 16a (transport phase)")

# 16b. Manipulation phase
find_and_replace(
    doc,
    "· The Manipulation Phase",
    "· The Manipulation Phase (Grasp & Smoothness): Grasping requires wrist extension and hand aperture to be executed in a synchronized manner to match the size and shape of the object. Adding the 'Full Cup' constraint (fiction of liquid) introduces a significant constraint on the smoothness and precision of movement. This was to practice motor control and coordination instead of raw pressure, forcing the nervous system to reduce segmented speed profiles, which will be quantified by the Phyphox-derived SPARC metric."
)
print("OK: para 16b (manipulation phase)")

# 17. Add references 14-17 after reference 13
ref13_keyword = "Uğur Y, Coşkun H"
new_refs = [
    "14. Wagh V, Scott MW, Andrushko JW, Jones CB, Larssen BC, Boyd LA, Kraeutner SN. Using MediaPipe to track upper-limb reaching movements after stroke: a proof-of-principle study. J Neuroeng Rehabil. 2025 Nov 25;22(1):268.",
    "15. Staacks S, Hütz S, Heinke H, Stampfer C. Advanced tools for smartphone-based experiments: phyphox. Phys Teach. 2021;59(3):214-215.",
    "16. Dobkin BH, Dorsch A. The promise of mHealth: physical activity, fitness, and stroke. Neurorehabil Neural Repair. 2011;25(8):711-715.",
    "17. Dobkin BH. Wearable motion sensors to continuously measure real-world physical activities. Curr Opin Neurol. 2013;26(6):602-608.",
]
for i, p in enumerate(doc.paragraphs):
    if ref13_keyword in p.text:
        parent = p._element.getparent()
        idx = list(parent).index(p._element)
        for j, txt in enumerate(new_refs):
            np = doc.add_paragraph(txt)
            np.runs[0].font.color.rgb = RED
            parent.insert(idx + 1 + j, np._element)
        print("OK: references 14-17 added")
        break

# Save
doc.save(str(modified))
print(f"Saved updated proposal to: {modified}")
print(f"Backup saved to: {backup}")
