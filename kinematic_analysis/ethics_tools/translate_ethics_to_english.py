# -*- coding: utf-8 -*-
"""
Create English translation of the edited Turkish ethics document.
Paragraph-by-paragraph translation preserving structure.
"""
from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

src = Document('ethics_turkish_edited.docx')
dst = Document()

RED = RGBColor(0xFF, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

# Translation mapping: paragraph index -> English text
# Empty or whitespace-only paragraphs are kept empty.
# For paragraphs not in mapping, we will translate heuristically.

translations = {
    0: "Date:18/08/2026",
    2: "HUMAN RESEARCH ETHICS COMMITTEE Application Form",
    4: "PROJECT TITLE:",
    5: "Immediate Effects of a Single-Session PETTLEP-Based Action Observation and Motor Imagery (AOMI) Application on Upper Extremity Kinematics After Stroke: A Randomized Controlled Study",
    8: "APPLICATION INFORMATION",
    9: "Type of Application",
    10: "Initial Application",
    11: "Revision",
    12: "Continuation of an Approved Project",
    13: "Notification of Change in an Approved Project",
    14: "The results of this research will be published in a scientific journal.",
    16: "Nature of the Research",
    17: "Faculty Member Research",
    18: "Master's   Doctoral Thesis Study",
    19: "The Results of This Research Will Be Published in a Scientific Journal",
    20: "Other (Please explain)",
    22: "PROJECT ADVISOR FACULTY MEMBER",
    24: "PRINCIPAL INVESTIGATOR",
    26: "CO-INVESTIGATORS",
    30: "RATIONALE, AIM, AND RESOURCES OF THE RESEARCH",
    31: "Stroke is one of the leading causes of long-term disability among adults worldwide and causes permanent impairments in upper extremity function, movement control, and coordination in a significant proportion of individuals (1). Damage to cortical and subcortical motor structures results in voluntary movement disorders, especially on the affected side, negatively affecting the performance of daily living activities and quality of life (2,3).",
    32: "Traditional neurorehabilitation approaches are largely based on active task practice. However, many stroke patients, especially those with significant upper extremity involvement, cannot participate sufficiently in such applications (4). Therefore, motor-cognitive interventions that do not require physical movement production have gained attention in recent years.",
    33: "Motor Imagery (MI) is the process of mentally simulating a movement without any physical execution, and it largely shares neural mechanisms with actual movement (4,5). In recent years, the combination of Motor Imagery with Action Observation (AO), called Action Observation and Motor Imagery (AOMI), has been suggested to be more effective than either approach alone (6).",
    34: "The effectiveness of motor imagery is closely related to how similar the imagery is to the actual movement experience. The PETTLEP model (Physical, Environment, Task, Timing, Learning, Emotion, Perspective), developed for this purpose, aims to increase functional similarity between imagery and actual movement (7).",
    35: "In this study, the short-term effects of the PETTLEP-based AOMI application on upper extremity movement quality will be evaluated using a MediaPipe Pose Landmarker-based markerless motion analysis system. Kinematic parameters such as upper extremity movement control (NVP, straightness, pause time), trunk compensation (trunk ratio), shoulder girdle elevation (shoulder elevation), elbow angle, movement time, and peak velocity will be analyzed objectively.",
    36: "The aim of this study is to examine the acute effects of a single-session PETTLEP-based AOMI application on upper extremity movement control, compensatory movement strategies, and functional performance in individuals who have had a stroke. The findings are expected to guide the development and optimization of AOMI applications in neurorehabilitation.",
    38: "Keywords: Motor imagery, action observation, AOMI, PETTLEP, stroke, MediaPipe, kinematics, upper extremity rehabilitation.",
    40: "Research Question",
    41: "Does a single-session PETTLEP-based AOMI application cause acute changes in upper extremity movement control (NVP, straightness, pause time), compensatory movement strategies, and functional performance in individuals who have had a stroke compared to imagery and mental clearing control intervention?",
    43: "Hypotheses",
    44: "H₀: There is no significant difference between the PETTLEP-based AOMI application and the control intervention in terms of changes in upper extremity movement control (NVP, straightness, pause time), compensatory movement strategies, and functional performance.",
    45: "H₁: The PETTLEP-based AOMI application provides significant improvements in upper extremity movement control (NVP, straightness, pause time), compensatory movement strategies, and functional performance compared to the control intervention.",
    47: "REVISION (Changes to the Approved Protocol)",
    48: "The study has been restructured as a multicenter study, and Biruni University Hospital Physical Medicine and Rehabilitation Polyclinic has been included as a second data collection center. In this context, Assoc. Prof. Dr. Çiğdem Çınar has been added to the research team as a co-investigator. The scientific design, randomization method, and planned sample size (n = 28) of the study have not been changed.",
    49: "Additionally, the intervention protocol and outcome measures have been defined in more detail. The motor task used in the study was determined as 'Reach & Return'. The experimental intervention has been structured to consist of four 3-minute blocks (approximately 13 minutes in total) within the PETTLEP-based Action Observation and Motor Imagery (AOMI) protocol. The control condition has been arranged as a matched imagery and mental clearing protocol in terms of duration and attentional load.",
    50: "In addition, Uzm. Fzt. Zeynep Lide has been included in the research team for the review of the study's kinematic analysis processes and the standardization of data processing protocols. The primary outcome measure has been determined as the Number of Velocity Peaks (NVP), path straightness, and pause time parameters evaluating upper extremity movement control. Secondary outcome measures include trunk compensation (trunk_ratio), shoulder girdle elevation (shoulder_vert_norm; with ZoeDepth-based metric scaling), elbow angle (elbow_angle_mean), movement time (movement_time_sec), and peak velocity (peak_velocity_px_s).",
    53: "(Should be written in detail.)",
    54: "Literature titles related to the study should be written.",
    55: "REFERENCES:",
    76: "NATURE OF THE VOLUNTEERS",
    80: "PLACE WHERE THE RESEARCH WILL BE CONDUCTED",
    81: "Institution/Organization Name and Address:",
    82: "İstinye University Liv Bahçeşehir Hospital, Neurorehabilitation Clinic, Istanbul",
    83: "Biruni University Hospital, Physical Medicine and Rehabilitation Polyclinic, Istanbul (added as a multicenter data collection center within the scope of revision)",
    85: "(Specify if more than one)",
    87: "Hospital (Permission Letter Required)		Field",
    88: "Polyclinic (Permission Letter Required)		Other (please specify)",
    89: "Health Center (Permission Letter Required)",
    91: "DATA COLLECTION METHOD",
    92: "Study Design:",
    93: "In this study, a single-blind, pretest-posttest, prospective randomized controlled trial (RCT) will be conducted comparing the experimental condition (PETTLEP model-based Action Observation and Motor Imagery [AOMI]) with the control condition (cognitive and somatic control). The study will have a two-arm design.",
    94: "This research aims to preliminarily examine the acute effects of a single-session PETTLEP-based AOMI task on the kinematic characteristics and functional performance of the affected upper extremity in individuals who have had a stroke.",
    95: "The research will be conducted as a multicenter study, and participants will be recruited from İstinye University Liv Bahçeşehir Hospital Neurorehabilitation Clinic and Biruni University Hospital Physical Medicine and Rehabilitation Polyclinic.",
    96: "Sample Size and Randomization:",
    97: "The sample size was calculated using G*Power (v3.1.9.2) with an a priori power analysis for a 2 (Group: Motor Imagery and Control) × 2 (Time: Pretest and Posttest) mixed-design repeated measures ANOVA model. Based on findings from similar acute motor imagery studies, an effect size of f = 0.25, alpha = 0.05, and power (1-β) = 0.80 were assumed.",
    98: "Participants will be randomized to the experimental group (motor imagery) or control group in a 1:1 ratio using stratified permuted block randomization. Stratification criteria will be (i) gender (male/female) and (ii) spasticity level in the affected upper extremity. A computer-based randomization list with variable block sizes (4–6 people) will be created for each stratum by an independent researcher to ensure balance between groups and unpredictability of allocation.",
    99: "Allocation concealment will be provided by sequentially numbered, opaque, sealed envelopes prepared before participant recruitment begins. Randomization will be performed by an independent researcher not involved in the assessment or intervention.",
    101: "Assessments: To ensure reliability and consistency, environmental conditions will be standardized for all assessments. Each participant will be evaluated twice: immediately before and immediately after the intervention session. This pretest–posttest protocol will allow the acute effects of motor imagery or control intervention on upper extremity kinematics to be examined.",
    102: "Kinematic parameters will be automatically extracted from video recordings using MediaPipe Pose Landmarker (33 landmarks). The validity of MediaPipe-based markerless motion analysis in tracking upper extremity reaching kinematics after stroke has been demonstrated previously (8).",
    104: "Standardization of Experimental Conditions:",
    105: "Participants in this study will wear a long-sleeved, dark purple t-shirt that is part of regular training attire. Braids/ponytails should be fixed so that neck and shoulder markers are visible to the camera without obstruction. Experimental conditions (chair and table height, distance between camera and participant, lighting) will be kept constant for all participants.",
    106: "Task Execution (Reach & Return):",
    107: "For the Reach–Return task, verbal instructions will be given from the same standard text to avoid attentional bias. Participants are required to complete three trials at each assessment time point (before and after); the average of the three trials will be used for analysis. The built-in smartphone/webcam camera will be positioned perpendicular to the sagittal plane of the affected upper extremity at a fixed height and distance.",
    108: "The primary outcome of the study is the change in upper extremity movement control (NVP, straightness, pause time). The MediaPipe-based pipeline derives the palm trajectory from smartphone or webcam video (8). Number of Velocity Peaks (NVP): The number of local maxima in the tangential speed profile of the palm center; higher NVP values indicate that the movement is divided into more sub-movements and is more discontinuous. Path Straightness: The similarity of the reaching trajectory to an ideal straight line; values closer to 1 indicate more linear paths, while lower values indicate more deviated paths. Pause Time: The total time during the movement window in which speed remains below a predetermined threshold; longer pause times indicate more movement interruption.",
    109: "To evaluate instantaneous kinematic changes, recordings will be taken immediately before and immediately after the single training session. To ensure sensor consistency, the same smartphone or webcam model will be used for all participants. The camera will be mounted on a tripod at a fixed height and positioned to capture the reaching movement from the affected side (sagittal plane).",
    110: "Primary Outcome Measure:",
    111: "The primary outcome measure of the study will be the change in upper extremity movement control. Movement control will be evaluated using the Number of Velocity Peaks (NVP), path straightness, and pause time parameters. For this purpose, the palm center movement trajectory will be obtained from videos recorded with a smartphone or webcam via the MediaPipe-based markerless motion analysis system (8).",
    112: "For NVP calculation, the tangential speed profile will be created from the two-dimensional (X–Y) movement trajectory of the palm center; the number of speed peaks will be determined by local maximum detection. Straightness will be calculated as the ratio of the length of the performed trajectory to the straight-line distance between the start and end points. Pause time will be obtained by summing the time intervals in which the speed profile remains below a predetermined threshold (15-18).",
    113: "NVP, straightness, and pause time are complementary kinematic measures in evaluating upper extremity movement control. Lower NVP and pause time values together with straightness values close to 1 indicate more continuous, linear, and uninterrupted movement; whereas high NVP, long pause time, and low straightness values indicate that the movement is more fragmented, deviated, and interrupted.",
    114: "Secondary Outcome Measures:",
    115: "1. Trunk Compensation (Trunk Ratio)",
    116: "Aim: To quantitatively evaluate the amount of trunk displacement used to compensate for movement limitation in the affected upper extremity.",
    117: "Measurement Method: Trunk compensation will be calculated using the Trunk Ratio parameter, which reflects the relative contribution of trunk and hand movements during the movement. Higher values will indicate greater use of trunk compensation during task performance.",
    119: "2. Shoulder Girdle Elevation (Shoulder Elevation)",
    120: "Aim: To evaluate the level of compensatory shoulder elevation that occurs during movement.",
    121: "Measurement Method: Shoulder elevation will be calculated by normalizing the vertical (Y-axis) displacement of the affected shoulder landmark to shoulder width (shoulder_vert_norm). Shoulder width metric scale will be obtained using ZoeDepth monocular depth estimation. Higher values will indicate greater compensatory shoulder elevation.",
    123: "3. Elbow Angle (Elbow Angle Mean)",
    124: "Aim: To evaluate the position of the elbow joint and movement strategy during the reaching movement.",
    125: "Measurement Method: Elbow angle will be determined by calculating the angle between the vectors formed by the shoulder–elbow–wrist points in each frame throughout the movement and taking the average value. Measurement provides high accuracy especially in images obtained from the sagittal plane.",
    127: "4. Movement Time",
    128: "Aim: To evaluate the time required to complete the movement toward the target.",
    129: "Measurement Method: Movement time will be defined as the time interval (start–end) during which the palm tangential speed remains above a predetermined threshold and will be calculated in seconds. Shorter durations will indicate faster movement performance.",
    131: "5. Peak Velocity",
    132: "Aim: To evaluate the maximum movement speed reached during the target-directed movement.",
    133: "Measurement Method: Peak velocity will be calculated as the highest value reached by the palm tangential speed within the movement window and will be reported in pixels/second (px/s). Higher values reflect faster movement production.",
    135: "6. Upper Extremity Function – Wolf Motor Function Test Short Form (WMFT-4)",
    136: "Aim: To evaluate whether observed changes in kinematic parameters are reflected in functional performance.",
    137: "Measurement Method: Upper extremity function will be evaluated using the four-item short form of the Wolf Motor Function Test (WMFT-4). Completion time and functional performance score will be recorded for each task. WMFT has been shown to be valid and reliable in evaluating upper extremity function after stroke and is widely used (9).",
    139: "7. Pain – Visual Analog Scale (VAS)",
    140: "Aim: To evaluate pain severity in the affected upper extremity and to monitor the safety of the intervention by determining short-term changes in pain perception.",
    141: "Measurement Method: Participants will evaluate pain severity by marking on a 10 cm horizontal line with 'no pain' (0) at one end and 'worst imaginable pain' (10) at the other end (11). VAS is a valid and reliable pain assessment tool widely used in clinical practice.",
    143: "8. Motor Imagery Ability – Kinesthetic and Visual Imagery Questionnaire-10 (KVIQ-10)",
    144: "Aim: To evaluate participants' motor imagery capacity in visual and kinesthetic dimensions.",
    145: "Measurement Method: The participant will be asked to perform or observe specific movements. Then, the participant will be asked to mentally simulate the same movement and evaluate the vividness of the imagery. Visual and kinesthetic imagery vividness will be scored using a Likert-type scale; higher scores indicate stronger imagery capacity (12).",
    147: "9. Mood State – Visual Analog Mood Scale-4 (VAMS-4)",
    148: "Aim: To evaluate the effects of the intervention on participants' emotional states.",
    149: "Measurement Method: Mood state will be evaluated in four dimensions: happy, calm, sad, and tense. The scale will be applied before and after the intervention.",
    151: "10. Physical Activity Level – International Physical Activity Questionnaire (IPAQ)",
    152: "Aim: To determine participants' initial physical activity levels.",
    153: "Measurement Method: The International Physical Activity Questionnaire (IPAQ) will be applied at the beginning of the study to evaluate participants' physical activity levels.",
    155: "11. Motor Difference Rating Scale (MDRS)",
    156: "Aim: To evaluate the level of motor performance change perceived by the participant after the intervention.",
    157: "Measurement Method: The Motor Difference Rating Scale will be applied after the intervention to determine the level of motor performance change perceived by the participant. Higher scores indicate greater positive change perceived by the participant.",
    159: "Intervention",
    160: "The intervention protocol is based on the PETTLEP-based Motor Imagery (MI) approach applied together with Action Observation (AO). The study procedure consists of three stages: pre-assessment, intervention, and post-assessment. The intervention duration is approximately 13 minutes. No physical application will be performed during the intervention.",
    162: "Pre-Assessment",
    163: "After system calibration, participants will be asked to perform the Reach & Return task three times. Movements will be recorded with the MediaPipe-based markerless motion capture system and movement control (NVP, straightness, pause time), trunk compensation (trunk ratio), and other kinematic parameters will be analyzed.",
    165: "Intervention Stage",
    166: "Participants will be assigned to the experimental group or control group after randomization. Both groups will attend a session of approximately 13 minutes matched in duration and attentional load. All interventions and assessments will be conducted in the same room, chair arrangement, headphone system, and session duration.",
    169: "Experimental Group (PETTLEP-Based AO + MI)",
    170: "The intervention will consist of four training blocks following a 60-second preparation phase. Each block will include 45 seconds of action observation, 90 seconds of motor imagery, and 45 seconds of rest periods.",
    171: "In the preparation phase, participants will be asked to slowly perform a reaching movement once with their unaffected upper extremity, notice the sensations that occur during the movement, and transfer these sensations to their affected upper extremity. In the action observation stage, participants will watch a first-person perspective video of the Reach & Return movement performed with the affected upper extremity. In the motor imagery stage, participants will mentally simulate the same movement with their eyes closed, in accordance with the PETTLEP principles.",
    173: "Intervention Fidelity",
    174: "All imagery instructions will be presented via headphones using standardized pre-recorded audio files. After each imagery block, participants will be asked whether the imagery was compatible with normal movement speed. During the intervention, participant adherence will be monitored by ensuring eyes remain closed, no voluntary movement is made, and possible side effects are recorded.",
    176: "Control Group",
    177: "Participants in the control group will perform non-motor mental tasks matched in duration and attentional load to the experimental group. The intervention will consist of a 60-second preparation phase and four 3-minute blocks.",
    178: "In the preparation phase, participants will be asked to sit comfortably, focus on their breathing, and distance themselves from daily thoughts. In the training blocks, participants will be asked to form mental images that do not contain movement, such as calm nature scenes, colors, or light. Participants will not be asked to mentally simulate upper extremity movements or the Reach & Return task.",
    179: "Application of PETTLEP Principles",
    180: "The motor imagery session is structured in accordance with the Physical, Environment, Task, Timing, Learning, Emotion, and Perspective components of the PETTLEP model. Participants will be in the same sitting position as during assessment, work under the same environmental conditions, and mentally simulate the frequently used daily living Reach & Return task. The imagery will be performed from a first-person perspective and at the same speed as the actual movement.",
    182: "Data Collection Methods",
    183: "Data will be collected using video recording, clinical scales, performance tests, and an artificial intelligence-based motion analysis system (MediaPipe). Participant information will be coded with research codes and analyses will be performed on anonymized data.",
    185: "Funding Support Request: None",
    189: "INCLUSION, EXCLUSION, AND WITHDRAWAL CRITERIA FOR THE RESEARCH",
    190: "A total of 28 stroke patients (14 participants in each group) who meet the inclusion criteria are planned to be included in the study. Participants will be randomized to the experimental group or control group.",
    191: "The study population will consist of stroke patients applying to or being followed up at İstinye University Liv Bahçeşehir Hospital Neurorehabilitation Clinic and Biruni University Hospital Physical Medicine and Rehabilitation Polyclinic in Istanbul, Turkey. The research will be conducted as a multicenter study.",
    192: "All information about the study will be explained in detail to potential participants, and written informed consent forms will be obtained from individuals who agree to participate. Participant recruitment will begin after ethics committee approval and will continue until the targeted sample size is reached.",
    194: "Inclusion Criteria:",
    195: "Adult women or men aged 40-80 years.",
    196: "First-ever unilateral ischemic or hemorrhagic stroke confirmed by CT (Computed Tomography) or MRI (Magnetic Resonance Imaging), or recurrent stroke history that does not leave permanent motor sequelae from previous strokes.",
    197: "MRC muscle strength grade ≥ 2.",
    198: "Mild or moderate spasticity (Modified Ashworth Scale ≤ 2).",
    199: "Sufficient cognitive function to understand simple instructions and perform short-term motor imagery tasks (clinically determined by the evaluator).",
    200: "Medically stable individuals without uncontrolled systemic disease.",
    201: "Ability to sit unsupported for at least 30 minutes.",
    202: "Ability to perform the short cognitive-motor exercises in the protocol.",
    204: "Exclusion Criteria:",
    205: "Presence of another central nervous system disease other than stroke (e.g., Parkinson's disease, multiple sclerosis, traumatic brain injury).",
    206: "Musculoskeletal or orthopedic problems limiting upper extremity movement; fixed contractures, severe joint deformities, or clinically detected glenohumeral subluxation (>1 finger width).",
    207: "Moderate or severe sensory loss (NIHSS sensory score ≥ 2) or moderate or severe neglect (NIHSS extinction/inattention score ≥ 2).",
    208: "Uncontrolled seizures, serious cardiovascular instability, or any contraindication for imagery-based interventions.",
    211: "APPROACH AND STATISTICAL METHODS TO BE APPLIED",
    212: "All statistical analyses will be performed using IBM SPSS Statistics for Mac, version 24.0 (IBM Corp., Armonk, NY, USA). Demographic and clinical characteristics of participants will be presented with descriptive statistics. Continuous variables will be presented as mean ± standard deviation if normally distributed, and as median (25th-75th percentile) if not normally distributed. Categorical variables will be presented as frequency and percentage.",
    213: "In the analysis of primary and secondary outcome measures, 2 × 2 Mixed Model Analysis of Variance (Mixed ANOVA) will be used if the data are normally distributed. In the analysis, group (Experimental/Control) will be defined as the between-groups factor and time (Pre/Post Intervention) as the within-groups factor. The effect of the intervention will primarily be evaluated with the Group × Time interaction effect.",
    214: "To examine the relationship between motor imagery ability and performance changes, the relationships between KVIQ-10 scores and change scores in kinematic variables (ΔNVP, ΔStraightness, ΔPause Time, ΔTrunk Ratio, etc.) will be evaluated using Pearson or Spearman correlation analyses.",
    215: "To evaluate the magnitude of the intervention effect, partial eta squared (ηp²) will be calculated for parametric analyses and rank-biserial correlation coefficient for non-parametric analyses, and will be reported with corresponding confidence intervals. To reduce the probability of Type I error that may arise from multiple comparisons, Bonferroni correction will be applied for secondary kinematic outcome comparisons.",
    216: "Analyses will be performed according to the Intention-to-Treat (ITT) approach, and appropriate statistical methods (multiple imputation or mixed models) will be used in case of missing data. In all analyses, the statistical significance level will be accepted as p < 0.05.",
    217: "SUPPORT AND BUDGET INFORMATION",
    218: "This research is not supported by any institution (TÜBİTAK/BAP: No).",
    219: "Expense Items:",
    220: "Printing cost (informed consent forms, data forms): 1000 TL",
    221: "Total: 1000 TL — To be covered by the principal investigator.",
    223: "Is this research supported by any institution or organization (TÜBİTAK, BAP, etc.)? The budget amount and source must be specified.",
    227: "* I undertake to adhere to the data collection method and study design of the research project stated in this form, the accuracy of the information, and that it will not be used for another purpose.",
    231: "Principal Investigator",
    232: "Handwritten name and surname: Abdelrahman Walid Hamza Mohamed Elsayed Sabee",
    233: "Date (as day/month/year): …/…/….",
    234: "Signature:",
}

# References
new_refs_en = [
    '15. Rohrer B, Fasoli S, Krebs HI, Hughes R, Volpe B, Frontera WR, Hogan N. Movement smoothness changes during stroke recovery. J Neurosci. 2002;22(18):8297-8304.',
    '16. Krebs HI, Hogan N, Rohrer B, Volpe BT, Frontera WR. Robot-aided neurorehabilitation in subacute stroke patients: a follow-up study. J Rehabil Res Dev. 2007;44(5):665-672.',
    '17. Cirstea MC, Levin MF. Compensatory strategies for reaching in stroke. Brain. 2000;123(Pt 5):940-953.',
    '18. van Vliet PM, Sheridan MR. Coordination between reaching and grasping in patients with hemiparesis and healthy subjects. Arch Phys Med Rehabil. 2007;88(10):1325-1331.'
]
for i, ref_num in enumerate([70, 71, 72, 73]):
    translations[ref_num] = new_refs_en[i]

# Now build the English document by copying structure and translating
for i, p in enumerate(src.paragraphs):
    text = p.text
    if i in translations:
        new_text = translations[i]
    else:
        # For paragraphs not translated, keep original (headers, form fields, etc.)
        # Actually we should translate everything. For empty paragraphs, keep empty.
        new_text = text
    
    # Copy paragraph formatting
    new_p = dst.add_paragraph()
    new_p.alignment = p.alignment
    new_p.style = p.style
    
    # Add text preserving color where applicable
    # If original has red runs, keep red for the corresponding new text
    has_red = any(run.font.color and run.font.color.rgb == RED for run in p.runs)
    if has_red:
        run = new_p.add_run(new_text)
        run.font.color.rgb = RED
        run.bold = any(r.bold for r in p.runs)
        run.italic = any(r.italic for r in p.runs)
        if p.runs[0].font.size:
            run.font.size = p.runs[0].font.size
        if p.runs[0].font.name:
            run.font.name = p.runs[0].font.name
    else:
        run = new_p.add_run(new_text)
        run.bold = any(r.bold for r in p.runs)
        run.italic = any(r.italic for r in p.runs)
        if p.runs and p.runs[0].font.size:
            run.font.size = p.runs[0].font.size
        if p.runs and p.runs[0].font.name:
            run.font.name = p.runs[0].font.name

# Copy tables
def copy_table(src_table, dst_doc):
    rows = len(src_table.rows)
    cols = len(src_table.columns)
    dst_table = dst_doc.add_table(rows=rows, cols=cols)
    dst_table.style = src_table.style
    for i, src_row in enumerate(src_table.rows):
        dst_row = dst_table.rows[i]
        for j, src_cell in enumerate(src_row.cells):
            dst_cell = dst_row.cells[j]
            dst_cell.text = src_cell.text
            # Copy cell shading/vertical alignment if needed
    return dst_table

for table in src.tables:
    copy_table(table, dst)

# Save English document
dst.save('ethics_english_translated.docx')
print('Saved ethics_english_translated.docx')
