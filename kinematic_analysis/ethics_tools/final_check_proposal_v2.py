from docx import Document
from pathlib import Path

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc = Document(str(p))

out = open(r"C:\Users\acer\AppData\Local\Temp\opencode\proposal_final_check_v2.txt", "w", encoding="utf-8")

out.write("=== Remaining OpenCap/OpenSim mentions ===\n")
found = False
for i, para in enumerate(doc.paragraphs):
    if 'OpenCap' in para.text or 'OpenSim' in para.text:
        out.write(f"[{i}] {para.text}\n\n")
        found = True
if not found:
    out.write("None found.\n")

out.write("\n=== References section ===\n")
for i, para in enumerate(doc.paragraphs):
    t = para.text.strip()
    if t and (t.startswith("8.") or t.startswith("14.") or t.startswith("15.") or t.startswith("16.")):
        out.write(f"[{i}] {t}\n")

out.write("\n=== Key tool-assignment paragraphs ===\n")
keywords = [
    "MediaPipe will quantify",
    "Phyphox will capture",
    "Movement Smoothness (Spectral Arc Length",
    "Trunk Ratio",
    "side-view camera",
    "wrist-mounted",
]
for kw in keywords:
    out.write(f"\n--- {kw} ---\n")
    for para in doc.paragraphs:
        if kw in para.text:
            out.write(para.text[:500] + "\n")
            break

out.close()
print("done")
