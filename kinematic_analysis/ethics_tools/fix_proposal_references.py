# -*- coding: utf-8 -*-
"""
Fix references in the updated proposal:
- Replace reference 8 (OpenCap) with MediaPipe paper
- Remove duplicate MediaPipe reference (currently #14)
- Renumber 15-17 to 14-16
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc = Document(str(p))
RED = RGBColor(0xFF, 0x00, 0x00)

# Replace reference 8 text
ref8_old = ("8.   Uhlrich SD, Falisse A, Kidziński Ł, Muccini J, Ko M, Chaudhari AS, et al. OpenCap: 3D human movement "
            "dynamics from smartphone videos. PLOS Comput Biol. 2023;19(10):e1011462.")
ref8_new = ("8.   Wagh V, Scott MW, Andrushko JW, Jones CB, Larssen BC, Boyd LA, Kraeutner SN. Using MediaPipe to track "
            "upper-limb reaching movements after stroke: a proof-of-principle study. J Neuroeng Rehabil. 2025 Nov 25;22(1):268.")

for p_para in doc.paragraphs:
    if ref8_old in p_para.text:
        # Clear and rewrite in red
        for run in p_para.runs:
            run._element.getparent().remove(run._element)
        r = p_para.add_run(ref8_new)
        r.font.color.rgb = RED
        print("OK: reference 8 updated to MediaPipe")
        break

# Find and remove duplicate MediaPipe reference (currently #14)
ref14_text = ("14. Wagh V, Scott MW, Andrushko JW, Jones CB, Larssen BC, Boyd LA, Kraeutner SN. Using MediaPipe to track "
              "upper-limb reaching movements after stroke: a proof-of-principle study. J Neuroeng Rehabil. 2025 Nov 25;22(1):268.")
for p_para in doc.paragraphs:
    if ref14_text in p_para.text:
        p_para._element.getparent().remove(p_para._element)
        print("OK: duplicate reference 14 removed")
        break

# Renumber 15-17 to 14-16
renumber_map = {
    "15. Staacks": "14. Staacks",
    "16. Dobkin BH, Dorsch": "15. Dobkin BH, Dorsch",
    "17. Dobkin BH. Wearable": "16. Dobkin BH. Wearable",
}
for old_prefix, new_prefix in renumber_map.items():
    for p_para in doc.paragraphs:
        if p_para.text.startswith(old_prefix):
            new_text = new_prefix + p_para.text[len(old_prefix):]
            for run in p_para.runs:
                run._element.getparent().remove(run._element)
            r = p_para.add_run(new_text)
            r.font.color.rgb = RED
            print(f"OK: renumbered {old_prefix.split('.')[0]} -> {new_prefix.split('.')[0]}")
            break

# Save
doc.save(str(p))
print(f"Saved fixed proposal to: {p}")
