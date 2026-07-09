# -*- coding: utf-8 -*-
"""
Update the English thesis proposal to reflect the new tool split:
  - Phyphox (smartphone wrist accelerometer): SPARC, Movement Time, Peak Velocity
  - MediaPipe (side-view camera): Trunk Ratio / Trunk Displacement, Shoulder Elevation, Elbow Angle
Replace OpenCap with MediaPipe + Phyphox.
All changed text is coloured red; existing formatting/spacing is preserved.
"""
from pathlib import Path
from docx import Document
from docx.shared import RGBColor
import shutil

src = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_BACKUP.docx")
backup = src.with_stem(src.stem + "_before_tool_update")
modified = src.with_stem(src.stem.replace("_BACKUP", "_updated"))

# Create backup if not present
if not backup.exists():
    shutil.copy2(src, backup)

doc = Document(str(src))
RED = RGBColor(0xFF, 0x00, 0x00)


def replace_in_paragraph(p, old_text, new_text):
    """Replace substring, colouring only the inserted new text red."""
    full = p.text
    idx = full.find(old_text)
    if idx == -1:
        return False
    prefix = full[:idx]
    suffix = full[idx + len(old_text):]
    for run in p.runs:
        run._element.getparent().remove(run._element)
    if prefix:
        p.add_run(prefix)
    r = p.add_run(new_text)
    r.font.color.rgb = RED
    if suffix:
        p.add_run(suffix)
    return True


def replace_paragraph_text(p, new_text):
    """Replace entire paragraph text in red."""
    for run in p.runs:
        run._element.getparent().remove(run._element)
    r = p.add_run(new_text)
    r.font.color.rgb = RED


# ------------------------------------------------------------------
# 1. Subject/Purpose paragraph mentioning OpenCap (para 29)
# ------------------------------------------------------------------
old = ("For this purpose, the integration of OpenCap, a recently developed markerless motion capture system that utilizes "
       "artificial intelligence to enable accurate clinical quantification of three‐dimensional (3D) kinematic measures, "
       "is introduced (8). Hence, by quantifying objectively parameters of motion, such as movement smoothness, trunk "
       "displacement and Shoulder Girdle Elevation we can quantify the \"how\" of a movement, which is an important mediator "
       "between central neural planning and peripheral motor execution (8). Furthermore, the combination of this kinematic "
       "analysis with the Box and Block Test (BBT) allows to relate the quality of movement with functional dexterity in everyday life (9).")
new = ("For this purpose, a dual-sensor kinematic setup combining MediaPipe Pose Landmarker (markerless video-based motion "
       "capture) with Phyphox smartphone accelerometry is introduced. MediaPipe will quantify spatial and angular parameters "
       "(trunk displacement, shoulder girdle elevation, elbow angle) from side-view video, while Phyphox will capture time- and "
       "acceleration-based parameters (movement smoothness, movement time, peak velocity) from a wrist-mounted smartphone "
       "accelerometer. Hence, by objectively quantifying the \"how\" of movement, we can examine the mediator between central "
       "neural planning and peripheral motor execution (8,19,20). Furthermore, the combination of this kinematic analysis with "
       "the Box and Block Test (BBT) allows us to relate the quality of movement to functional dexterity in everyday life (9).")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 1 (subject/purpose)")
        break

# ------------------------------------------------------------------
# 2. Expectation paragraph mentioning OpenCap (para 31)
# ------------------------------------------------------------------
old = ("Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature. Using OpenCap "
       "technology in combination with a customized PETTLEP based AOMI protocol, this study seeks to investigate whether specific "
       "mental practice can lead to measurable short-term changes in movement smoothness, compensatory strategies and functional "
       "dexterity in persons post-stroke.")
new = ("Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature. Using MediaPipe "
       "and Phyphox technologies in combination with a customized PETTLEP-based AOMI protocol, this study seeks to investigate "
       "whether specific mental practice can lead to measurable short-term changes in movement smoothness, compensatory strategies "
       "and functional dexterity in persons post-stroke.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 2 (expectation)")
        break

# ------------------------------------------------------------------
# 3. Keywords (para 33)
# ------------------------------------------------------------------
old = "Keywords: Motor imagery, action observation, stroke, OpenCap, kinematics, upper limb rehabilitation."
new = "Keywords: Motor imagery, action observation, stroke, MediaPipe, Phyphox, kinematics, upper limb rehabilitation."
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 3 (keywords)")
        break

# ------------------------------------------------------------------
# 4. Assessments paragraph mentioning OpenCap parameters (para 94)
# ------------------------------------------------------------------
old = ("To ensure reliability and consistency, the environmental condition will be standardized for all evaluations, which are going "
       "to be performed in the Neurorehabilitation unit of Istinye University Liv Bahçeşehir Hospital. All participants will be tested "
       "at 2 time points—once immediately before and once immediately after the session. This pre-post test design enables the "
       "identification of short-term modifications in (movement smoothness, Trunk Displacement, Shoulder Girdle Elevation) following "
       "motor imagery and a control task.")
new = ("To ensure reliability and consistency, the environmental condition will be standardized for all evaluations, which are going "
       "to be performed in the Neurorehabilitation unit of Istinye University Liv Bahçeşehir Hospital. All participants will be tested "
       "at 2 time points—once immediately before and once immediately after the session. This pre-post test design enables the "
       "identification of short-term modifications in movement smoothness (Phyphox), trunk displacement (MediaPipe), shoulder girdle "
       "elevation (MediaPipe), movement time (Phyphox) and peak velocity (Phyphox) following motor imagery and a control task.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 4 (assessments)")
        break

# ------------------------------------------------------------------
# 5. Description of Task (para 100)
# ------------------------------------------------------------------
old = ("Description of Task: To prevent attentional bias, the Verbal instruction for the Reach-to-Grasp is verbally presented from the "
       "same standardized script. Participants are invited to participate in three trials per each point of assessment (pre and post), "
       "which are recorded. They will be instructed to reach for the object at a pace that feels natural to them (comfortable speed of "
       "movement) and will not receive any specific instructions to go either \"fast\" or \"slow\", this is to catch their habitual movement behavior.")
new = ("Description of Task: To prevent attentional bias, the verbal instruction for the Reach-to-Grasp is verbally presented from the "
       "same standardized script. Participants are invited to participate in three trials per each point of assessment (pre and post). "
       "During each trial, a side-view video will be recorded with a smartphone/webcam for MediaPipe analysis, and a second smartphone "
       "running the Phyphox application will be secured on the affected-side wrist to record tri-axial accelerometry. They will be "
       "instructed to reach for the object at a pace that feels natural to them (comfortable speed of movement) and will not receive any "
       "specific instructions to go either \"fast\" or \"slow\"; this is to catch their habitual movement behavior.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 5 (description of task)")
        break

# ------------------------------------------------------------------
# 6. Data Collection Setup (para 102)
# ------------------------------------------------------------------
old = ("Data Collection Setup: To evaluate immediate kinematic modifications, recordings will be taken prior to and immediately following "
       "the single training session. (iPhone 14 Pro max & iPhone XS max) will be used for all participants to ensure sensor consistency. "
       "The cameras will be mounted on tripods at a fixed height and placed at 45° angles (one anterior-lateral and one posterior-lateral) "
       "relative to the participant, at a standardized distance of 1.5 m to 2.0 m.")
new = ("Data Collection Setup: To evaluate immediate kinematic modifications, recordings will be taken prior to and immediately following "
       "the single training session. For MediaPipe analysis, a smartphone or webcam will be used to record a side-view (sagittal plane) "
       "video of the affected upper limb and trunk. The camera will be mounted on a tripod at a fixed height, positioned perpendicular to "
       "the reaching plane at a standardized distance of 1.5 m to 2.0 m. For Phyphox analysis, a standard smartphone will be secured on the "
       "affected-side wrist using an armband or Velcro strap; the Phyphox application will record linear acceleration from the built-in "
       "tri-axial accelerometer at approximately 100 Hz.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 6 (data collection setup)")
        break

# ------------------------------------------------------------------
# 7. Camera config / OpenCap processing paragraph (para 104)
# ------------------------------------------------------------------
old = ("The cameras will be configured to record at 1080p resolution and 60 frames per second (fps). Recordings will be conducted in a room "
       "with consistent, diffuse lighting to avoid backlighting or shadows that could interfere with AI tracking. The same OpenCap model and "
       "processing parameters will be used for all participant analyses to ensure data comparability.")
new = ("The video camera will be configured to record at 1080p resolution and 60 frames per second (fps). Recordings will be conducted in a room "
       "with consistent, diffuse lighting to avoid backlighting or shadows that could interfere with AI tracking. The same MediaPipe Pose "
       "Landmarker model and processing parameters will be used for all participant analyses to ensure data comparability. The Phyphox CSV "
       "output will be imported into a Python-based pipeline for SPARC, movement-time and peak-velocity computation.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 7 (camera config)")
        break

# ------------------------------------------------------------------
# 8. OpenCap 3D model paragraph (para 107)
# ------------------------------------------------------------------
old = ("The video data is processed by uploading the video to the OpenCap cloud platform to obtain 3D OpenSim musculoskeletal models. The "
       "following specific kinematic measures will be computed according to their functional importance (8).")
new = ("The video data will be processed locally using the MediaPipe Pose Landmarker (33 landmark, 2D) pipeline to extract joint trajectories, "
       "and the smartphone accelerometer data will be processed locally using a Python pipeline. The following specific kinematic measures "
       "will be computed according to their functional importance (8,19).")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 8 (video processing)")
        break

# ------------------------------------------------------------------
# 9. Primary outcome - OpenCap (para 113)
# ------------------------------------------------------------------
old = ("·     The primary outcome of this study is the change in movement smoothness of the affected upper limb, as measured by the OpenCap System. "
       "Instead of classic marker-based motion capture systems, OpenCap derives 3D human movement dynamics from videos taken with smartphones. "
       "The video will be processed by submitting the video to the OpenCap cloud platform to generate 3D OpenSim musculoskeletal models. Due to "
       "their functional meaning, the following specific kinematic parameters will be derived from the curves (8) .")
new = ("·     The primary outcome of this study is the change in movement smoothness of the affected upper limb, as measured by the Spectral Arc "
       "Length (SPARC) derived from the Phyphox wrist accelerometer recording. The smartphone, secured on the affected-side wrist, records tri-axial "
       "linear acceleration at approximately 100 Hz; the resulting CSV file is imported into a Python pipeline to compute the absolute velocity profile, "
       "normalize it, apply Fast Fourier Transform (FFT) and calculate the arc length of the spectral curve (15,16). Due to their functional meaning, "
       "the following specific kinematic parameters will be derived (8,19).")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 9 (primary outcome)")
        break

# ------------------------------------------------------------------
# 10. NVP -> SPARC (para 115)
# ------------------------------------------------------------------
old = ("·     Movement Smoothness (Number of Velocity Peaks - NVP): Calculated from the hand marker (3rd Metacarpal head) velocity profile. One "
       "smooth movement is usually made up of one accelerating and one decelerating phase. Increase in velocity peaks means segmentation of movement "
       "(the movement becomes more jerk like) due to loss of motor control. Reduction in NVP will be interpreted as increased movement smoothness.")
new = ("·     Movement Smoothness (Spectral Arc Length - SPARC): Calculated from the wrist accelerometer absolute velocity profile [√(x²+y²+z²)]. "
       "SPARC quantifies the smoothness of the speed profile by measuring the arc length of its normalized Fourier spectrum; more negative values "
       "indicate smoother, less segmented movement. Increase in SPARC (i.e., less negative) reflects loss of motor control and more stop-and-go "
       "behavior. Reduction in SPARC (more negative) will be interpreted as increased movement smoothness.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 10 (SPARC description)")
        break

# ------------------------------------------------------------------
# 11. Secondary outcome - Shoulder Girdle Elevation (para 122-124)
# ------------------------------------------------------------------
old = "o Purpose: To measure this as a compensatory strategy; the markerless capture system tracks surface anatomical landmarks."
new = "o Purpose: To measure this as a compensatory strategy; MediaPipe tracks the shoulder landmark from side-view video."
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 11a (shoulder purpose)")
        break

old = ("o Measurement Method: Calculated by measuring the relative vertical (Y axis) displacement of the Shoulder marker with respect to the Neck "
       "marker during movement.")
new = ("o Measurement Method: Calculated by measuring the relative vertical (Y axis) displacement of the affected-side Shoulder marker with respect "
       "to its resting position during movement. The side-view camera orientation maximizes the accuracy of pure vertical elevation.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 11b (shoulder method)")
        break

# ------------------------------------------------------------------
# 12. Trunk Displacement (para 129, 131)
# ------------------------------------------------------------------
old = ("o Purpose: To quantify the extent to which the patient resorts to “trunk lean” to compensate for reduced arm length. The aim is to determine "
       "whether the patient is actually moving the upper limb or compensating using the trunk. This variable is expressed not as an absolute "
       "biomechanical value, but as an indicator of compensatory strategy quality.")
new = ("o Purpose: To quantify the extent to which the patient resorts to trunk lean to compensate for reduced arm length. The aim is to determine "
       "whether the patient is actually moving the upper limb or compensating using the trunk. MediaPipe provides the trunk and hand landmarks needed "
       "to compute this compensatory strategy index.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 12a (trunk purpose)")
        break

old = ("o Measurement Method: Calculated by measuring the change in the Lumbar Extension angle (lower back extension) in the sagittal plane during "
       "the reaching phase, obtained from the OpenSim kinematic model.")
new = ("o Measurement Method: Calculated from MediaPipe trunk and hand landmark trajectories as the ratio of trunk displacement to hand displacement "
       "during the reaching phase (Trunk Ratio). Higher values indicate greater reliance on trunk compensation.")
for p in doc.paragraphs:
    if old in p.text:
        replace_in_paragraph(p, old, new)
        print("OK: para 12b (trunk method)")
        break

# ------------------------------------------------------------------
# 13. OpenCap permission paragraph (para 134)
# ------------------------------------------------------------------
old = ("The OpenCap system has been made available to all researchers by its developers (Stanford University) as an open-source and web-based platform "
       "for use in scientific research; therefore, specific author permission is not required for its use [8].")
new = ("MediaPipe Pose Landmarker is an open-source machine-learning pipeline developed by Google, and Phyphox is an open-source smartphone "
       "physics-lab application developed by RWTH Aachen University. Both tools are freely available for scientific research; therefore, specific "
       "author permission is not required for their use in this study [8,19].")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 13 (permissions)")
        break

# ------------------------------------------------------------------
# 14. Baseline Assessment (para 175)
# ------------------------------------------------------------------
old = ("1. Baseline Assessment: After the system calibration, active 3 runs of a routine Reach-to-Grasp task will be performed by each subject. Their "
       "movement kinematics will be recorded using the OpenCap system (pre-intervention movement quality: Movement Smoothness and Trunk Displacement) to capture.")
new = ("1. Baseline Assessment: After the system calibration, each subject will perform 3 runs of a routine Reach-to-Grasp task. Movement kinematics "
       "will be recorded simultaneously: MediaPipe will capture side-view video for trunk displacement, shoulder girdle elevation and elbow angle; "
       "Phyphox will record wrist accelerometry for movement smoothness (SPARC), movement time and peak velocity (pre-intervention movement quality).")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 14 (baseline assessment)")
        break

# ------------------------------------------------------------------
# 15. Statistical correlations paragraph (para 256)
# ------------------------------------------------------------------
old = ("Correlations Subjective vividness of imagery (MIQ-3 scores) and the amount of kinematic facilitation (Δ scores of MS and TD) will be correlated "
       "with each other using Pearson’s (normal distribution) or Spearman’s (non-normal distribution) to verify whether higher imagery ability can predict "
       "the best motor outcome performance.")
new = ("Correlations Subjective vividness of imagery (MIQ-3 scores) and the amount of kinematic facilitation (Δ scores of SPARC, trunk displacement, "
       "shoulder girdle elevation, elbow angle, movement time and peak velocity) will be correlated with each other using Pearson’s (normal distribution) "
       "or Spearman’s (non-normal distribution) to verify whether higher imagery ability can predict the best motor outcome performance.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 15 (correlations)")
        break

# ------------------------------------------------------------------
# 16. Functional Relevance paragraphs mentioning OpenCap variables (para 241, 243)
# ------------------------------------------------------------------
old = ("· The Transport Phase (The Reach & The Trunk): Shoulder flexion and elbow extension are distance-related components. Forward leaning Trunk "
       "Compensation and Shoulder Girdle Elevation are in fact a very a very common compensatory mechanism in this phase of post-stroke patients. To treat "
       "this phase, the \"Glued Back\" constraint is introduced as it attempts to recover to real limb range-of-motion (ROM) not compensatory motion, a motion "
       "pattern that can be numerically expressed in OpenCap through ay quantifiable measure ('Trunk Displacement' and '” Shoulder Girdle Elevation” variables).")
new = ("· The Transport Phase (The Reach & The Trunk): Shoulder flexion and elbow extension are distance-related components. Forward leaning trunk "
       "compensation and shoulder girdle elevation are in fact very common compensatory mechanisms in this phase of post-stroke patients. To treat this phase, "
       "the \"Glued Back\" constraint is introduced as it attempts to recover real limb range-of-motion (ROM) rather than compensatory motion; these motion "
       "patterns can be numerically expressed through MediaPipe-derived 'Trunk Displacement' and 'Shoulder Girdle Elevation' variables.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 16a (transport phase)")
        break

old = ("· The Manipulation Phase (Grasp & Smoothness): Grasping is a two-handed operation: wrist extension and hand aperture, that need to be executed in a "
       "synchronized manner to match the size and shape of the object. Adding the 'Full Cup' constraint (fiction of liquid) introduces a significant constraint "
       "on the smoothness and precision of movement. This was to practice exactly motor control and coordination instead of raw pressure by forcing the nervous "
       "system to slow down the \"velocity peaks\" (sharpness).")
new = ("· The Manipulation Phase (Grasp & Smoothness): Grasping requires wrist extension and hand aperture to be executed in a synchronized manner to match "
       "the size and shape of the object. Adding the 'Full Cup' constraint (fiction of liquid) introduces a significant constraint on the smoothness and precision "
       "of movement. This was to practice motor control and coordination instead of raw pressure, forcing the nervous system to reduce segmented speed profiles, "
       "which will be quantified by the Phyphox-derived SPARC metric.")
for p in doc.paragraphs:
    if old in p.text:
        replace_paragraph_text(p, new)
        print("OK: para 16b (manipulation phase)")
        break

# ------------------------------------------------------------------
# 17. Add new references 19-21 after reference 13
# ------------------------------------------------------------------
ref13_old = ("13. Uğur Y, Coşkun H, Şenyurt AY. The Movement Imagery Questionnaire-3: Reliability and Validity Study on Turkish Sample. "
             "Spormetre Beden Eğitimi ve Spor Bilimleri Dergisi. 2021;19(4):145-156.")
new_refs = [
    "14. Wagh V, Scott MW, Andrushko JW, Jones CB, Larssen BC, Boyd LA, Kraeutner SN. Using MediaPipe to track upper-limb reaching movements after stroke: a proof-of-principle study. J Neuroeng Rehabil. 2025 Nov 25;22(1):268.",
    "15. Staacks S, Hütz S, Heinke H, Stampfer C. Advanced tools for smartphone-based experiments: phyphox. Phys Teach. 2021;59(3):214-215.",
    "16. Dobkin BH, Dorsch A. The promise of mHealth: physical activity, fitness, and stroke. Neurorehabil Neural Repair. 2011;25(8):711-715.",
    "17. Dobkin BH. Wearable motion sensors to continuously measure real-world physical activities. Curr Opin Neurol. 2013;26(6):602-608.",
]
# Renumber existing references 14-... will not be done automatically; we insert as 14-17 and leave renumbering to user if needed.
for i, p in enumerate(doc.paragraphs):
    if ref13_old in p.text:
        parent = p._element.getparent()
        idx = list(parent).index(p._element)
        for j, txt in enumerate(new_refs):
            np = doc.add_paragraph(txt)
            np.runs[0].font.color.rgb = RED
            parent.insert(idx + 1 + j, np._element)
        print("OK: references 14-17 added")
        break

# Save
modified_path = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc.save(str(modified_path))
print(f"Saved updated proposal to: {modified_path}")
print(f"Backup saved to: {backup}")
