from docx import Document
from pathlib import Path

p = Path(r"D:\Thesis app\phyphox\ethics commitee\proposal\Abdelrahman Sabee Proposal (1)_updated.docx")
doc = Document(str(p))

out = open(r"C:\Users\acer\AppData\Local\Temp\opencode\proposal_final_check_v3.txt", "w", encoding="utf-8")

out.write("=== References 8 and 14-18 ===\n")
for para in doc.paragraphs:
    t = para.text.strip()
    if t and (t.startswith("8.") or t.startswith("14.") or t.startswith("15.") or t.startswith("16.") or t.startswith("17.") or t.startswith("18.")):
        out.write(t + "\n")

out.write("\n=== In-text citations containing 14-18 ===\n")
for i, para in enumerate(doc.paragraphs):
    if any(f"({x}" in para.text or f", {x})" in para.text or f", {x}," in para.text for x in ["14", "15", "16", "17", "18"]):
        out.write(f"[{i}] {para.text[:300]}\n\n")

out.write("\n=== OpenCap/OpenSim remaining ===\n")
found = False
for para in doc.paragraphs:
    if 'OpenCap' in para.text or 'OpenSim' in para.text:
        out.write(para.text + "\n")
        found = True
if not found:
    out.write("None.\n")

out.close()
print("done")
