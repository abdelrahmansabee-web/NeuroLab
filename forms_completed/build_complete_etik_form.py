# -*- coding: utf-8 -*-
"""Build complete updated ethics application form (.docx) — single final file."""
from pathlib import Path

from docx import Document
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt

OUT = Path(r"D:\Thesis app\manuscript f\REVIZYON_PAKETI")
OUT2 = Path(r"D:\Thesis app\NeuroLab\forms_completed")
FILENAME = "insan-arastirmalari-etik-kurul-basvuru-formu-DOLDURULMUS.docx"


def heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(4)
    return p


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build():
    doc = Document()
    para(doc, "Tarih: 18/06/2026", bold=True)
    doc.add_paragraph()

    heading(doc, "İNSAN ARAŞTIRMALARI ETİK KURULU BAŞVURU FORMU", 0)
    para(doc, "PROJENİN ADI:", bold=True)
    para(
        doc,
        "İnme Sonrası Tek Seans PETTLEP Temelli Eylem Gözlemi ve Motor İmgeleme (AOMI) "
        "Uygulamasının Üst Ekstremite Kinematiği Üzerindeki Anlık Etkileri: "
        "Randomize Kontrollü Çalışma",
    )
    doc.add_paragraph()

    heading(doc, "BAŞVURU BİLGİLERİ", 1)
    bullets(
        doc,
        [
            "Başvurunun şekli: Düzeltme / Onaylanan Projede Değişiklik Bildirimi",
            "Araştırmanın niteliği: Yüksek Lisans Tez Çalışması",
            "Bu araştırmanın sonuçları bilimsel bir dergide yayınlanacaktır.",
        ],
    )

    heading(doc, "PROJE DANIŞMANI ÖĞRETİM ÜYESİ", 1)
    para(doc, "Dr. Öğr. Üyesi Begüm Kara Kaya | Biruni Üniversitesi, Fizyoterapi ve Rehabilitasyon")
    para(doc, "begum.kara@biruni.edu.tr")

    heading(doc, "SORUMLU ARAŞTIRMACI", 1)
    para(doc, "Yüksek Lisans Öğrencisi / Fizyoterapist Abdelrahman Walid Hamza Mohamed Elsayed Sabee")
    para(doc, "İstinye Üniversitesi, Lisansüstü Eğitim Enstitüsü, Fizyoterapi ve Rehabilitasyon")
    para(doc, "abdelrahman.sabee@stu.istinye.edu.tr")

    heading(doc, "YARDIMCI ARAŞTIRMACILAR", 1)
    para(doc, "Doç. Dr. Çiğdem Çınar | Biruni Üniversitesi Hastanesi, FTR")
    para(doc, "0507 783464 | ccinar@biruni.edu.tr")

    heading(doc, "ARAŞTIRMANIN GEREKÇESİ, AMACI VE KAYNAKLARI", 1)
    para(
        doc,
        "Stroke sonrası üst ekstremite motor bozukluğu yetişkinlerde önemli bir fonksiyonel "
        "kısıtlılık nedenidir. Hareket pürüzsüzlüğündeki bozulma (duraklama yüzdesi), gövde "
        "kompensasyonu ve omuz yükselmesi, merkezi motor planlama ve yürütme bozukluklarını "
        "yansıtır. PETTLEP çerçevesinde yapılandırılmış Eylem Gözlemi ve Motor İmgeleme (AOMI), "
        "fiziksel efor gerektirmeden motor yolları hedefleyebilir; ancak tek seanslı AOMI'nin "
        "anlık kinematik etkileri objektif markerless analizle yeterince araştırılmamıştır.",
    )
    para(
        doc,
        "Bu randomize kontrollü çalışma, inme sonrası bireylerde tek seans PETTLEP temelli AOMI'nin "
        "çapraz vücut ulaşma–dönüş (Reach & Return) görevinde hareket pürüzsüzlüğü "
        "(smoothness_pause_pct; birincil sonuç), toplam hareket süresi, gövde-kompanzasyon oranı, "
        "omuz dikey yer değiştirmesi, tepe el hızı ve üst ekstremite fonksiyonu (WMFT-4) üzerindeki "
        "anlık etkilerini, zaman eşleştirilmiş imgeleme ve zihinsel arınma kontrol koşulu ile "
        "karşılaştırmayı amaçlamaktadır.",
    )
    para(doc, "Literatür ve kaynaklar:", bold=True)
    bullets(
        doc,
        [
            "Holmes & Collins (2001) — PETTLEP motor imgeleme çerçevesi",
            "Guerra et al. (2017); Kim & Lee (2022) — AOMI ve inme rehabilitasyonu",
            "Schwarz et al. (2022); Lakshminarayanan et al. (2023) — motor imgeleme etkinliği",
            "Schuster et al. (2011); Di Rienzo et al. (2016) — motor olmayan kontrol koşulları",
            "Uğur et al. (2021) — KVIQ-10 Türkçe geçerlilik",
            "MediaPipe markerless kinematik analiz (Lugaresi et al., 2019)",
        ],
    )
    para(doc, "REVİZYON (onaylanmış protokolde değişiklik):", bold=True)
    bullets(
        doc,
        [
            "Çalışma çok merkezli hale getirilmiş; Biruni Üniversitesi Hastanesi ikinci veri "
            "toplam merkezi olarak eklenmiş ve Doç. Dr. Çiğdem Çınar yardımcı araştırmacı "
            "olarak ekibe dahil edilmiştir.",
            "Motor görev Reach & Wipe yerine Reach & Return (ulaşma–dönüş) olarak güncellenmiştir.",
            "Deneysel müdahale: 4 blok × 3 dk (Watch 45 sn + Imagine 90 sn + Rest 45 sn); "
            "kalibrasyon 60 sn; toplam ~13 dk.",
            "Kontrol koşulu: motor olmayan imgeleme + zihinsel arınma (video ve frekans tonu yok); "
            "toplam ~13 dk.",
        ],
    )

    heading(doc, "GÖNÜLLÜLERİN NİTELİĞİ", 1)
    para(doc, "Hasta (İnme / Stroke): n=28, yaş 40–80, toplam 28")
    para(doc, "Etik Kurul onayı sonrası veri toplama süresi: 12 ay (Şubat 2026 – Şubat 2027)")

    heading(doc, "ARAŞTIRMANIN YAPILACAĞI YER", 1)
    bullets(
        doc,
        [
            "İstinye Üniversitesi Liv Bahçeşehir Hastanesi, Nörorehabilitasyon Kliniği, İstanbul",
            "Biruni Üniversitesi Hastanesi, Fiziksel Tıp ve Rehabilitasyon Polikliniği, İstanbul",
        ],
    )
    para(doc, "Hastane / Poliklinik (izin yazısı alınacaktır)")

    heading(doc, "VERİ TOPLAMA YÖNTEMİ", 1)
    bullets(doc, ["Ölçek", "Test", "Gözlem", "Bilgisayar ortamında uygulama", "Görüntü kaydı", "Ses kaydı"])

    heading(doc, "ARAŞTIRMAYA DAHİL / HARİÇ TUTULMA KRİTERLERİ", 1)
    para(doc, "Dahil olma kriterleri:", bold=True)
    bullets(
        doc,
        [
            "40–80 yaş arası yetişkinler",
            "Tek taraflı iskemik veya hemorajik inme",
            "Etkilenen üst ekstremitede artık gönüllü hareket",
            "Modified Ashworth Skalası ≤ 2",
            "MMSE ≥ 21",
            "Tıbbi olarak stabil; en az 30 dk desteksiz oturabilme",
            "Reach & Return görevini ve seans protokolüne uyum",
        ],
    )
    para(doc, "Hariç tutulma kriterleri:", bold=True)
    bullets(
        doc,
        [
            "İnme dışı santral sinir sistemi hastalıkları",
            "Üst ekstremite hareketini kısıtlayan ortopedik durumlar",
            "Görevi veya motor imgelemeyi engelleyen ciddi duyu kaybı / ihmal",
            "Kontrolsüz nöbet, kardiyovasküler instabilite",
            "Protokole uyum sağlayamama veya gönüllü geri çekilme",
        ],
    )

    heading(doc, "UYGULANACAK YAKLAŞIM VE İSTATİSTİKSEL YÖNTEM", 1)
    para(
        doc,
        "Tasarım: Tek kör, ön test–son test, prospektif paralel grup randomize kontrollü "
        "çalışma (RCT); çok merkezli.",
    )
    para(
        doc,
        "Örneklem: n=28 (grup başına 14); G*Power; permütasyon blok randomizasyon (1:1); "
        "cinsiyet ve MAS stratifikasyonu; kinematik çıkarımı otomatik/kör.",
    )
    para(
        doc,
        "Değerlendirme görevi (pre/post): Çapraz vücut Reach & Return — etkilenen kol ile "
        "uyluktan öne ulaşma, kısa duraklama, uyluğa dönüş; üç deneme, ortalama analiz.",
    )

    heading(doc, "PETTLEP Çerçevesi (Holmes & Collins, 2001)", 2)
    bullets(
        doc,
        [
            "P (Physical): Gerçek oturma düzeni, normal kıyafet, aynı sandalye",
            "E (Environment): Değerlendirme odası; tablet ekranında ayna video",
            "T (Task): Reach & Return; tek akıcı hareket dizisi",
            "T (Timing): Gerçek zamanlı kinestetik MI; Blok 3–4 frekans tonları",
            "L (Learning): Bloklar arası düzeltici rehberlik",
            "E (Emotion): Rahatlık, güven, hazır anından sonra ulaşma",
            "P (Perspective): Watch = 3. kişi; Imagine = 1. kişi kinestetik",
        ],
    )

    heading(doc, "DENEYSEL GRUP — PETTLEP-AOMI Protokolü (~13 dk)", 2)
    bullets(
        doc,
        [
            "Kalibrasyon: 60 sn",
            "4 blok × 3 dk: Watch 45 sn + Imagine 90 sn + Rest 45 sn",
            "Fiziksel hareket yalnızca pre/post değerlendirmede",
        ],
    )

    para(doc, "Kalibrasyon (60 sn):", bold=True)
    para(
        doc,
        "Rahat oturun, sırtınızı sandalyeye yaslayın. İki ayağınız yere bassın. Etkilenen "
        "elinizi avucunuz açık şekilde uyluğunuzda dinlendirin. Etkilenmemiş kolunuzla bir kez "
        "yavaşça öne uzanın — bir bardağa uzanır gibi. Bu his için tek bir kelime seçin "
        "(sabit, kolay veya yumuşak) ve sessizce söyleyin. Gözlerinizi kapatın. Aynı kelimeyi "
        "ve hissi etkilenen kolunuza gönderin.",
    )

    exp_blocks = [
        (
            "Blok 1 — Başlatma ve Gövde",
            [
                ("Watch (45 sn)", "Elinizi uyluğunuzda bırakın ve ekrana bakın. Etkilenen kolunuzu izleyin. Gövde öne eğilmez; kol kendi başına uzanır."),
                ("Imagine (90 sn)", "Gözlerinizi kapatın. Gövde sabit; zihninizde hazır deyin. Sonra el yumuşakça ileri başlar; gövde sandalyede kalır."),
                ("Rest (45 sn)", "Gözlerinizi açın. Kolunuzu uyluğunuzda bırakın. Sadece dinlenin."),
            ],
        ),
        (
            "Blok 2 — Omuz Kontrolü",
            [
                ("Watch (45 sn)", "Omuz alçak ve gevşek kalır, kulaktan uzak."),
                ("Imagine (90 sn)", "Omuz ağır ve aşağıda; dirsek açılır, el öne uzanır; omuz kalkmaz."),
                ("Rest (45 sn)", "Gözlerinizi açın ve dinlenin."),
            ],
        ),
        (
            "Blok 3 — Dirsek + Kronometri",
            [
                ("Watch (45 sn)", "Dirsek yumuşakça açılır; kaslar sakin."),
                ("Imagine (90 sn)", "Dirsek açılır, el ulaşır, kısa durur, uyluğa döner. Frekans seslerini takip edin."),
                ("Rest (45 sn)", "Gözlerinizi açın ve dinlenin. Uzanma hızı rahat mıydı? Evet/hayır."),
            ],
        ),
        (
            "Blok 4 — Tam Ulaşma–Dönüş + Kronometri",
            [
                ("Watch (45 sn)", "El uylukta → ulaşma → duraklama → dönüş; önce vücut hazır, sonra kol."),
                ("Imagine (90 sn)", "Tam uzanma hayal edin; frekans seslerini takip edin; zihinde bir kez daha tekrarlayın."),
                ("Rest (45 sn)", "Gözlerinizi açın. Vücut hazır mıydı? Hız doğal mıydı? Evet/hayır. Seans bitti."),
            ],
        ),
    ]
    for title, steps in exp_blocks:
        para(doc, title, bold=True)
        for st, tx in steps:
            para(doc, f"{st}: {tx}")

    para(doc, "Frekans tonları (Blok 3–4): ulaşma (düşük) | duraklama (orta) | dönüş (farklı) | 4 sn sessizlik", bold=True)

    heading(doc, "KONTROL GRUBU — İmgeleme ve Zihinsel Arınma (~13 dk)", 2)
    bullets(
        doc,
        [
            "Motor imgeleme YOK; video YOK; frekans tonu YOK",
            "60 sn giriş + 4 blok × 3 dk (Zihinsel Arınma 45 sn + İmgeleme 90 sn + Dinlenme 45 sn)",
            "Dinlenme dönemlerinde soru sorulmaz",
        ],
    )

    para(doc, "Giriş — Zihin Hazırlığı (60 sn):", bold=True)
    para(
        doc,
        "Rahat oturun. Gözlerinizi kapatın. Derin nefes alın ve verin. Düşünceleri bulut gibi "
        "bırakın. Hareket, kol veya uzanma hayal etmeyin.",
    )

    ctrl_blocks = [
        (
            "Blok 1 — Nefes ve Zihinsel Arınma",
            [
                ("Zihinsel Arınma (45 sn)", "Nefes al… ver… Düşünceleri bırakın; hareket imgeleri gelirse nefese dönün."),
                ("İmgeleme (90 sn)", "Sakin açık alan — gökyüzü veya yumuşak ışık; sahne hareketsiz."),
                ("Dinlenme (45 sn)", "Gözlerinizi açın. Sadece dinlenin."),
            ],
        ),
        (
            "Blok 2 — Renk ve Işık",
            [
                ("Zihinsel Arınma (45 sn)", "Endişe ve planları bırakın; sadeleşin."),
                ("İmgeleme (90 sn)", "Sakin bir renk ve yumuşak ışık; vücut imgeleri yok."),
                ("Dinlenme (45 sn)", "Gözlerinizi açın ve dinlenin."),
            ],
        ),
        (
            "Blok 3 — Sakin Doğa",
            [
                ("Zihinsel Arınma (45 sn)", "Gerginliği fark edin ve bırakın."),
                ("İmgeleme (90 sn)", "Durgun göl, ağaçlar, bulutlar — donmuş sahne."),
                ("Dinlenme (45 sn)", "Gözlerinizi açın. Sadece dinlenin."),
            ],
        ),
        (
            "Blok 4 — Kapanış",
            [
                ("Zihinsel Arınma (45 sn)", "Motor düşünceleri bırakın; sakin nefes."),
                ("İmgeleme (90 sn)", "Blok 1 veya 3 sahnesini tekrarlayın; hareketsiz."),
                ("Dinlenme (45 sn)", "Gözlerinizi açın. Sadece dinlenin. Teşekkürler. Seans bitti."),
            ],
        ),
    ]
    for title, steps in ctrl_blocks:
        para(doc, title, bold=True)
        for st, tx in steps:
            para(doc, f"{st}: {tx}")

    heading(doc, "ÖLÇÜMLER VE İSTATİSTİK", 2)
    para(
        doc,
        "Ölçümler: MediaPipe + NeuroLab markerless kinematik; birincil sonuç smoothness_pause_pct; "
        "ikincil: total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity; "
        "WMFT-4, KVIQ-10, VAMS-4, IPAQ, MDRS, VAS.",
    )
    para(
        doc,
        "İstatistik: 2 (Grup) × 2 (Zaman) karma ANOVA; ITT; LOCF; Holm–Bonferroni.",
    )

    heading(doc, "DESTEK VE BÜTÇE BİLGİSİ", 1)
    para(doc, "Hayır — harici kurum desteği yok.")

    doc.add_paragraph()
    para(
        doc,
        "*** Bu formda yer alan araştırma projesine ait veri toplama yöntemi ve çalışma "
        "dizaynına sadık kalınacağını taahhüt ederim.",
    )
    doc.add_paragraph()
    para(doc, "Sorumlu Araştırmacı: Abdelrahman Walid Hamza Mohamed Elsayed Sabee")
    para(doc, "Tarih: …/…/2026")
    para(doc, "İmza: _________________")

    for folder in (OUT, OUT2):
        folder.mkdir(parents=True, exist_ok=True)
        path = folder / FILENAME
        doc.save(path)
        print("SAVED", path)


if __name__ == "__main__":
    build()
