# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document

p = Path(r"D:/Thesis app/ايتيك كرول/insan-arastirmalari-etik-kurul-basvuru-formu-DOLDURULMUS.docx")
d = Document(str(p))
out = Path(r"D:/Thesis app/NeuroLab/forms_completed/_etik_original_dump.txt")
lines = []
lines.append(f"PARAGRAPHS: {len(d.paragraphs)}")
lines.append(f"TABLES: {len(d.tables)}")
for i, para in enumerate(d.paragraphs):
    runs_info = []
    for r in para.runs:
        fn = r.font.name or "-"
        fs = r.font.size or "-"
        runs_info.append(f"[{fn}/{fs}]")
    lines.append(f"--- P{i} style={para.style.name} ---")
    lines.append(para.text)
    if runs_info:
        lines.append("RUNS: " + " ".join(runs_info[:5]))
for ti, table in enumerate(d.tables):
    lines.append(f"=== TABLE {ti} rows={len(table.rows)} cols={len(table.columns)} ===")
    for ri, row in enumerate(table.rows):
        cells = [c.text.replace("\n", " | ") for c in row.cells]
        lines.append(f"R{ri}: " + " || ".join(cells))
out.write_text("\n".join(lines), encoding="utf-8")
print("WROTE", out)
