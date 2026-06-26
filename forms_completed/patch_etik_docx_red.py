# -*- coding: utf-8 -*-
"""Patch full ethics docx (275 para) preserving layout; changed text in red."""
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

FOLDER = Path(r"D:\Thesis app\ايتيك كرول")
SRC = FOLDER / "insan-arastirmalari-etik-kurul-basvuru-formu-DOLDURULMUS.docx"
BACKUP = FOLDER / "insan-arastirmalari-etik-kurul-basvuru-formu-DOLDURULMUS.orijinal.bak.docx"
OUT = FOLDER / "insan-arastirmalari-etik-kurul-basvuru-formu-GUNCELLENMIS-SON.docx"
OUT2 = Path(r"D:\Thesis app\NeuroLab\forms_completed\insan-arastirmalari-etik-kurul-basvuru-formu-GUNCELLENMIS.docx")
RED = RGBColor(0xFF, 0x00, 0x00)

EXP_START = "Toplam doz: 2 dk kalibrasyon + 5"
EXP_END = "Müdahale Uyumu:"
CTRL_START = "Katılımcılar, motor olmayan bilişsel sistemi uyararak"
CTRL_END = "PETTLEP Çerçevesinin Uygulanması:"

# 20 slots: P160–P179 (same paragraph count as original)
EXP_LINES = [
    "Toplam doz: 60 sn kalibrasyon + 4 × 3 dk eğitim bloğu = ~13 dk. Her 3 dakikalık blok: Watch 45 sn + Imagine 90 sn + Rest 45 sn.",
    "",
    "Faz 1 — Kalibrasyon (60 sn):",
    "Amaç: Notice–name–transfer; etkilenmemiş kol ile bir kez yavaş ulaşma, hissi etkilenen tarafa aktarma.",
    'Script (kulaklık, sakin ses): "Rahat oturun, sırtınızı sandalyeye yaslayın. İki ayağınız yere bassın. Etkilenen elinizi avucunuz açık şekilde uyluğunuzda dinlendirin. Şimdi etkilenmemiş kolunuzla bir kez yavaşça öne uzanın — bir bardağa uzanır gibi. Nasıl hissettiğinize dikkat edin. Bu his için tek bir kelime seçin — sabit, kolay veya yumuşak — ve sessizce söyleyin. Gözlerinizi kapatın. Aynı kelimeyi ve hissi etkilenen kolunuza gönderin."',
    "",
    "Faz 2 — Eğitim Blokları (4 tekrar):",
    "Her blok aynı Watch → Imagine → Rest sırasını izler. Blok 1: Başlatma + gövde | Blok 2: Omuz kontrolü | Blok 3: Dirsek + kronometri | Blok 4: Tam ulaşma–dönüş + kronometri.",
    "",
    "Watch / Eylem Gözlemi (45 sn; Perspektif & Görev):",
    "Katılımcı, etkilenen tarafın Reach & Return (ulaşma–dönüş) hareketini gösteren ayna ters çevrilmiş birinci şahıs videosunu izler. Bloklara göre gövde sabitliği, omuz kontrolü, dirsek açılımı ve tam ulaş–duraklama–dönüş dizisi vurgulanır.",
    'Script (örnek, Blok 1): "Gözlerinizi açın ve videoyu izleyin. Elinizi uyluğunuzda bırakın. Gövde öne eğilmez; kol kendi başına uzanır — sakin ve net, hızlı değil. Doğal zamanlama: ulaş… bekle… dön."',
    "",
    "Imagine / Motor İmgeleme (90 sn; Fiziksel, Zamanlama & Duygu):",
    "Gözler kapalı; iç perspektif; gerçek zamanlı kinestetik imgeleme kritiktir. Blok 3–4'te frekans tonları ile ulaşma–duraklama–dönüş kronometrisi uygulanır (ulaşma: düşük ton; duraklama: orta ton; dönüş: farklı ton; 4 sn sessizlik).",
    'Script (örnek, Blok 4): "Gözlerinizi kapatın. Vücudunuz hazır, gövde sabit, omuz aşağı. El önünüze uzanır, kısa durur, yumuşakça uyluğunuza döner. Frekans seslerini takip edin. Zihninizde bir kez daha aynı sakin hızda tekrarlayın."',
    "",
    "Rest / Nöral Dinlenme (45 sn):",
    'Script: "Gözlerinizi açın. Kolunuzu uyluğunuzda bırakın. Normal nefes alın. Şimdi bir şey çalışmayın — sadece dinlenin." Blok 3–4\'te kronometri uyumu için kısa evet/hayır sorusu sorulabilir.',
    "",
]

# 4 slots before PETTLEP: title P183 + P184–P186
CTRL_LINES = [
    "Katılımcılar, motor olmayan imgeleme ve zihinsel arınma protokolünü takip edeceklerdir. Süre ve dikkat yükü deneysel grupla eşleştirilmiştir (~13 dk): 60 sn giriş + 4 blok × 3 dk (Zihinsel Arınma 45 sn + İmgeleme 90 sn + Dinlenme 45 sn). Video izleme, frekans tonları ve motor imgeleme uygulanmaz.",
    "Giriş — Zihin Hazırlığı (60 sn): Rahat oturun; nefes al–ver; düşünceleri bırakın. Motor imgeler yok. Blok 1–2: Zihinsel arınma (nefes, düşünce bırakma) + imgeleme (açık gökyüzü/ışık veya sakin renk–ışık); vücut, kol veya ulaşma imgeleri yasaktır.",
    "Blok 3–4: Sakin doğa sahnesi (durgun göl, ağaçlar; donmuş sahne) ve kapanış. Dinlenme dönemlerinde soru sorulmaz — katılımcı yalnızca dinlenir. Aynı kulaklık, oda, sandalye ve seans süresi deneysel grupla eşleştirilir.",
    "",
]

REPLACEMENTS = [
    ("bilişsel ve somatik kontrol grubuna", "imgeleme ve zihinsel arınma kontrol grubuna"),
    ("Kontrol Grubu (Bilişsel ve Somatik Kontrol):", "Kontrol Grubu (İmgeleme ve Zihinsel Arınma Kontrolü):"),
    ("Reach & Wipe — Ulaş-Sil", "Reach & Return — Ulaşma–Dönüş"),
    ("Reach & Wipe", "Reach & Return"),
    ("Ulaş-Sil", "Ulaşma–Dönüş"),
    ("yaklaşık 17 dakika", "yaklaşık 13 dakika"),
    ("(~17 dk)", "(~13 dk)"),
    ("~17 dk", "~13 dk"),
    ("= 17 dk", "= ~13 dk"),
    ("2 dk kalibrasyon + 5", "60 sn kalibrasyon + 4"),
    ("5 × 3 dk", "4 × 3 dk"),
    ("5 blok × 3 dk", "4 blok × 3 dk"),
    ("(5 tekrar)", "(4 tekrar)"),
    ("Eğitim Blokları (5 tekrar)", "Eğitim Blokları (4 tekrar)"),
    ("Imagine 75 sn", "Imagine 90 sn"),
    ("Rest 60 sn", "Rest 45 sn"),
    ("Rest / Nöral Dinlenme (60 sn)", "Rest / Nöral Dinlenme (45 sn)"),
    ("forward, wipe, return", "forward, pause, return"),
    ("ulaş… sil… dön", "ulaş… bekle… dön"),
    ("gerçek havlu altında el", "etkilenen el uylukta dinlenmiş"),
    ("hedef nesneyi (havlu)", "değerlendirme ortamını"),
    ("masaya ulaşıp silme hareketi", "uyluktan öne ulaşma ve uyluğa dönüş hareketi"),
    ("Silme Fazı (Pürüzsüzlük):", "Dönüş Fazı (Pürüzsüzlük):"),
    ("Yana silme, koordineli dirsek ve el hareketi gerektirir.", "Ulaşma sonrası yumuşak dönüş, koordineli dirsek fleksiyonu gerektirir."),
    ("Bir yüzeye ulaşıp temizleme yeteneği", "Uyluktan hedefe ulaşıp kontrollü dönüş yeteneği"),
    (
        "birincil sonuç (smoothness_pause_pct) değiştirilmemiştir.",
        "birincil sonuç (smoothness_pause_pct) değiştirilmemiştir. Ek revizyon: motor görev Reach & Return (ulaşma–dönüş); deneysel müdahale 4 blok × 3 dk (~13 dk); kontrol koşulu imgeleme + zihinsel arınma.",
    ),
    (
        "Tedavi, değerlendirmede kullanılan aynı masa, sandalye ve hedef nesneyi (havlu) kullanarak gerçekleşir.",
        "Tedavi, değerlendirmede kullanılan aynı sandalye ve oda düzeninde; tablet ekranında etkilenen tarafın ayna videosu ile gerçekleşir.",
    ),
    (
        "T – Zaman (Zamansallık — kritik): Katılımcılardan o anda hareket ettiklerini hayal etmeleri istenir. Zihinsel Kronometri, hayal edilen sürenin fiziksel sürenin %75'i (±%25) içinde olup olmadığı kontrol edilir.",
        "T – Zaman (Zamansallık — kritik): Gerçek zamanlı kinestetik motor imgeleme; Blok 3–4'te frekans tonları ile ulaşma–duraklama–dönüş kronometrisi. Zihinsel Kronometri, hayal edilen sürenin fiziksel sürenin ±%25 içinde olup olmadığı kontrol edilir.",
    ),
]

REF_NEW = "14. Di Rienzo F, et al. Motor imagery training in stroke: a systematic review. Front Hum Neurosci. 2016."


def apply_rules(text: str) -> str:
    out = text
    for old, new in REPLACEMENTS:
        out = out.replace(old, new)
    return out


def set_para_text_red(para, new_text: str) -> None:
    if para.runs:
        ref = para.runs[0]
        name = ref.font.name
        size = ref.font.size
        bold = ref.font.bold
        italic = ref.font.italic
    else:
        name = size = bold = italic = None

    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    run = para.runs[0] if para.runs else para.add_run()
    run.text = new_text
    run.font.color.rgb = RED
    if name:
        run.font.name = name
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def replace_block_range(paragraphs, start_key, end_key, lines):
    start_i = end_i = None
    for i, p in enumerate(paragraphs):
        if start_i is None and start_key in p.text:
            start_i = i
        elif start_i is not None and end_key in p.text:
            end_i = i
            break
    if start_i is None or end_i is None:
        return False
    slots = end_i - start_i
    for idx in range(slots):
        line = lines[idx] if idx < len(lines) else ""
        set_para_text_red(paragraphs[start_i + idx], line)
    return True


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    set_para_text_red(para, text)
    return para


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def verify(doc_path: Path) -> dict:
    doc = Document(str(doc_path))
    full = "\n".join(p.text for p in doc.paragraphs)
    red_count = 0
    for p in iter_all_paragraphs(doc):
        for run in p.runs:
            if run.font.color and run.font.color.rgb == RED:
                red_count += 1
                break
    return {
        "paragraphs": len(doc.paragraphs),
        "red_paragraphs": red_count,
        "wipe_left": full.count("Reach & Wipe") + full.count("Ulaş-Sil"),
        "return_count": full.count("Reach & Return"),
        "ctrl_ok": "zihinsel arınma" in full.lower(),
        "exp_ok": "60 sn kalibrasyon" in full and "4 × 3 dk" in full,
        "di_rienzo": "Di Rienzo" in full,
    }


def main():
    if not SRC.exists():
        print("SOURCE NOT FOUND:", SRC, file=sys.stderr)
        sys.exit(1)
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)

    shutil.copy2(BACKUP, OUT)
    doc = Document(str(OUT))
    paras = list(doc.paragraphs)

    ok_exp = replace_block_range(paras, EXP_START, EXP_END, EXP_LINES)
    ok_ctrl = replace_block_range(paras, CTRL_START, CTRL_END, CTRL_LINES)

    skip_keys = (EXP_START, CTRL_START)
    for para in iter_all_paragraphs(doc):
        if any(k in para.text for k in skip_keys):
            continue
        t = para.text
        if not t.strip():
            continue
        new_t = apply_rules(t)
        if new_t != t:
            set_para_text_red(para, new_t)

    full = "\n".join(p.text for p in doc.paragraphs)
    if "Di Rienzo" not in full:
        for para in doc.paragraphs:
            if para.text.startswith("13. Schuster"):
                insert_after(para, REF_NEW)
                break

    doc.save(str(OUT))
    shutil.copy2(OUT, OUT2)

    stats = verify(OUT)
    print("SAVED:", OUT)
    print("COPY :", OUT2)
    print("EXP block patched:", ok_exp)
    print("CTRL block patched:", ok_ctrl)
    for k, v in stats.items():
        print(f"  {k}: {v}")


if __name__ == "__main__":
    main()
