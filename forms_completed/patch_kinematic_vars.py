# -*- coding: utf-8 -*-
"""Update kinematic variables in ethics form (red text)."""
import shutil
import sys
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)

SRC = Path(r"D:\Thesis app\ETIK_ACIL\ETIK_KURUL_GUNCELLENMIS.docx")
WORK = Path(r"D:\Thesis app\ETIK_ACIL\_kinematic_patch_work.docx")
OUT = Path(r"D:\Thesis app\ETIK_ACIL\ETIK_KURUL_KINEMATIK.docx")
OUT2 = Path(r"D:\Thesis app\NeuroLab\forms_completed\ETIK_KURUL_KINEMATIK.docx")
ETH = Path(r"D:\Thesis app\ايتيك كرول\insan-arastirmalari-etik-kurul-basvuru-formu-GUNCELLENMIS-SON.docx")

# Exact paragraph replacements by index (verified on 276-para doc)
PARA_MAP = {
    36: (
        "Bu amaçla, MediaPipe Pose Landmarker tabanlı işaretsiz hareket yakalama sistemi kullanılacaktır. "
        "Kamera yan görünümde (90°) konumlandırılarak ulaşma kinematiği optimize edilecektir (8). "
        "Hareket pürüzsüzlüğü (SPARC), gövde kompanzasyonu (trunk_ratio), omuz elevasyonu (shoulder_vert_norm; "
        "ZoeDepth metrik ölçekleme), dirsek açısı (elbow_angle_mean), hareket süresi (movement_time_sec) ve "
        "tepe hız (peak_velocity_px_s) objektif olarak ölçülecektir. WMFT-4 ile kombinasyon, hareket kalitesinin "
        "fonksiyonel üst ekstremite performansı ile ilişkilendirilmesine olanak tanır (9)."
    ),
    43: (
        "Tek seanslık PETTLEP tabanlı AOMI; imgeleme ve zihinsel arınma kontrol grubuna kıyasla, "
        "inme geçiren bireylerde etkilenen üst ekstremitenin hareket pürüzsüzlüğü (SPARC), "
        "gövde kompanzasyonu (trunk_ratio), omuz kuşağı elevasyonu (shoulder_vert_norm), "
        "dirsek açısı, hareket süresi, tepe hız ve üst ekstremite fonksiyonunda (WMFT-4) "
        "anlık iyileşmelere yol açar mı?"
    ),
    67: (
        "REVİZYON (onaylanmış protokolde değişiklik): Çalışma çok merkezli hale getirilmiş; "
        "Biruni Üniversitesi Hastanesi ikinci veri toplama merkezi olarak eklenmiş ve Doç. Dr. Çiğdem Çınar "
        "yardımcı araştırmacı olarak ekibe dahil edilmiştir. Bilimsel tasarım ve örneklem (n=28) "
        "değiştirilmemiştir. Ek revizyon: motor görev Reach & Return; deneysel müdahale 4 blok × 3 dk (~13 dk); "
        "kontrol koşulu imgeleme + zihinsel arınma; birincil kinematik sonuç SPARC (Spectral Arc Length); "
        "ikincil sonuçlar trunk_ratio, shoulder_vert_norm (ZoeDepth), elbow_angle_mean, movement_time_sec, "
        "peak_velocity_px_s."
    ),
    115: (
        "Anlık kinematik değişiklikleri değerlendirmek amacıyla, tek eğitim seansının öncesinde ve hemen "
        "sonrasında kayıtlar alınacaktır. Sensör tutarlılığını sağlamak için tüm katılımcılarda aynı akıllı "
        "telefon veya web kamerası modeli kullanılacaktır. Kamera, sabit bir yüksekliğe monte edilen tripod "
        "üzerine yerleştirilecek ve katılımcıya göre yan görünümde (90°) yaklaşık 1,5–2,0 m standart mesafede "
        "konumlandırılacaktır. Kayıtlar 1080p çözünürlükte ve saniyede 30 kare (fps) yapılacaktır. Yapay zeka "
        "takibini engelleyebilecek arkadan aydınlatma veya gölgeleri önlemek için kayıtlar tutarlı, dağınık "
        "aydınlatmaya sahip bir odada gerçekleştirilecektir. Video verileri MediaPipe Pose Landmarker "
        "(33 landmark, 2D) ile işlenecek; omuz elevasyonu için ZoeDepth monoküler derinlik ağı ile omuz "
        "genişliği metrik ölçeklemesi uygulanacaktır."
    ),
    118: (
        "Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişimdir (SPARC). "
        "MediaPipe tabanlı pipeline akıllı telefon veya web kamerası videosundan avuç yörüngesi türetir (8). "
        "SPARC (Spectral Arc Length — Spektral Ark Uzunluğu): Avuç merkezinin X–Y yörüngesinden teğetsel hız "
        "profili hesaplanır; hız sinyali normalize edilir ve hızlı Fourier dönüşümü (FFT) ile frekans spektrumu "
        "elde edilir (fc ≤ 10 Hz). Spektral eğrinin ark uzunluğu alınır (Balasubramanian et al., 2012, 2015). "
        "Daha negatif SPARC değerleri daha pürüzsüz, kesintisiz hareketi; daha az negatif (veya pozitif) "
        "değerler stop-and-go davranışını yansıtır."
    ),
    121: "Gövde Kompanzasyonu — Trunk Ratio (Telafi Edici Strateji):",
    122: (
        "Amaç: Hastanın azalmış kol uzunluğunu telafi etmek için gövde yer değiştirmesine ne derece "
        "başvurduğunu nicelleştirmek."
    ),
    123: (
        "Ölçüm Yöntemi (trunk_ratio): Aktif hareket penceresinde (avuç hızı > hız eşiği) gövde landmark'ının "
        "başlangıç–bitiş yer değiştirmesinin, aynı pencerede avuç yer değiştirmesine oranı olarak hesaplanır "
        "(Wagh et al., 2025; Schwarz et al., 2022)."
    ),
    125: "Omuz Kuşağı Elevasyonu — Shoulder Elevation (Telafi Edici Strateji):",
    126: (
        "Amaç: Omuz kuşağının yukarı kalkması (telafi edici elevasyon) düzeyini ölçmek."
    ),
    127: (
        "Ölçüm Yöntemi (shoulder_vert_norm): ZoeDepth monoküler derinlik ağı ile elde edilen omuz genişliği "
        "metrik ölçeği kullanılarak, hareket penceresinde etkilenen omuz landmark'ının dikey (Y ekseni) "
        "yer değiştirmesi omuz genişliğine normalize edilir (mevcut NeuroLab yöntemi). Yan görünüm "
        "piksel tabanlı omuz analizi bu değişken için kullanılmaz."
    ),
    129: (
        "Diğer ikincil kinematik değişkenler: (4) Dirsek Açısı (elbow_angle_mean) — omuz–dirsek–bilek "
        "vektörleri arasındaki açının hareket penceresinde frame bazında hesaplanıp ortalamasının alınması; "
        "yan görünümde yüksek doğruluk. (5) Hareket Süresi (movement_time_sec) — avuç hızının eşik üstünde "
        "kaldığı süre (onset–offset), saniye cinsinden. (6) Tepe Hız (peak_velocity_px_s) — hareket "
        "penceresinde avuç teğetsel hızının maksimumu (px/s)."
    ),
    156: (
        "Temel Değerlendirme: Sistem kalibrasyonundan sonra, her denek tarafından Reach & Return görevinin "
        "üç aktif tekrarı gerçekleştirilecektir. Hareket kinematikleri, MediaPipe tabanlı işaretsiz hareket "
        "yakalama sistemi ile kaydedilecektir (müdahale öncesi hareket kalitesi: SPARC ve trunk_ratio)."
    ),
    202: (
        "Uzanma ve Gövde Fazı: Omuz fleksiyonu ve dirsek ekstansiyonu mesafe bileşenine katkıda bulunur. "
        "Gövde Kompanzasyonu ve Omuz Kuşağı Elevasyonu, inme sonrası hastaların bu aşamasında gerçekten çok "
        "sık görülen telafi edici davranışlardır. İmgeleme scriptleri, gövdeyi sabit tutmayı ve omuzları "
        "rahat bırakmayı vurgular; bu, trunk_ratio ve shoulder_vert_norm ile nicel olarak ölçülür."
    ),
    203: (
        "Dönüş Fazı (Pürüzsüzlük): Ulaşma sonrası yumuşak dönüş, koordineli dirsek fleksiyonu gerektirir. "
        "Gerçek zamanlı imgeleme ve \"rahatlık\" vurgusu, hareket pürüzsüzlüğü ve doğruluğu üzerinde hassas "
        "motor kontrolü eğitmeyi hedefler; SPARC, elbow_angle_mean ve peak_velocity_px_s ile ölçülür."
    ),
}

GLOBAL_REPLACEMENTS = [
    ("smoothness_pause_pct", "SPARC"),
    ("total_trunk_palm_ratio", "trunk_ratio"),
    ("total_duration_s", "movement_time_sec"),
    ("total_peak_velocity", "peak_velocity_px_s"),
    ("Pause Time %", "Spectral Arc Length (SPARC)"),
    ("birincil sonuç (smoothness_pause_pct)", "birincil sonuç (SPARC)"),
    ("Birincil sonuç smoothness_pause_pct", "Birincil sonuç SPARC"),
    ("ikincil kinematik değişken ailesi (k = 8)", "ikincil kinematik değişken ailesi (k = 5)"),
    (
        "smoothness_pause_pct ve total_trunk_palm_ratio Δ skorları",
        "SPARC ve trunk_ratio Δ skorları",
    ),
    (
        "frontal açıda yaklaşık 1.5–2.0 m",
        "yan görünümde (90°) yaklaşık 1,5–2,0 m",
    ),
]


def set_red(para, text: str) -> None:
    if para.runs:
        ref = para.runs[0]
        name, size, bold, italic = ref.font.name, ref.font.size, ref.font.bold, ref.font.italic
    else:
        name = size = bold = italic = None
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    run = para.runs[0] if para.runs else para.add_run()
    run.text = text
    run.font.color.rgb = RED
    if name:
        run.font.name = name
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def apply_global(text: str) -> str:
    out = text
    for old, new in GLOBAL_REPLACEMENTS:
        out = out.replace(old, new)
    return out


def patch_doc(path: Path) -> None:
    doc = Document(str(path))
    done = set()
    for idx, new_text in PARA_MAP.items():
        if idx < len(doc.paragraphs):
            set_red(doc.paragraphs[idx], new_text)
            done.add(idx)

    for i, para in enumerate(doc.paragraphs):
        if i in done or not para.text.strip():
            continue
        new_t = apply_global(para.text)
        if new_t != para.text:
            set_red(para, new_t)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.text.strip():
                        continue
                    new_t = apply_global(para.text)
                    if new_t != para.text:
                        set_red(para, new_t)

    doc.save(str(path))


def main():
    src = SRC if SRC.exists() else OUT2
    if not src.exists():
        print("Missing source", file=sys.stderr)
        sys.exit(1)
    shutil.copy2(src, WORK)
    patch_doc(WORK)
    shutil.copy2(WORK, OUT)
    shutil.copy2(WORK, OUT2)
    print("OK", OUT)


if __name__ == "__main__":
    main()
