# -*- coding: utf-8 -*-
"""Regenerate ethics form + export docx/doc/pdf/html to simple ASCII paths."""
import shutil
import subprocess
import sys
from pathlib import Path

# Run patch logic
sys.path.insert(0, str(Path(__file__).parent))
import patch_etik_docx_red as patch

OUT_DIR = Path(r"D:\Thesis app\ETIK_ACIL")
OUT_DIR.mkdir(parents=True, exist_ok=True)

NAMES = {
    "docx": OUT_DIR / "ETIK_KURUL_GUNCELLENMIS.docx",
    "doc": OUT_DIR / "ETIK_KURUL_GUNCELLENMIS.doc",
    "pdf": OUT_DIR / "ETIK_KURUL_GUNCELLENMIS.pdf",
    "html": OUT_DIR / "ETIK_KURUL_GUNCELLENMIS.html",
    "bat": OUT_DIR / "AC_ETIK_FORMU.bat",
}


def export_html(docx_path: Path, html_path: Path) -> None:
    from docx import Document

    doc = Document(str(docx_path))
    parts = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            parts.append("<br>")
            continue
        red = any(
            r.font.color and r.font.color.rgb and str(r.font.color.rgb).upper() == "FF0000"
            for r in p.runs
        )
        esc = (
            t.replace("&", "&amp;")
            .replace("<", "&lt;")
            .replace(">", "&gt;")
        )
        tag = f'<p style="color:#c00">{esc}</p>' if red else f"<p>{esc}</p>"
        parts.append(tag)

    html = (
        "<!DOCTYPE html><html><head><meta charset='utf-8'>"
        "<title>Etik Kurul Formu - Guncellenmis</title>"
        "<style>body{font-family:Calibri,Segoe UI,Arial;max-width:920px;"
        "margin:32px auto;line-height:1.55;padding:24px;font-size:11pt}"
        "p{margin:0.35em 0}</style></head><body>"
        + "\n".join(parts)
        + "</body></html>"
    )
    html_path.write_text(html, encoding="utf-8")


def export_word_formats(docx_src: Path) -> None:
    import win32com.client

    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = 0
    try:
        doc = word.Documents.Open(str(docx_src.resolve()), ReadOnly=False)
        doc.SaveAs2(str(NAMES["docx"].resolve()), FileFormat=16)  # docx
        doc.SaveAs2(str(NAMES["doc"].resolve()), FileFormat=0)  # .doc
        doc.ExportAsFixedFormat(str(NAMES["pdf"].resolve()), 17)  # pdf
        doc.Close(False)
    finally:
        word.Quit()


def main():
    # Patch from original backup
    if not patch.SRC.exists():
        raise FileNotFoundError(patch.SRC)
    if not patch.BACKUP.exists():
        shutil.copy2(patch.SRC, patch.BACKUP)

    work = OUT_DIR / "_work_patch.docx"
    shutil.copy2(patch.BACKUP, work)

    from docx import Document

    doc = Document(str(work))
    paras = list(doc.paragraphs)
    patch.replace_block_range(paras, patch.EXP_START, patch.EXP_END, patch.EXP_LINES)
    patch.replace_block_range(paras, patch.CTRL_START, patch.CTRL_END, patch.CTRL_LINES)
    skip_keys = (patch.EXP_START, patch.CTRL_START)
    for para in patch.iter_all_paragraphs(doc):
        if any(k in para.text for k in skip_keys):
            continue
        t = para.text
        if not t.strip():
            continue
        new_t = patch.apply_rules(t)
        if new_t != t:
            patch.set_para_text_red(para, new_t)
    full = "\n".join(p.text for p in doc.paragraphs)
    if "Di Rienzo" not in full:
        for para in doc.paragraphs:
            if para.text.startswith("13. Schuster"):
                patch.insert_after(para, patch.REF_NEW)
                break
    doc.save(str(work))

    src = work

    shutil.copy2(src, NAMES["docx"])
    export_word_formats(NAMES["docx"])
    export_html(NAMES["docx"], NAMES["html"])

    NAMES["bat"].write_text(
        '@echo off\r\n'
        f'start "" "{NAMES["doc"]}"\r\n',
        encoding="utf-8",
    )

    # Also copy to ethics folder
    eth = Path(r"D:\Thesis app\ايتيك كرول")
    if eth.exists():
        for ext, p in [("docx", NAMES["docx"]), ("doc", NAMES["doc"]), ("pdf", NAMES["pdf"])]:
            dst = eth / f"insan-arastirmalari-etik-kurul-basvuru-formu-GUNCELLENMIS-SON.{ext}"
            shutil.copy2(p, dst)

    log = OUT_DIR / "OK.txt"
    log.write_text(
        "\n".join(f"{k}: {v} ({v.stat().st_size} bytes)" for k, v in NAMES.items() if v.exists()),
        encoding="utf-8",
    )
    subprocess.Popen(["explorer", str(OUT_DIR)])


if __name__ == "__main__":
    main()
