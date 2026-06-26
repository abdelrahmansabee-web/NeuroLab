# -*- coding: utf-8 -*-
"""Add SPARC/MediaPipe references to kinematic ethics form (red text)."""
import shutil
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.shared import RGBColor
from docx.text.paragraph import Paragraph

RED = RGBColor(0xFF, 0x00, 0x00)
SRC = Path(r"D:\Thesis app\NeuroLab\forms_completed\ETIK_KURUL_KINEMATIK.docx")
WORK = Path(r"D:\Thesis app\ETIK_ACIL\_refs_patch_work.docx")
OUT = Path(r"D:\Thesis app\ETIK_ACIL\ETIK_KURUL_KINEMATIK.docx")
OUT2 = Path(r"D:\Thesis app\NeuroLab\forms_completed\ETIK_KURUL_KINEMATIK.docx")

NEW_REFS = [
    "15. Balasubramanian S, Melendez-Calderon A, Burdet E. A robust and sensitive metric for quantifying movement smoothness. IEEE Trans Biomed Eng. 2012;59(8):2126-2136.",
    "16. Balasubramanian S, Melendez-Calderon A, Roby-Brami A, Burdet E. On the analysis of movement smoothness. J NeuroEngineering Rehabil. 2015;12:112.",
    "17. Refai MIM, Saes M, Scheltinga BL, van Kordelaar J, Bussmann JBJ, Veltink PH, et al. Smoothness metrics for reaching performance after stroke. Part 1: which one to choose? J Neuroeng Rehabil. 2021;18(1):154.",
    "18. Saes M, Refai MIM, van Kordelaar J, Scheltinga BL, van Beijnum B-JF, Bussmann JB, et al. Smoothness metric during reach-to-grasp after stroke. Part 2. Longitudinal association with motor impairment. J Neuroeng Rehabil. 2021;18(1):144.",
]

PARA_UPDATES = {
    36: (
        "Bu amaçla, MediaPipe Pose Landmarker tabanlı işaretsiz hareket yakalama sistemi kullanılacaktır. "
        "Kamera yan görünümde (90°) konumlandırılarak ulaşma kinematiği optimize edilecektir (8). "
        "Hareket pürüzsüzlüğü (SPARC; 15,16,17,18), gövde kompanzasyonu (trunk_ratio; 8), "
        "omuz elevasyonu (shoulder_vert_norm; 8, ZoeDepth metrik ölçekleme), dirsek açısı (elbow_angle_mean), "
        "hareket süresi (movement_time_sec) ve tepe hız (peak_velocity_px_s) objektif olarak ölçülecektir. "
        "WMFT-4 ile kombinasyon, hareket kalitesinin fonksiyonel üst ekstremite performansı ile "
        "ilişkilendirilmesine olanak tanır (9)."
    ),
    118: (
        "Çalışmanın birincil sonucu, etkilenen üst ekstremitenin hareket pürüzsüzlüğündeki değişimdir (SPARC). "
        "MediaPipe tabanlı pipeline akıllı telefon veya web kamerası videosundan avuç yörüngesi türetir (8). "
        "SPARC (Spectral Arc Length — Spektral Ark Uzunluğu): Avuç merkezinin X–Y yörüngesinden teğetsel hız "
        "profili hesaplanır; hız sinyali normalize edilir ve hızlı Fourier dönüşümü (FFT) ile frekans "
        "spektrumu elde edilir (fc ≤ 10 Hz). Spektral eğrinin ark uzunluğu alınır (15,16). "
        "İnme sonrası ulaşma görevlerinde değerlendirilen 32 pürüzsüzlük metriği arasından yalnızca SPARC'nin "
        "matematiksel olarak geçerli bulunduğu gösterilmiştir (17). SPARC'nin FM-UE ile uzunlamasına pozitif "
        "ilişkisi doğrulanmıştır (18). Daha negatif SPARC değerleri daha pürüzsüz, kesintisiz hareketi; "
        "daha az negatif (veya pozitif) değerler stop-and-go davranışını yansıtır."
    ),
    123: (
        "Ölçüm Yöntemi (trunk_ratio): Aktif hareket penceresinde (avuç hızı > hız eşiği) gövde landmark'ının "
        "başlangıç–bitiş yer değiştirmesinin, aynı pencerede avuç yer değiştirmesine oranı olarak hesaplanır "
        "(8,3)."
    ),
    127: (
        "Ölçüm Yöntemi (shoulder_vert_norm): ZoeDepth monoküler derinlik ağı ile elde edilen omuz genişliği "
        "metrik ölçeği kullanılarak, hareket penceresinde etkilenen omuz landmark'ının dikey (Y ekseni) "
        "yer değiştirmesi omuz genişliğine normalize edilir (mevcut NeuroLab yöntemi; telafi edici omuz "
        "elevasyonu MediaPipe ile gösterilmiştir — 8). Yan görünüm piksel tabanlı omuz analizi bu değişken "
        "için kullanılmaz."
    ),
}


def set_red(para, text: str) -> None:
    if para.runs:
        ref = para.runs[0]
        name, size, bold, italic = ref.font.name, ref.font.size, ref.font.bold, ref.font.italic
    else:
        name = size = bold = italic = None
    for run in para.runs[1:]:
        run._element.getparent().remove(run._element)
    run = para.runs[0] if para.runs else para.add_run()
    run.text = text
    run.font.color.rgb = RED
    if name:
        run.font.name = name
    if size:
        run.font.size = size
    if bold is not None:
        run.font.bold = bold
    if italic is not None:
        run.font.italic = italic


def insert_after(paragraph, text):
    new_p = OxmlElement("w:p")
    paragraph._p.addnext(new_p)
    para = Paragraph(new_p, paragraph._parent)
    set_red(para, text)
    return para


def main():
    shutil.copy2(SRC, WORK)
    doc = Document(str(WORK))

    full = "\n".join(p.text for p in doc.paragraphs)
    if "Refai MIM" not in full:
        anchor = None
        for para in doc.paragraphs:
            if para.text.strip().startswith("14. Di Rienzo"):
                anchor = para
                break
        if anchor:
            prev = anchor
            for ref in NEW_REFS:
                prev = insert_after(prev, ref)

    for idx, text in PARA_UPDATES.items():
        if idx < len(doc.paragraphs):
            set_red(doc.paragraphs[idx], text)

    doc.save(str(WORK))
    shutil.copy2(WORK, OUT)
    shutil.copy2(WORK, OUT2)
    print("saved", OUT2)


if __name__ == "__main__":
    main()
