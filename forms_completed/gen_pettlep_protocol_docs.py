# -*- coding: utf-8 -*-
"""Generate updated Turkish PETTLEP Reach-Return protocol + ethics form section text."""
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

OUT_DIR = Path(r"D:\Thesis app\manuscript f\REVIZYON_PAKETI")
OUT_DIR.mkdir(parents=True, exist_ok=True)

PROTOCOL_DOCX = OUT_DIR / "PETTLEP_Reach_Return_Protokol_TR.docx"
CONTROL_DOCX = OUT_DIR / "Kontrol_Grubu_Motor_Olmayan_Imgeleme_TR.docx"
ETHICS_TXT = OUT_DIR / "ETIK_FORM_UYGULANACAK_YAKLASIM_GUNCELLENMIS.txt"

ETHICS_SECTION = """Tasarım: Tek kör, ön test–son test, prospektif paralel grup randomize kontrollü çalışma (RCT); çok merkezli (İstinye Üniversitesi Liv Bahçeşehir Hastanesi Nörorehabilitasyon Kliniği + Biruni Üniversitesi Hastanesi FTR Polikliniği).

Örneklem: n=28 (grup başına 14); G*Power ile planlanmış; permütasyon blok randomizasyon (1:1); cinsiyet ve MAS (0–1+ / 2) stratifikasyonu; körleme: katılımcılar kör olamaz, kinematik çıkarımı otomatik/kör.

Görev: Çapraz vücut ulaşma–dönüş (Reach & Return); etkilenen üst ekstremite ile uyluktan öne ulaşma, kısa duraklama ve uyluğa dönüş. Pre/post değerlendirmede üç deneme, ortalama analiz. Müdahale sırasında fiziksel hareket uygulanmaz.

PETTLEP çerçevesi (Holmes & Collins, 2001):
P (Physical): Gerçek oturma düzeni, normal kıyafet, değerlendirme ile aynı sandalye ve oda düzeni.
E (Environment): Değerlendirme odası ile aynı ortam; tablet ekranında etkilenen tarafın ayna videosu.
T (Task): Çapraz vücut ulaşma–dönüş (Reach & Return); günlük yaşama uygun, tek akıcı hareket dizisi.
T (Timing): Gerçek zamanlı kinestetik motor imgeleme; Blok 3–4'te frekans tonları ile ulaşma–duraklama–dönüş kronometrisi.
L (Learning): Bloklar arası kısa düzeltici rehberlik (gövde sabitliği, omuz kontrolü, dirsek açılımı).
E (Emotion): Rahatlık, güven, doğal kas hissi; "hazır" anından sonra ulaşma.
P (Perspective): Watch aşamasında dış perspektif (3. kişi video); Imagine aşamasında iç perspektif (1. kişi kinestetik MI).

Deneysel müdahale (PETTLEP-AOMI; toplam ~13 dk; müdahale sırasında fiziksel uygulama YOK):
• Kalibrasyon: 60 sn (notice–name–transfer; etkilenmemiş kol ile bir kez yavaş ulaşma).
• 4 eğitim bloku × 3 dk: Watch 45 sn + Imagine 90 sn + Rest 45 sn.
• Blok 1: Başlatma + gövde | Blok 2: Omuz kontrolü | Blok 3: Dirsek + kronometri | Blok 4: Tam ulaşma–dönüş + kronometri.
• Blok 3–4'te frekans tonları: ulaşma (düşük), duraklama (orta), dönüş (farklı ton), 4 sn sessizlik.
• Ayrıntılı Türkçe scriptler ek protokol dosyasında (PETTLEP_Reach_Return_Protokol_TR.docx).
• Fiziksel Reach & Return yalnızca pre/post değerlendirmede.

Kontrol koşulu (imgeleme + zihinsel arınma; motor imgeleme YOK; zaman eşleştirilmiş ~13 dk):
• 60 sn giriş: nefes ve zihin sakinleştirme (motor imgeler YOK).
• 4 blok × 3 dk: Zihinsel Arınma 45 sn + İmgeleme 90 sn + Dinlenme 45 sn (dinlenme dönemlerinde soru sorulmaz).
• Video izleme ve frekans tonları YOK; aynı kulaklık, oda, sandalye ve seans süresi.

Kontrol grubu scriptleri (Türkçe):

Giriş — Zihin Hazırlığı (60 sn): Rahat oturun, sırtınızı sandalyeye yaslayın. İki ayağınız yere bassın. Ellerinizi uyluklarınızda dinlendirin. Gözlerinizi yavaşça kapatın. Derin bir nefes alın… ve yavaşça verin. Omuzlarınızı gevşetin. Zihninizdeki düşünceleri bir bulut gibi düşünün — gelir, geçer; onları itmeyin, sadece bırakın. Vücudunuzda hareket, kol veya uzanma hayal etmeyin. Sadece nefesinize ve sakinliğe odaklanın.

Blok 1 — Nefes ve Zihinsel Arınma:
Zihinsel Arınma (45 sn): Gözleriniz kapalı. Nefesinize dikkat edin: nefes al… nefes ver. Her nefes verişte zihniniz biraz daha sakinleşsin. Aklınıza gelen düşünceleri isimlendirmeden bırakın — sanki nehirde akan yapraklar gibi geçsin gitsin. Kol, omuz veya hareket imgeleri gelirse, yumuşakça nefese geri dönün.
İmgeleme (90 sn): Gözleriniz kapalı. Zihninizde sakin, açık bir alan hayal edin — belki açık bir gökyüzü, belki yumuşak bir ışık. Renklere ve ışığa odaklanın; sahne tamamen hareketsiz. Vücudunuz sandalyede; sadece zihninizdeki görüntü değişiyor. Ulaşma, yürüme veya kol hareketi DÜŞÜNMEYİN.
Dinlenme (45 sn): Gözlerinizi açın. Birkaç normal nefes alın. Şimdi bir şey hayal etmeyin — sadece dinlenin.

Blok 2 — Renk ve Işık İmgeleme:
Zihinsel Arınma (45 sn): Gözlerinizi kapatın. Omuzlarınızı bir kez daha bırakın. Nefes al… ver… Zihninizi sadeleştirin: bugünkü endişeleri, planları, görevleri bir kenara bırakın. Bu birkaç dakika yalnızca sizin sakin alanınız.
İmgeleme (90 sn): Zihninizde yavaşça bir renk belirsin — mavi, yeşil veya altın sarısı, hangisi size sakin geliyorsa. Rengin tonlarını, derinliğini hayal edin. Renk yavaşça yumuşak bir ışığa dönüşsün. Her şey sabit ve hareketsiz; sadece renk ve ışık var. Vücut imgeleri kullanmayın.
Dinlenme (45 sn): Gözlerinizi açın ve dinlenin.

Blok 3 — Sakin Doğa Sahnesi:
Zihinsel Arınma (45 sn): Gözleriniz kapalı. Nefesiniz doğal ritminde aksın. Zihninizde biriken gerginliği fark edin — yargılamadan, sadece fark edin — ve nefes verirken bırakın. Zihinsel arınma: düşünceler geçer, siz kalırsınız.
İmgeleme (90 sn): Zihninizde sakin bir doğa sahnesi canlandırın: durgun bir göl, yeşil ağaçlar, hafif bulutlar. Suyun yüzeyi, ağaçların rengi, gökyüzünün tonu — ayrıntılara odaklanın. Sahne donmuş bir fotoğraf gibi; rüzgâr, yürüme veya el hareketi yok.
Dinlenme (45 sn): Gözlerinizi açın. Normal nefes alın. Şimdi bir şey hayal etmeyin — sadece dinlenin.

Blok 4 — Tam Zihinsel Arınma ve Kapanış:
Zihinsel Arınma (45 sn): Gözleriniz kapalı. Derin nefes al… yavaş ver… Zihninizi temizleyin: motor düşünceler, görevler, hareket planları — hepsini şimdilik bırakın. Sadece sakin nefes ve boş, huzurlu bir zihin.
İmgeleme (90 sn): Blok 1'deki açık alanı veya Blok 3'teki sakin sahneyi zihninizde tekrar canlandırın. Renk, ışık ve sessizliğe odaklanın. Kol, omuz, uzanma, yürüme — hiçbiri yok. Zihniniz sakin, sahne hareketsiz.
Dinlenme (45 sn): Gözlerinizi açın. Normal nefes alın. Hiçbir şey hayal etmeyin — sadece dinlenin. Teşekkürler. Seans bitti.

Ölçümler: MediaPipe + ZoeDepth (NeuroLab); birincil sonuç smoothness_pause_pct; ikincil kinematikler (total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity); WMFT-4, KVIQ-10, VAMS-4, IPAQ, MDRS, VAS.

İstatistik: 2 (Grup) × 2 (Zaman) karma ANOVA; ITT; LOCF; Holm–Bonferroni (ikincil kinematik ailesi)."""


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE
    p.paragraph_format.space_after = Pt(6)
    return p


def build_experimental_protocol():
    doc = Document()
    add_heading(doc, "PETTLEP Temelli Eylem Gözlemi ve Motor İmgeleme (AOMI)", 0)
    add_heading(doc, "Çapraz Vücut Ulaşma–Dönüş (Reach & Return) Protokolü", level=1)
    add_para(doc, "Seans süresi: ~13 dk | Kalibrasyon: 60 sn | 4 blok × 3 dk (Watch 45 sn + Imagine 90 sn + Rest 45 sn)")

    add_heading(doc, "Kalibrasyon (60 sn)", level=2)
    add_para(doc, "Rahat oturun, sırtınızı sandalyeye yaslayın. İki ayağınız yere bassın. Etkilenen elinizi avucunuz açık şekilde uyluğunuzda dinlendirin. Şimdi etkilenmemiş kolunuzla bir kez yavaşça öne uzanın — bir bardağa uzanır gibi. Nasıl hissettiğinize dikkat edin. Belki vücudunuz sabit kalır, belki başlangıç kolay gelir. Bu his için tek bir kelime seçin — sabit, kolay veya yumuşak — ve sessizce söyleyin. Gözlerinizi kapatın. Aynı kelimeyi ve hissi etkilenen kolunuza gönderin; sanki o kol da aynı kalitede hareket edebilirmiş gibi.")

    blocks = [
        ("Blok 1 — Başlatma ve Gövde (Initiation + Trunk)", [
            ("Watch (45 sn)", "Elinizi uyluğunuzda bırakın ve ekrana bakın. Kendi kolunuzu — etkilenen tarafı — izliyorsunuz. Önce gövdenin sandalyede nasıl sessiz kaldığına bakın. Gövde öne eğilmez. Sonra el hareket etmeye başlar — sakin ve net, hızlı değil. Kol kendi başına uzanır; vücut onun peşinden gitmez."),
            ("Imagine (90 sn)", "Gözlerinizi kapatın. Ayaklar yerde, el uyluğunuzda. Önce vücudunuzun hazırlandığını hissedin — gövde sabit, sırt destekli. Zihninizde hazır deyin. Hareket etmeden önce vücut hazırlanır. Sonra, o hazırlık anından sonra eliniz yumuşakça ileri başlar. Gövdeniz sandalyede geride kalır; sadece kolunuz uzanır. Vücudunuzun öne eğildiğini hayal ederseniz durun ve hazırdan tekrar başlayın. Bir süre bu hisste kalın — önce hazırlık, sonra uzanma."),
            ("Rest (45 sn)", "Gözlerinizi açın. Kolunuzu uyluğunuzda bırakın. Normal nefes alın. Şimdi bir şey çalışmayın — sadece dinlenin."),
        ]),
        ("Blok 2 — Omuz Kontrolü (Shoulder Control)", [
            ("Watch (45 sn)", "Ekrandaki omuza bakın. Kol öne uzanırken omuz alçak ve gevşek kalır, kulaktan uzak."),
            ("Imagine (90 sn)", "Gözlerinizi kapatın. Omuğunuzun ağır ve aşağıda olduğunu hayal edin — önce yerleşsin. Sonra dirseğiniz açılır ve eliniz öne uzanır. Omuz kulağa kalkmaz; kol hareket ederken sakin ve alçak kalır."),
            ("Rest (45 sn)", "Gözlerinizi açın ve dinlenin. Şimdilik hayal kurmayın."),
        ]),
        ("Blok 3 — Dirsek ve Kronometri (Elbow + Chronometry)", [
            ("Watch (45 sn)", "El ilerlerken dirseğin nasıl yumuşakça açıldığını izleyin. Kaslar sakin görünür — sert değil, zorlanmış değil."),
            ("Imagine (90 sn)", "Gözlerinizi kapatın. Omuz aşağıda. Kolunuzun önünün yumuşak ve sakin olduğunu hissedin. Dirseğiniz yumuşakça açılır, eliniz önünüzdeki noktaya uzanır, kısa durur, sonra yumuşakça uyluğunuza döner. Lütfen duyduğunuz frekans seslerini takip edin ve hareketi bu ritimle eş zamanlı olarak zihninizde canlandırın."),
            ("Rest + Kronometri (45 sn)", "Gözlerinizi açın ve dinlenin. Kısa bir soru: o uzanmayı hayal ederken, gerçekten rahat bir uzanma ile aynı hızda mı hissettiniz — çok hızlı değil, donmuş değil? Evet veya hayır diyebilirsiniz."),
        ]),
        ("Blok 4 — Tam Ulaşma–Dönüş (Full Reach + Chronometry)", [
            ("Watch (45 sn)", "Tüm hareketi izleyin: el uylukta, öne uzanma, kısa duraklama, uyluğa dönüş. Tek akıcı sıra — önce vücut hazır, sonra kol."),
            ("Imagine (90 sn)", "Gözlerinizi kapatın. Aynı sandalye, aynı duruş. Bir tam uzanma hayal edin: vücudunuz hazır ve gövde sabit, omuz aşağı, dirsek yumuşakça açılıyor, el önünüzdeki noktaya uzanıyor, kısa bir duraklama, sonra yumuşakça uyluğunuza dönüş. Zihninizde bir kez daha yapın — aynı sıra, aynı sakin hız. Lütfen duyduğunuz frekans seslerini takip edin ve hareketi bu ritimle eş zamanlı olarak zihninizde canlandırın."),
            ("Rest + Kronometri (45 sn)", "Gözlerinizi açın ve dinlenin. İki kısa soru. Birincisi: kolunuz hareket etmeden önce vücudunuz hazır hissetti mi? İkincisi: tüm uzanma gerçek zamanlı hızda mıydı — doğal, aceleci değil? Her biri için evet veya hayır. Teşekkürler. Seans bitti."),
        ]),
    ]

    for title, steps in blocks:
        add_heading(doc, title, level=2)
        for step_title, step_text in steps:
            add_para(doc, step_title, bold=True)
            add_para(doc, step_text)

    add_heading(doc, "Frekans Tonları (Blok 3–4)", level=2)
    add_para(doc, "Ulaşma (uzanma): düşük ton | Duraklama (bekleme): orta ton | Dönüş: farklı ton | Dinlenme: 4 sn tam sessizlik")

    doc.save(PROTOCOL_DOCX)


def build_control_protocol():
    doc = Document()
    add_heading(doc, "Kontrol Grubu Protokolü", 0)
    add_heading(doc, "İmgeleme ve Zihinsel Arınma (Motor Olmayan Kontrol)", level=1)
    add_para(
        doc,
        "Amaç: Deneysel grupla aynı süre, ortam ve dikkat yükü. Görevler yalnızca "
        "zihinsel arınma (sakinleştirme, nefes, düşünce bırakma) ve motor olmayan "
        "görsel imgelemeden oluşur. Kol, gövde, ulaşma veya herhangi bir vücut hareketi "
        "hayal edilmez; eylem gözlemi ve frekans tonları uygulanmaz.",
    )
    add_para(doc, "Seans süresi: ~13 dk | Giriş: 60 sn | 4 blok × 3 dk (Zihinsel Arınma 45 sn + İmgeleme 90 sn + Dinlenme 45 sn)")

    add_heading(doc, "Giriş — Zihin Hazırlığı (60 sn)", level=2)
    add_para(
        doc,
        "Rahat oturun, sırtınızı sandalyeye yaslayın. İki ayağınız yere bassın. "
        "Ellerinizi uyluklarınızda dinlendirin, avuçlar yukarı veya aşağı — nasıl rahatsa. "
        "Gözlerinizi yavaşça kapatın. Derin bir nefes alın… ve yavaşça verin. "
        "Omuzlarınızı gevşetin. Zihninizdeki düşünceleri bir bulut gibi düşünün — "
        "gelir, geçer; onları itmeyin, sadece bırakın. "
        "Vücudunuzda hareket, kol veya uzanma hayal etmeyin. "
        "Sadece nefesinize ve sakinliğe odaklanın.",
    )

    blocks = [
        (
            "Blok 1 — Nefes ve Zihinsel Arınma",
            [
                (
                    "Zihinsel Arınma (45 sn)",
                    "Gözleriniz kapalı. Nefesinize dikkat edin: nefes al… nefes ver. "
                    "Her nefes verişte zihniniz biraz daha sakinleşsin. "
                    "Aklınıza gelen düşünceleri isimlendirmeden bırakın — "
                    "sanki nehirde akan yapraklar gibi geçsin gitsin. "
                    "Kol, omuz veya hareket imgeleri gelirse, yumuşakça nefese geri dönün.",
                ),
                (
                    "İmgeleme (90 sn)",
                    "Gözleriniz kapalı. Zihninizde sakin, açık bir alan hayal edin — "
                    "belki açık bir gökyüzü, belki yumuşak bir ışık. "
                    "Renklere ve ışığa odaklanın; sahne tamamen hareketsiz. "
                    "Vücudunuz sandalyede; sadece zihninizdeki görüntü değişiyor. "
                    "Ulaşma, yürüme veya kol hareketi DÜŞÜNMEYİN.",
                ),
                (
                    "Dinlenme (45 sn)",
                    "Gözlerinizi açın. Birkaç normal nefes alın. "
                    "Şimdi bir şey hayal etmeyin — sadece dinlenin.",
                ),
            ],
        ),
        (
            "Blok 2 — Renk ve Işık İmgeleme",
            [
                (
                    "Zihinsel Arınma (45 sn)",
                    "Gözlerinizi kapatın. Omuzlarınızı bir kez daha bırakın. "
                    "Nefes al… ver… Zihninizi sadeleştirin: "
                    "bugünkü endişeleri, planları, görevleri bir kenara bırakın. "
                    "Bu birkaç dakika yalnızca sizin sakin alanınız.",
                ),
                (
                    "İmgeleme (90 sn)",
                    "Zihninizde yavaşça bir renk belirsin — mavi, yeşil veya altın sarısı, "
                    "hangisi size sakin geliyorsa. Rengin tonlarını, derinliğini hayal edin. "
                    "Renk yavaşça yumuşak bir ışığa dönüşsün. "
                    "Her şey sabit ve hareketsiz; sadece renk ve ışık var. "
                    "Vücut imgeleri kullanmayın.",
                ),
                (
                    "Dinlenme (45 sn)",
                    "Gözlerinizi açın ve dinlenin.",
                ),
            ],
        ),
        (
            "Blok 3 — Sakin Doğa Sahnesi",
            [
                (
                    "Zihinsel Arınma (45 sn)",
                    "Gözleriniz kapalı. Nefesiniz doğal ritminde aksın. "
                    "Zihninizde biriken gerginliği fark edin — yargılamadan, "
                    "sadece fark edin — ve nefes verirken bırakın. "
                    "Zihinsel arınma: düşünceler geçer, siz kalırsınız.",
                ),
                (
                    "İmgeleme (90 sn)",
                    "Zihninizde sakin bir doğa sahnesi canlandırın: "
                    "durgun bir göl, yeşil ağaçlar, hafif bulutlar. "
                    "Suyun yüzeyi, ağaçların rengi, gökyüzünün tonu — "
                    "ayrıntılara odaklanın. Sahne donmuş bir fotoğraf gibi; "
                    "rüzgâr, yürüme veya el hareketi yok.",
                ),
                (
                    "Dinlenme (45 sn)",
                    "Gözlerinizi açın. Normal nefes alın. "
                    "Şimdi bir şey hayal etmeyin — sadece dinlenin.",
                ),
            ],
        ),
        (
            "Blok 4 — Tam Zihinsel Arınma ve Kapanış",
            [
                (
                    "Zihinsel Arınma (45 sn)",
                    "Gözleriniz kapalı. Derin nefes al… yavaş ver… "
                    "Zihninizi temizleyin: motor düşünceler, görevler, "
                    "hareket planları — hepsini şimdilik bırakın. "
                    "Sadece sakin nefes ve boş, huzurlu bir zihin.",
                ),
                (
                    "İmgeleme (90 sn)",
                    "Blok 1'deki açık alanı veya Blok 3'teki sakin sahneyi "
                    "zihninizde tekrar canlandırın. "
                    "Renk, ışık ve sessizliğe odaklanın. "
                    "Kol, omuz, uzanma, yürüme — hiçbiri yok. "
                    "Zihniniz sakin, sahne hareketsiz.",
                ),
                (
                    "Dinlenme (45 sn)",
                    "Gözlerinizi açın. Normal nefes alın. "
                    "Hiçbir şey hayal etmeyin — sadece dinlenin. Teşekkürler. Seans bitti.",
                ),
            ],
        ),
    ]

    for title, steps in blocks:
        add_heading(doc, title, level=2)
        for step_title, step_text in steps:
            add_para(doc, step_title, bold=True)
            add_para(doc, step_text)

    add_heading(doc, "Kontrol Grubu İlkeleri", level=2)
    add_para(doc, "• Görev türü: zihinsel arınma + motor olmayan görsel imgeleme")
    add_para(doc, "• Aynı sandalye, oda, kulaklık ve toplam süre (~13 dk)")
    add_para(doc, "• Video izleme YOK | Frekans tonları YOK | Motor imgeler YASAK")
    add_para(doc, "• Deneysel gruptan fark: içerik bilişsel-sakinleştirici, motor planlama yok")
    add_para(doc, "• Literatür: Schuster et al. (2011); Di Rienzo et al. (2016)")

    doc.save(CONTROL_DOCX)


def main():
    build_experimental_protocol()
    build_control_protocol()
    ETHICS_TXT.write_text(ETHICS_SECTION, encoding="utf-8")
    print("WROTE", PROTOCOL_DOCX)
    print("WROTE", CONTROL_DOCX)
    print("WROTE", ETHICS_TXT)


if __name__ == "__main__":
    main()
