# -*- coding: utf-8 -*-
"""Generate English-only PETTLEP / AOMI session protocol docx (NeuroLab RCT)."""
from pathlib import Path

from docx import Document
from docx.shared import Pt

OUT = Path(r"D:\Thesis app\manuscript f\forms_completed\AOMI_Session_Protocol_NeuroLab.docx")
OUT_COPY = Path(r"D:\Thesis app\NeuroLab\forms_completed\AOMI_Session_Protocol_NeuroLab.docx")

SECTIONS = [
    (
        "PETTLEP Motor Imagery Protocol in Stroke Rehabilitation",
        "The PETTLEP model (Holmes & Collins, 2001) is grounded in functional equivalence: "
        "imagined movement activates overlapping neural pathways with executed movement. "
        "PETTLEP denotes seven elements required for effective, ecologically valid motor imagery.",
    ),
    (
        "PETTLEP seven elements",
        "P – Physical: Same seated posture, clothing context, and towel/table setup as performance.\n"
        "E – Environment: Same room, chair, lighting, and tablet placement as assessment.\n"
        "T – Task: Imagery matched to current ability (e.g., limiting shoulder elevation during reach).\n"
        "T – Timing: Real-time imagery duration matching actual movement (~3 s reach + ~3 s return).\n"
        "L – Learning: Scripts updated as performance improves (e.g., finger opening, smoothness).\n"
        "E – Emotion: Comfort, confidence, satisfaction; kinesthetic muscle sensation without frustration.\n"
        "P – Perspective: External (3rd person) for observation/video; internal (1st person) for kinesthetic MI.",
    ),
    (
        "Therapist scripts — Reach & Wipe (step-by-step)",
        "Step 1 — Setup (Physical & Environment): "
        "'Sit upright on the chair. Feet flat on the floor. Feel your back supported. "
        "Place the towel under your affected hand on the table in front of you. "
        "Close your eyes and take a slow, deep breath.'\n\n"
        "Step 2 — External observation (3rd person; block minute 1): "
        "'Imagine standing beside yourself, watching your body. Notice stable sitting. "
        "Look at your affected shoulder — relaxed and low, not rising toward your ear. "
        "Your trunk is upright and provides a solid base. You are ready to move.'\n\n"
        "Step 3 — Internal imagery (1st person; block minutes 2–3): "
        "'Return inside your body. Look through your own eyes at the table and towel under your hand. "
        "Focus on the muscles of your upper arm and behind your shoulder.'\n\n"
        "Step 4 — Timing & emotion: "
        "'On my signal, imagine sending a clear message from your brain to your arm to reach forward. "
        "The movement takes only 3 seconds. Feel the towel slide on the table. "
        "Feel comfort and confidence as your arm moves smoothly without resistance. "
        "Ready? Begin… (count silently 1…2…3)… stop. "
        "Now imagine returning slowly to the start (1…2…3)… relax.'\n\n"
        "Step 5 — Learning: Repeat imagery within the block; if finger opening is difficult, add: "
        "'Feel your fingers relax and open gently over the towel before you begin to push.'\n\n"
        "Vividness check (after each block): "
        "'On a scale of 1–10, how vivid was the movement? How much did you feel your muscles working?'",
    ),
    (
        "22-minute RCT intervention session (Experimental — PETTLEP-AOMI)",
        "Physical Reach & Wipe execution occurs ONLY during pre/post assessment (3 trials, mean analysed). "
        "The 22-minute intervention contains observation and imagery only.\n\n"
        "Calibration (2 min): notice–name–transfer on the unaffected limb.\n"
        "5 blocks × 4 min:\n"
        "  Minute 1 — Action observation (mirror-reversed personal video)\n"
        "  Minutes 2–3 — Corrective motor imagery (standardized headphone audio)\n"
        "  Minute 4 — Rest (eyes open; no movement or imagery)\n"
        "Block targets (proximal → distal): "
        "(1) movement initiation (2) trunk stability (3) shoulder girdle control "
        "(4) elbow opening / muscle quieting (5) integrated smooth Reach & Wipe.\n"
        "Dose: 2 min calibration + 5 min observation + 10 min imagery + 5 min rest = 22 min.\n"
        "After each block: mental chronometry, vividness rating, physiotherapist fidelity checklist.",
    ),
    (
        "Control group (matched 22 min)",
        "2-minute introductory relaxation followed by five 4-minute blocks alternating "
        "Body Scanning (sequential somatic attention, no movement imagery) and "
        "Spatial Navigation (familiar route/layout, no human limb movement). "
        "Same room, tablet, headphones, and audio delivery as the experimental group.",
    ),
    (
        "Clinic-to-RCT session mapping",
        "Video observation (1–2 min in clinic guides) → 5 min total AO (1 min × 5 blocks).\n"
        "Mental practice (10–15 min) → 10 min MI (2 min × 5 blocks).\n"
        "Rest (1–2 min) → 5 min rest (1 min × 5 blocks).\n"
        "Physical execution (10–15 min in clinic guides) → pre/post assessment only in this RCT.\n"
        "Feedback (2 min) → chronometry + vividness + fidelity checklist per block.",
    ),
]


def _style(doc: Document) -> None:
    s = doc.styles["Normal"]
    s.font.name = "Times New Roman"
    s.font.size = Pt(12)


def build() -> None:
    doc = Document()
    _style(doc)
    doc.add_heading("AOMI Session Protocol — NeuroLab RCT", level=1)
    for heading, body in SECTIONS:
        doc.add_heading(heading, level=2)
        doc.add_paragraph(body)
    for path in (OUT, OUT_COPY):
        path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(path))
    print(f"Written: {OUT}")


if __name__ == "__main__":
    build()
