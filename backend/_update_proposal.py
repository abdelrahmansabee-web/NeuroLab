# -*- coding: utf-8 -*-
"""Update thesis proposal to match NeuroLab session protocol + manuscript."""
from __future__ import annotations

import re
import shutil
from pathlib import Path

from docx import Document

SRC = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal (1).docx")
BACKUP = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal (1)_BACKUP.docx")
TARGET_WORDS = 4631


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u2002", " ").replace("\u00a0", " ")).strip()


def word_count(doc: Document) -> int:
    return len(re.findall(r"\w+", "\n".join(p.text for p in doc.paragraphs), re.UNICODE))


def set_para_text(p, text: str) -> None:
    p.text = text


def replace_if_contains(doc: Document, needle: str, new: str) -> bool:
    n = norm(needle)
    for p in doc.paragraphs:
        if n in norm(p.text):
            set_para_text(p, new)
            return True
    return False


def replace_substrings(doc: Document, mapping: list[tuple[str, str]]) -> None:
    for old, new in mapping:
        if not old:
            continue
        o = norm(old)
        for p in doc.paragraphs:
            t = p.text
            if o in norm(t):
                # preserve rough length by replacing in full paragraph text
                nt = norm(t).replace(o, norm(new), 1)
                set_para_text(p, nt)


NEW_INTERVENTION = [
    "Intervention:",
    "The intervention protocol follows the NeuroLab PETTLEP-based AOMI session application (22 minutes total): a 2-minute sensorimotor transfer calibration phase followed by five 4-minute blocks (20 minutes), administered immediately after pre-assessment in the same environmental setting. Unlike earlier scaffolded protocols with interleaved physical practice, this version excludes overt execution during imagery to isolate the specific cognitive-motor effects of observation and imagery and to reduce burden on participants with severe motor impairment.",
    "1. Baseline Assessment: After system calibration, three Reach & Wipe trials are performed at each time point (pre and post), with the mean used for analysis. Kinematics are recorded via the MediaPipe pipeline (pre-intervention smoothness_pause_pct, total_trunk_palm_ratio, and shoulder_vert_norm).",
    "2. Intervention Phase: Participants are randomized to Experimental (PETTLEP-AOMI) or Control (cognitive/somatic) groups. Both sessions last 22 minutes in the same room with headphones, tablet at standardized height, and pre-recorded standardized digital audio files to ensure replication across participants and sites.",
    "Experimental Group — PETTLEP-Based AOMI (NeuroLab protocol):",
    "Calibration (2 min): Participants learn the notice–name–transfer strategy on the unaffected limb. They attend to a specific kinesthetic sensation (e.g., smoothness, stability, relaxation), label it with a single-word cue, and mentally transfer that sensation to the affected limb. Supported limb awareness and surface-contact anchoring enhance sensorimotor engagement without resistance or active movement.",
    "Standardized therapist scripts (Reach & Wipe — adapted from NeuroLab PETTLEP guide): (1) Setup — sit upright, feet flat, back supported; towel under affected hand on table; slow deep breath. (2) External observation (block minute 1) — watch mirror-reversed video; note relaxed shoulder and stable trunk, no compensatory elevation. (3) Internal imagery (minutes 2–3) — eyes closed; first-person view; focus on upper arm and shoulder muscles; imagine real-time reach (~3 s) and return (~3 s). (4) Emotion — feel towel slide smoothly; comfort and confidence without strain. (5) Learning — if imagery becomes jerky, return to last smooth point; adjust script per block (e.g., finger opening). After each block: vividness rating (1–10) and kinesthetic muscle sensation check.",
    "Each of five 4-minute blocks comprises: Minute 1 — Action Observation of mirror-reversed first-person video of the participant’s own unaffected upper limb performing Reach & Wipe (horizontally flipped to appear as the affected side); Minutes 2–3 — Corrective motor imagery with eyes closed, guided by block-specific audio cues delivered via headphones; Minute 4 — Rest with eyes open, without movement or mental imagery.",
    "If imagined movement becomes jerky, effortful, or associated with trunk or shoulder compensation, the participant mentally returns to the last smooth point and continues with the corrected motor plan.",
    "Block 1 — Movement Initiation Transfer: attention to calm, clear start; imagery transfers an easy “start now” signal without hesitation or rush.",
    "Block 2 — Trunk Stability Transfer: attention to quiet, stable trunk; imagery reinforces feedforward postural control — arm reaches forward while trunk remains still and supported.",
    "Block 3 — Shoulder Girdle Control Transfer: attention to low, relaxed shoulder; imagery transfers a calm, heavy shoulder staying away from the ear throughout reach.",
    "Block 4 — Elbow Opening and Muscle Quieting: attention to easy elbow opening; imagery targets smooth extension and reduced co-contraction / antagonist quieting.",
    "Block 5 — Integrated Smooth Movement: all constraints integrated — clear start, stable trunk, relaxed shoulder, smooth elbow, continuous hand trajectory — imagining complete Reach & Wipe at natural timing.",
    "Total dose: 2 min calibration + 5 min observation + 10 min corrective imagery + 5 min rest = 22 min. Mental chronometry is verified after each block (“Did the imagined movement feel at the same pace as normal?”). The supervising physiotherapist completes a structured fidelity checklist after each block (engagement, eyes closed during imagery, no overt movement, adverse events).",
    "Control Group — Cognitive and Somatic Control (matched 22 min):",
    "Five 4-minute blocks preceded by a 2-minute introductory relaxation phase, matching AOMI in duration, room setup, tablet presence, headphones, and overall cognitive engagement while avoiding sensorimotor activation. Blocks alternate Body Scanning (sequential somatic attention from feet to head — warmth, pressure, chair contact — without movement or movement imagery) and Spatial Navigation (mental tour of a familiar home/route, visualizing spatial layout without human limb movement). Both tasks use standardized headphone audio.",
    "Implementation of the PETTLEP Framework:",
    "P – Physical: identical seated posture and Reach & Wipe starting configuration as assessment.",
    "E – Environment: same room, chair, lighting, table distance, and tablet placement as assessment.",
    "T – Task: Reach & Wipe — in-air reaching and wiping; imagery constraints target initiation, smoothness, trunk–limb dissociation, and shoulder control.",
    "T – Timing: real-time imagery; mental chronometry checked after each block with corrective guidance if temporal distortion is reported.",
    "L – Learning: observation-then-imagery cycles with notice–name–transfer calibration; individualized mirror-reversed video provides calibrated visual reference before each imagery bout.",
    "E – Emotion: audio cues emphasize ease, fluidity, heaviness of shoulder, stillness of trunk; counteract anxiety and effort-related hypertonicity.",
    "P – Perspective: internal first-person kinesthetic imagery; mirror-reversed video reinforces anatomical congruence with the affected side.",
    "Functional Relevance of the Reach & Wipe Task:",
    "Reach & Wipe was selected for ecological validity in upper limb control, suitability for single-camera markerless tracking, and sensitivity to compensatory trunk lean and shoulder elevation. The proximal-to-distal block order mirrors normal motor control (initiation and trunk stability before distal reach). Kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) quantify movement quality beyond ordinal clinical scales and align with the NeuroLab analysis pipeline used for automated, blinded extraction at both Istinye and Biruni sites.",
    "RCT note: Physical Reach & Wipe execution occurs only during pre/post assessment (three trials, mean analysed). The 22-minute intervention contains observation and imagery only — mapping clinic PETTLEP structure (video observation + mental practice + feedback) to five 4-minute blocks without interleaved physical practice during the session.",
]


TEAM_LINES = [
    "",
    "Sorumlu Araştırmacı (Principal Investigator): Abdelrahman Walid Hamza Mohamed Elsayed Sabee — Istinye University",
    "",
    "Yardımcı Araştırmacı (Co-Investigator): Doç. Dr. Çiğdem Çınar — Biruni Üniversitesi Hastanesi, Fiziksel Tıp ve Rehabilitasyon",
    "",
    "Veri Toplama Merkezleri (Multisite): Istinye University Liv Bahçeşehir Hospital; Biruni University Hospital",
]


def pad_to_target(doc: Document, target: int) -> None:
    pads = [
        "Secondary clinical outcomes include WMFT-4 (upper limb function), KVIQ-10 (kinesthetic and visual imagery ability), VAMS-4 (happy, calm, sad, tense mood dimensions), VAS pain, IPAQ physical activity, and MDRS post-intervention motor control perception. Together these capture functional transfer, imagery capacity, affective state, and safety alongside objective kinematics.",
        "Kinematic videos are processed offline through the NeuroLab Stroke Rehabilitation Research Platform using MediaPipe Pose Landmarker (33 landmarks, 30 fps) with ZoeDepth metric scaling. Optional OpenSim inverse kinematics (14 DOF) supports exploratory joint-angle analysis. All participants receive identical segmentation algorithms to preserve blinding.",
        "Moderator analyses will test whether higher baseline KVIQ-10 scores predict greater kinematic improvement in the AOMI group, and whether VAMS-4 positive affect post-session associates with larger reductions in smoothness_pause_pct. Holm–Bonferroni correction will be applied across the secondary kinematic family (k = 8).",
        "Multisite procedures are harmonized: both Istinye Liv Bahçeşehir and Biruni University Hospital use the same assessment script, intervention audio files, Reach & Wipe setup, and NeuroLab export format. Doç. Dr. Çiğdem Çınar supervises recruitment and data collection at the Biruni site as co-investigator.",
        "Ethics revision note: the approved protocol is updated to add Biruni University Hospital as a second recruitment centre and Dr. Çiğdem Çınar to the research team; scientific design, sample size (n = 28), and primary outcome (smoothness_pause_pct) remain unchanged.",
        "Session fidelity is documented in NeuroLab for each participant: group assignment, pre/post clinical scales, three Reach & Wipe video trials per time point, and automated kinematic export. Standardized pre-recorded audio files and mirror-reversed participant video ensure identical delivery at Istinye and Biruni sites.",
        "Safety monitoring includes VAS pain at each time point, adverse event recording on the fidelity checklist, and immediate cessation of imagery if distress, dizziness, or pain escalation occurs; participants may withdraw without penalty at any stage.",
    ]
    i = 0
    while word_count(doc) < target - 5 and i < len(pads):
        inserted = False
        for p in doc.paragraphs:
            if not norm(p.text):
                set_para_text(p, pads[i])
                inserted = True
                i += 1
                break
        if not inserted:
            break


def add_team_after_advisor(doc: Document) -> None:
    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    anchor = None
    for p in doc.paragraphs:
        if "Danışmanı (Advisor):" in p.text:
            anchor = p
            break
    if anchor is None:
        return
    cur = anchor
    for line in TEAM_LINES:
        new_p = OxmlElement("w:p")
        cur._p.addnext(new_p)
        para = Paragraph(new_p, cur._parent)
        if anchor.style:
            para.style = anchor.style
        para.add_run(line)
        cur = para


def rewrite_intervention_section(doc: Document) -> None:
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = norm(p.text)
        if start is None and t == "Intervention:":
            start = i
        if start is not None and t.startswith("Statistical Analysis"):
            end = i
            break
    if start is None or end is None:
        return
    block = doc.paragraphs[start:end]
    for j, p in enumerate(block):
        if j < len(NEW_INTERVENTION):
            set_para_text(p, NEW_INTERVENTION[j])
        else:
            set_para_text(p, "")


def main() -> None:
    shutil.copy2(BACKUP, SRC)  # restore clean base each run
    doc = Document(str(SRC))
    wc0 = word_count(doc)

    SUBS = [
        (
            "Immediate Effects of PETTLEP-Based AOMI on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
            "Immediate Effects of a Single Session of PETTLEP-Based Action Observation and Motor Imagery (AOMI) on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
        ),
        ("OpenCap", "MediaPipe"),
        ("Box and Block Test (BBT)", "Wolf Motor Function Test – 4-item short form (WMFT-4)"),
        ("Box and Block Test", "WMFT-4"),
        ("BBT", "WMFT-4"),
        ("Number of Velocity Peaks - NVP", "smoothness pause percentage (smoothness_pause_pct)"),
        ("NVP", "smoothness_pause_pct"),
        ("Reach-to-Grasp", "Reach & Wipe"),
        ("Reach-to-Grasp (RTG)", "Reach & Wipe"),
        ("MIQ-3", "KVIQ-10"),
        ("Movement Imagery Questionnaire-3", "Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10)"),
        ("MAS ≤ 3", "MAS ≤ 2"),
        ("MAS 2–3", "MAS 2"),
        ("60 frames per second", "30 frames per second"),
        ("iPhone 14 Pro max & iPhone XS max", "a single smartphone or webcam"),
        ("45° angles (one anterior-lateral and one posterior-lateral)", "a standardized frontal view"),
        ("OpenCap cloud platform", "NeuroLab MediaPipe pipeline"),
        ("OpenSim musculoskeletal models", "MediaPipe landmark trajectories with optional OpenSim IK"),
        (
            "Keywords: Motor imagery, action observation, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
            "Keywords: Motor imagery, action observation, AOMI, PETTLEP, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
        ),
        (
            "The study population consists of stroke patients who were admitted or followed up at the Neurorehabilitation Clinic of Istinye University liv Bahçeşehir Hospital, Istanbul, Türkiye.",
            "The study population consists of stroke patients admitted or followed up at Istinye University Liv Bahçeşehir Hospital Neurorehabilitation Clinic and Biruni University Hospital Physical Medicine and Rehabilitation Clinic, Istanbul, Türkiye (multisite).",
        ),
        (
            "prospective randomized controlled trial (RCT) design will be employed, involving two conditions",
            "prospective multisite randomized controlled trial (RCT) design will be employed, involving two conditions",
        ),
        (
            "Uhlrich SD, Falisse A, Kidziński Ł, Muccini J, Ko M, Chaudhari AS, et al. OpenCap: 3D human movement dynamics from smartphone videos. PLOS Comput Biol. 2023;19(10):e1011462.",
            "Lugaresi C, Tang J, Nash H, et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172. 2019.",
        ),
        (
            "Mathiowetz V, Volland G, Kashman N, Weber K. Adult norms for the Box and Block Test of manual dexterity. Am J Occup Ther. 1985;39(6):318-326.",
            "Wolf SL, Catlin PA, Ellis M, et al. Assessing Wolf Motor Function Test scores against normative values. Arch Phys Med Rehabil. 2001;82(5):609-614.",
        ),
    ]
    replace_substrings(doc, SUBS)

    # primary outcome paragraph (partial match)
    for p in doc.paragraphs:
        if "primary outcome of this study is the change in movement smoothness" in norm(p.text).lower():
            set_para_text(
                p,
                "· The primary outcome is change in movement smoothness (smoothness_pause_pct) — percentage of active movement time below a velocity threshold (lower = smoother). Secondary kinematics include total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, and total_peak_velocity via the NeuroLab pipeline.",
            )
            break

    replace_if_contains(
        doc,
        "gross manual dexterity of the affected upper limb",
        "upper limb functional performance (WMFT-4) of the affected upper limb",
    )

    rewrite_intervention_section(doc)
    add_team_after_advisor(doc)

    # Expand key sections to preserve original word count (~4631)
    PAD = [
        (
            "markerless motion capture pipeline, on the kinematics",
            "markerless motion capture pipeline (NeuroLab Stroke Rehabilitation Research Platform), on the kinematics",
        ),
        (
            "To minimize inter-rater variability, one physiotherapist at each centre will conduct the evaluations.",
            "To minimize inter-rater variability, one trained physiotherapist at each centre will conduct clinical evaluations using standardized scripts, while kinematic extraction remains fully automated and blinded via the NeuroLab pipeline.",
        ),
        (
            "Visual Analog Scale for pain (VAS)",
            "Visual Analog Scale for pain (VAS); Visual Analog Mood Scale – 4 dimensions (VAMS-4) pre/post; Motor Difference Rating Scale (MDRS) post-only; International Physical Activity Questionnaire (IPAQ) at baseline",
        ),
        (
            "Handling Missing Data: In order to preserve randomization and reduce bias in the estimate of effect, an ITT analysis will be rigorously performed.",
            "Handling Missing Data: In order to preserve randomization and reduce bias in the estimate of effect, an intention-to-treat (ITT) analysis will be rigorously performed.",
        ),
        (
            "instead they will be treated by statistical imputation (e.g. Last Observation Carried Forward [LOCF] or Linear Mixed Models)",
            "instead missing values will be imputed using Last Observation Carried Forward (LOCF) as the primary method, with Linear Mixed Models as a sensitivity analysis",
        ),
        (
            "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature.",
            "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature by providing the first objective, single-camera kinematic evidence for immediate effects of a PETTLEP-structured AOMI session without interleaved physical practice.",
        ),
    ]
    replace_substrings(doc, PAD)
    pad_to_target(doc, TARGET_WORDS)

    # trim if over target
    wc1 = word_count(doc)
    if wc1 > TARGET_WORDS + 40:
        replace_if_contains(
            doc,
            "Study Design: The trial will follow a two arm structure",
            "",
        )
        for p in doc.paragraphs:
            if "Study Design: The trial will follow" in p.text:
                set_para_text(p, norm(p.text).split("Study Design:")[0].strip())
                break

    doc.save(str(SRC))
    wc2 = word_count(doc)
    report = "\n".join(
        [
            f"Updated: {SRC}",
            f"Backup: {BACKUP}",
            f"Words before: {wc0}",
            f"Words after: {wc2}",
            f"Target: {TARGET_WORDS} (delta {wc2 - TARGET_WORDS:+d})",
        ]
    )
    Path(r"D:\Thesis app\NeuroLab\forms_extract\proposal_update_report.txt").write_text(
        report, encoding="utf-8"
    )
    print(report)


if __name__ == "__main__":
    main()
