# -*- coding: utf-8 -*-
"""Extract proposal structure and sample formatting."""
import sys
from pathlib import Path
from docx import Document
from docx.enum.text import WD_COLOR_INDEX

sys.stdout.reconfigure(encoding="utf-8")
p = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal .docx")
d = Document(str(p))
out = Path(r"D:\Thesis app\NeuroLab\forms_extract\proposal_original_indexed.txt")
lines = []
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if not t:
        continue
    bold = any(r.bold for r in para.runs if r.text.strip())
    lines.append(f"{i:3d} | bold={bold} | {t[:150]}")
out.write_text("\n".join(lines), encoding="utf-8")
print(f"Wrote {len(lines)} lines to {out}")
# section markers
for i, para in enumerate(d.paragraphs):
    t = para.text.strip()
    if t in ("Research Question", "Hypotheses", "Intervention:", "Statistical Analysis", "References", "Study Design") or t.startswith("2.") or t.startswith("1."):
        print(i, t[:80])
