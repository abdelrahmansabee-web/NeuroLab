# -*- coding: utf-8 -*-
import re
import sys
from pathlib import Path
from docx import Document

sys.stdout.reconfigure(encoding="utf-8")
p = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal (1).docx")
d = Document(str(p))
text = "\n".join(x.text for x in d.paragraphs)
print("Words:", len(re.findall(r"\w+", text, re.UNICODE)))
print("Has 25 min:", "25 min" in text)
print("Has 1:3:", "1:3" in text)
print("Has NeuroLab protocol:", "NeuroLab protocol" in text)
print("Has notice-name-transfer:", "notice" in text.lower() and "transfer" in text.lower())
start = None
for i, x in enumerate(d.paragraphs):
    if x.text.strip() == "Intervention:":
        start = i
        break
if start is not None:
    print(f"\nIntervention section from para {start}:")
    for j in range(start, min(start + 30, len(d.paragraphs))):
        t = d.paragraphs[j].text.strip()
        print(f"  {j}: {t[:160] if t else '(empty)'}")
else:
    print("Intervention: heading NOT FOUND")
