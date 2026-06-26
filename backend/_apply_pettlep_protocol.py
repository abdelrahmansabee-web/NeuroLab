# -*- coding: utf-8 -*-
"""Apply user PETTLEP session protocol to proposal and manuscript."""
from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

PROP = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal .docx")
MAN = Path(
    r"D:\Thesis app\manuscript f\Immediate Effects of a Single Session of PETTLEP_UPDATED_CORRECTED.docx"
)
MAN_ALT = MAN.with_name(MAN.stem + "_PROTOCOL_UPDATED.docx")
MAN_OUT = MAN
RED = RGBColor(0xFF, 0x00, 0x00)

PROPOSAL_INTERVENTION = [
    "Intervention:",
    "The intervention follows a standardized PETTLEP-based AOMI session for the Reach & Wipe task. Total session duration is approximately 17 minutes (within the recommended 15–20 minute range): a 2-minute calibration phase followed by five 3-minute training blocks. Overt physical practice is excluded during the session to isolate cognitive-motor effects of observation and imagery.",
    "1. Baseline Assessment: After system calibration, three Reach & Wipe trials are performed at each time point (pre and post), with the mean used for analysis. Kinematics are recorded via the MediaPipe pipeline (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm).",
    "2. Intervention Phase: Participants are randomized to Experimental (PETTLEP-AOMI) or Control groups. Both sessions are matched in duration (~17 min), room setup, headphones, tablet height, and pre-recorded standardized audio delivery at Istinye and Biruni sites.",
    "Experimental Group — PETTLEP-Based AOMI:",
    "Session timing structure: Calibration (120 s) + five blocks × 180 s = 17 minutes. Each 3-minute block comprises Watch 45 s, Imagine 75 s, and Rest 60 s. Clinical practice allows 4–6 blocks; this study uses five blocks for dose standardization.",
    "Phase 1 — Calibration and Priming (120 seconds): Goal — deep mental focus and sensorimotor transfer from the unaffected to the affected limb before imagery.",
    "Calibration script (delivered in a calm voice via headphones): “Close your eyes. Feel the chair supporting your body weight. Take a deep breath in… and release it slowly.” (15 s). “Now focus on your unaffected hand. Feel the towel beneath it. Imagine wiping the table smoothly with that hand — see through your own eyes (internal perspective). Feel natural shoulder and arm contraction and the comfort of smooth movement. We will call this sensation ‘comfort’.” (45 s). “Now clearly imagine this strong ‘comfort’ signal travelling from your brain to your affected arm, awakening muscle fibres there. You are ready.” (60 s).",
    "Phase 2 — Training Blocks (repeated five times): Each block uses the same Watch → Imagine → Rest sequence applied to Reach & Wipe on the affected side.",
    "Step 1 — Watch / Action Observation (45 s; Perspective & Task): Participant views mirror-reversed first-person video of their own unaffected limb performing Reach & Wipe. External (third-person) monitoring is encouraged: “Open your eyes and watch the video. Observe as an external monitor — smooth arm extension, relaxed shoulders not rising toward the neck, and natural reach–wipe–return timing.”",
    "Step 2 — Imagine / Motor Imagery (75 s; Physical, Timing & Emotion): Eyes closed; internal (first-person) perspective; real-time imagery is critical. Script: “Return inside your body. Look through your own eyes at the table and towel under your affected hand. We will imagine the movement in real time.” The therapist paces with counting: Reach forward 1–2–3 (feel the towel slide); Wipe sideways 1–2–3 (feel arm weight moving smoothly with confidence); Return 1–2–3 (feel muscle relaxation). Three to five complete cycles are performed within 75 seconds according to each participant’s target pace.",
    "Step 3 — Neural Rest (60 s): “Stop imagining. Allow mind and muscles to relax completely. Do not think about movement for one full minute. Breathe calmly.” This prevents mental fatigue and preserves imagery quality in subsequent blocks.",
    "Implementation of the PETTLEP Framework:",
    "P – Physical: Participants adopt the actual task posture — seated on a real chair, normal clothing, real towel under the hand — matching assessment configuration.",
    "E – Environment: Imagery occurs in the same functional environment as assessment (real table setup with everyday objects), not an abstract empty clinical space.",
    "T – Task: Reach & Wipe is tailored to current upper-limb capacity and daily-life relevance.",
    "T – Timing (critical): Imagery duration matches real movement time; paced verbal counting during Imagine enforces 3-second reach, wipe, and return phases.",
    "L – Learning: Imagery content evolves across blocks as vividness improves; corrective guidance is given when imagery becomes jerky or compensatory.",
    "E – Emotion: Instructions integrate comfort, confidence, and awareness of natural muscular effort without strain.",
    "P – Perspective: External (third-person) during Watch for posture monitoring; internal (first-person, kinesthetic) during Imagine for muscle sensation.",
    "Total dose: 2 min calibration + 5 × 45 s observation + 5 × 75 s imagery + 5 × 60 s rest = 17 min. Mental chronometry is verified after each block. The physiotherapist completes a fidelity checklist (engagement, eyes closed during imagery, no overt movement, adverse events).",
    "Control Group — Cognitive and Somatic Control (matched ~17 min):",
    "Two-minute introductory relaxation followed by five 3-minute blocks alternating Body Scanning (sequential somatic attention from feet to head — warmth, pressure, chair contact — without movement or movement imagery) and Spatial Navigation (mental tour of a familiar home/route without human limb movement). Headphones deliver standardized audio; setup matches the experimental group while avoiding sensorimotor activation.",
    "Functional Relevance of the Reach & Wipe Task:",
    "Reach & Wipe provides ecological validity for upper-limb control, suits single-camera markerless tracking, and is sensitive to compensatory trunk lean and shoulder elevation. Kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) quantify movement quality via the NeuroLab pipeline at both sites.",
    "RCT note: Physical Reach & Wipe execution occurs only during pre/post assessment (three trials, mean analysed). The 17-minute intervention contains observation and imagery only — Watch (45 s), Imagine (75 s), and Rest (60 s) per block — without interleaved physical practice.",
]

ABSTRACT_METHODS_NEW = (
    "Methods: A single-blind, pretest–posttest, prospective randomized controlled trial (RCT) was designed for conduct at "
    "Istinye University Liv Bahçeşehir Hospital and Biruni University Hospital, Istanbul, Türkiye (multisite). "
    "A total of 28 stroke survivors (14 per group) aged 40–80 years, with mild-to-moderate spasticity (Modified Ashworth Scale ≤ 2) "
    "and residual voluntary upper limb movement, were to be recruited and randomized to either the PETTLEP-based AOMI group or a "
    "cognitive and somatic control group. The AOMI protocol consisted of a 2-minute calibration phase followed by five standardized "
    "3-minute blocks (17 minutes total): each block comprised 45 seconds of action observation (Watch), 75 seconds of real-time "
    "corrective motor imagery (Imagine), and 60 seconds of rest. Participants used a notice–name–transfer strategy during calibration "
    "to map kinesthetic sensations from the unaffected to the affected limb. The control group performed time-matched body scanning "
    "and spatial navigation tasks (~17 minutes total). Two-dimensional pose-landmark kinematics of the affected upper limb were to be "
    "captured using a custom markerless motion capture pipeline based on MediaPipe Pose Landmarker (Google) for 2D landmark extraction "
    "from a single smartphone or webcam at 30 fps and 1080p resolution, combined with the ZoeDepth monocular depth estimation network "
    "for metric scaling. The primary outcome was movement smoothness (smoothness_pause_pct) — the percentage of time within the active "
    "movement window where hand velocity falls below a fixed normalized threshold (0.03), indexing movement intermittency and stop-and-go "
    "behavior. Secondary outcomes included total movement duration (total_duration_s), peak and mean hand velocity, trunk-to-palm "
    "displacement ratio, maximum elbow angle, total path length, lateral hand range, shoulder vertical displacement range "
    "(shoulder_vert_norm; cm equivalent via ZoeDepth metric scaling, exported as total_depression_cm), phase-specific metrics "
    "(forward, wipe_right, wipe_left, return), Wolf Motor Function Test – 4-item short form (WMFT-4), Visual Analog Mood Scale – "
    "4 dimensions (VAMS-4), Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10), International Physical Activity "
    "Questionnaire (IPAQ), Motor Difference Rating Scale (MDRS, post-intervention only), and Visual Analog Scale for pain (VAS). "
    "All kinematic variables were collected and managed using the NeuroLab Stroke Rehabilitation Research Platform "
    "(custom web application; kinematics_analyzer.py v6). Data were to be analyzed using a 2 × 2 Mixed-Model ANOVA, with non-parametric "
    "alternatives applied where indicated."
)

MANUSCRIPT_REPLACEMENTS: list[tuple[str, str]] = [
    (
        "Understanding this proximal-to-distal principle further informed the sequential structure of the PETTLEP-based AOMI protocol employed in the present study. The five progressive imagery blocks were ordered to mirror the physiological hierarchy of normal reaching: movement initiation first, followed by trunk stability, then shoulder girdle control, then elbow opening, and finally integrated smooth movement. This ordering reflects the neurophysiological logic of APA-first, limb-action-second, and was chosen to maximize the ecological validity and motor specificity of the imagery protocol.",
        "Understanding this proximal-to-distal principle informed the Reach & Wipe task selection and imagery content within the PETTLEP-based AOMI protocol. Rather than separate block-specific motor targets, each standardized 3-minute cycle (Watch 45 s → Imagine 75 s → Rest 60 s) rehearsed the complete reach–wipe–return sequence with explicit emphasis on trunk stability, relaxed shoulder posture, and real-time timing — reflecting the APA-first, limb-action-second hierarchy embedded in the task instructions and therapist scripts.",
    ),
    (
        "The present study addressed these gaps by delivering a single session of PETTLEP-based AOMI — incorporating a sensorimotor transfer calibration phase, individualized mirror-reversed first-person observation, and progressive corrective motor imagery targeting initiation, postural control, and smoothness — and measuring the immediate kinematic response using the custom MediaPipe-based markerless motion capture pipeline with monocular depth estimation. This trial was designed to provide the first direct, objective evidence for the acute neuromotor effects of this approach in persons post-stroke.",
        "The present study addressed these gaps by delivering a single session of PETTLEP-based AOMI — incorporating a 2-minute sensorimotor transfer calibration phase, individualized mirror-reversed action observation (45 s per block), and real-time corrective motor imagery (75 s per block) structured in five standardized 3-minute cycles — and measuring the immediate kinematic response using the custom MediaPipe-based markerless motion capture pipeline with monocular depth estimation. This trial was designed to provide the first direct, objective evidence for the acute neuromotor effects of this approach in persons post-stroke.",
    ),
    (
        "Imagery of trunk stability reinstates feedforward postural control: The trunk stability imagery in Block 2 was specifically designed to reinstate, at the level of cortical motor representation, the feedforward postural organization that normally precedes arm movement. By imagining the trunk as stable and grounded before the arm reaches forward, the participant rehearsed the APA-first, arm-action-second sequence that characterizes normal reaching (20). This cortical rehearsal of the anticipatory postural hierarchy was hypothesized to reduce compensatory trunk lean in the subsequent physical assessment.",
        "Imagery of trunk stability reinstates feedforward postural control: During each Imagine phase, participants were instructed to keep the trunk quiet and supported while reaching forward in real time — rehearsing, at the cortical level, the APA-first, arm-action-second sequence that characterizes normal reaching (20). This feedforward postural emphasis within the standardized scripts was hypothesized to reduce compensatory trunk lean in the subsequent physical assessment.",
    ),
    (
        "Imagery of reciprocal muscle quieting targets spinal reflex dysregulation: The elbow opening and muscle quieting imagery in Block 4 specifically addressed the co-contraction and impaired reciprocal inhibition that underlie movement fragmentation in spastic hemiparesis (27). By imagining smooth elbow extension with relaxed antagonist musculature, the participant rehearsed the neural pattern of agonist activation with antagonist inhibition — a pattern that MI has been shown to reinforce through modulation of spinal interneuronal circuits (31).",
        "Imagery of smooth elbow extension targets spinal reflex dysregulation: Real-time Imagine scripts emphasized smooth reach and return with relaxed muscular effort, addressing co-contraction and impaired reciprocal inhibition that underlie movement fragmentation in spastic hemiparesis (27). By imagining smooth elbow extension and wipe with natural muscle sensation, participants rehearsed agonist activation with antagonist quieting — a pattern that MI may reinforce through modulation of spinal interneuronal circuits (31).",
    ),
    (
        "Integration into Early Rehabilitation: A 22-minute, single-session PETTLEP-based AOMI intervention requiring only a tablet and headphones is highly feasible for integration into the early rehabilitation phase, even in individuals with severe motor impairment who cannot participate in task-based physical training.",
        "Integration into Early Rehabilitation: A 17-minute, single-session PETTLEP-based AOMI intervention (within the recommended 15–20 minute clinical range) requiring only a tablet and headphones is highly feasible for integration into the early rehabilitation phase, even in individuals with severe motor impairment who cannot participate in task-based physical training.",
    ),
]

MANUSCRIPT_PETTLEP_BULLETS = [
    "Physical: Adopting the actual physical posture of the task — e.g., seated on a real chair, wearing normal clothing, holding or contacting a real towel where possible.",
    "Environment: Performing imagery in an environment similar to real-world performance (standardized table with everyday objects) rather than an abstract empty hospital room.",
    "Task: Selecting an imagined movement appropriate to the patient's current abilities and daily-life needs (Reach & Wipe).",
    "Timing: Performing imagery in real time so imagined duration matches actual movement duration; verified via paced counting and mental chronometry (critical dimension).",
    "Learning: Evolving the imagery scenario as the patient improves and acquires skill across repeated blocks.",
    "Emotion: Integrating movement-linked feelings such as comfort, confidence, and natural muscular effort without excessive strain.",
    "Perspective: Using internal (first-person) perspective to feel muscles during Imagine, and external (third-person) perspective to monitor posture and timing during Watch.",
]

MANUSCRIPT_SECTION_27 = """2.7 Intervention Protocol
2.7.1 Overview
The intervention consisted of a single session lasting approximately 17 minutes (within the recommended 15–20 minute clinical range), comprising a 2-minute calibration phase followed by five standardized 3-minute training blocks (15 minutes), administered immediately following the pre-intervention assessment. Both groups received the intervention in the same environmental setting. The functional task exemplar throughout was Reach & Wipe — reaching forward and wiping the table with the affected upper limb.

2.7.2 Experimental Group: PETTLEP-Based AOMI
The experimental protocol employed an AOMI paradigm — structured observation (Watch) followed immediately by real-time motor imagery (Imagine), without interleaved physical execution. Physical practice was excluded deliberately to (1) isolate cognitive-motor effects of observation and imagery from sensory-motor effects of overt movement, and (2) reduce burden on participants with severe impairment.

Exact session timings:
• Calibration: 120 seconds (2 minutes)
• Each training block: 180 seconds (3 minutes)
  – Watch (Action Observation): 45 seconds
  – Imagine (Motor Imagery): 75 seconds
  – Rest: 60 seconds
• Number of blocks: five (clinical practice allows 4–6; five blocks were selected for dose standardization)
• Total session: 2 + (5 × 3) = 17 minutes

Phase 1 — Calibration and Priming (120 seconds)
Goal: Establish deep mental focus and link somatic sensation from the unaffected limb before transferring attention to the affected limb.
Standardized script (calm voice, headphones):
(0–15 s) "Close your eyes. Feel the chair you are sitting on and your stable body weight. Take a deep breath… and release it slowly."
(15–60 s) "Now focus attention on your unaffected hand. Feel the towel beneath it. Imagine wiping the table smoothly with that hand — see the movement through your own eyes (internal perspective). Feel natural contraction in shoulder and arm muscles and the comfort of smooth movement. We will name this sensation 'comfort'."
(60–120 s) "Now vividly imagine this strong 'comfort' signal travelling from your brain… crossing to your affected arm. Feel it awakening muscle fibres there. You are ready."

Phase 2 — Training Blocks (five repetitions)
Each 3-minute block followed the same Watch → Imagine → Rest sequence for Reach & Wipe.

Step 1 — Watch / Visual Monitoring (45 seconds; Perspective & Task)
Goal: Activate mirror-neuron networks and calibrate the correct movement path.
Procedure: Participant views mirror-reversed first-person video of their own unaffected upper limb performing Reach & Wipe (flipped to appear as the affected side).
Script: "Open your eyes and watch this video. Observe carefully as an external monitor (third-person perspective). Notice smooth arm extension, relaxed shoulders that do not rise toward the neck, and natural timing — reach… wipe… return."

Step 2 — Imagine / Real-Time Motor Imagery (75 seconds; Physical, Timing & Emotion)
Goal: Maximize motor-cortex activation using real-time temporal structure.
Script: "Close your eyes again. Return inside your body. Look through your own eyes (internal perspective) at the table and towel under your affected hand. We will imagine the movement in exactly real time."
The therapist paces with regular counting:
• Reach forward: 1… 2… 3 (feel the towel slide)
• Wipe sideways: 1… 2… 3 (feel arm weight moving smoothly with confidence)
• Return to start: 1… 2… 3 (feel muscle relaxation)
Three to five complete cycles are completed within 75 seconds according to each participant's target movement pace. If imagery becomes jerky or compensatory, the participant returns mentally to the last smooth point and continues with the corrected plan (Learning dimension).

Step 3 — Neural Rest (60 seconds)
Goal: Prevent mental fatigue that could degrade imagery quality in subsequent blocks.
Script: "Stop imagining. Allow your mind and muscles to relax completely. Do not think about any movement for one full minute. Just breathe calmly."

Total intervention dose:
• 2 minutes calibration
• 5 × 45 s = 225 s (3.75 min) action observation
• 5 × 75 s = 375 s (6.25 min) motor imagery
• 5 × 60 s = 300 s (5 min) rest
• Total: 17 minutes

Intervention Fidelity:
All scripts were delivered via pre-recorded standardized digital audio files through headphones. Mental chronometry was verified after each Imagine phase: "Did the imagined movement feel like it happened at the same pace as you would normally move?" Corrective guidance was provided if temporal distortion was reported. The supervising physiotherapist completed a structured fidelity checklist after each block (engagement, eyes closed during Imagine, no overt movement, adverse events).

2.7.3 Control Group: Cognitive and Somatic Control
The control condition matched the AOMI protocol in: (a) total duration (~17 minutes); (b) physical setup (same room, chair, tablet); (c) headphones and audio delivery; and (d) overall cognitive engagement — while avoiding sensorimotor activation.
The control protocol consisted of a 2-minute introductory relaxation phase followed by five 3-minute blocks alternating:
• Body Scanning (Somatic Attention): Sequential attention from feet to head noting warmth, pressure, and chair contact without movement or movement imagery.
• Spatial Navigation (Visuospatial Control): Mental tour of a familiar home or route, visualizing spatial layout without human limb movement.
Both tasks used standardized headphone audio.

2.7.4 Implementation of the PETTLEP Framework: Detailed Exposition
The seven PETTLEP dimensions were operationalized as summarized in Table 1."""

TABLE_ROWS = [
    ("PETTLEP Dimension", "Operationalization in This Study"),
    (
        "P – Physical",
        "Real task posture: seated on actual chair, normal clothing, real towel under hand; identical configuration to assessment.",
    ),
    (
        "E – Environment",
        "Same functional table setup and everyday objects as assessment; avoids abstract empty-bed imagery context.",
    ),
    (
        "T – Task",
        "Reach & Wipe calibrated to current upper-limb capacity and ADL relevance.",
    ),
    (
        "T – Timing",
        "Real-time imagery with paced 3-s reach / 3-s wipe / 3-s return counting; mental chronometry verified after each Imagine phase.",
    ),
    (
        "L – Learning",
        "Repeated Watch–Imagine–Rest cycles; corrective return-to-smooth-point guidance when imagery quality declines across blocks.",
    ),
    (
        "E – Emotion",
        "Scripts integrate comfort, confidence, and natural muscular sensation without excessive strain.",
    ),
    (
        "P – Perspective",
        "External (third-person) monitoring during Watch; internal first-person kinesthetic imagery during Imagine.",
    ),
]

GLOBAL_SUBS = [
    (r"\b22-minute\b", "17-minute"),
    (r"\b22 minute\b", "17-minute"),
    (r"five 4-minute blocks \(20 minutes total\)", "five 3-minute blocks (15 minutes total)"),
    (r"five 4-minute blocks", "five 3-minute blocks"),
    (r"5 blocks × 4 minutes = 20 minutes", "5 blocks × 3 minutes = 15 minutes"),
    (r"Minute 1 — Action Observation", "Watch phase (45 seconds) — Action Observation"),
    (r"Minutes 2–3 — Corrective Motor Imagery", "Imagine phase (75 seconds) — Real-Time Motor Imagery"),
    (r"Minute 4 — Rest", "Rest phase (60 seconds)"),
    (r"Total: 22 minutes", "Total: 17 minutes"),
    (r"total session duration \(22 minutes\)", "total session duration (~17 minutes)"),
    (r"matched 22 min", "matched ~17 min"),
    (r"22 min calibration \+ 5 × 4-min blocks", "2 min calibration + 5 × 3-min blocks"),
    (r"22-minute intervention", "17-minute intervention"),
    (r"22 min\)", "17 min)"),
]


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u2002", " ").replace("\u00a0", " ")).strip()


def set_para_red(p, text: str, bold: bool | None = None) -> None:
    p.text = ""
    if not text:
        return
    run = p.add_run(text)
    run.font.color.rgb = RED
    if bold is not None:
        run.bold = bold


def set_para_plain(p, text: str, bold: bool | None = None) -> None:
    p.text = ""
    if not text:
        return
    run = p.add_run(text)
    if bold is not None:
        run.bold = bold


def replace_para_text(p, old: str, new: str, *, red: bool = True) -> bool:
    if norm(old) not in norm(p.text) and norm(p.text) != norm(old):
        return False
    bold = p.runs[0].bold if p.runs else None
    if red:
        set_para_red(p, new, bold=bold)
    else:
        set_para_plain(p, new, bold=bold)
    return True


def apply_global_subs(text: str) -> str:
    out = text
    for pat, repl in GLOBAL_SUBS:
        out = re.sub(pat, repl, out, flags=re.IGNORECASE)
    return out


def update_proposal() -> None:
    doc = Document(str(PROP))
    start = end = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "Intervention:" and start is None:
            start = i
        if start is not None and p.text.strip().startswith("Statistical Analysis"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Intervention section not found in proposal")

    nonempty_indices = [i for i in range(start, end) if norm(doc.paragraphs[i].text)]
    for j, idx in enumerate(nonempty_indices):
        p = doc.paragraphs[idx]
        old = p.text
        new = PROPOSAL_INTERVENTION[j] if j < len(PROPOSAL_INTERVENTION) else ""
        bold = p.runs[0].bold if p.runs else None
        if p.text.strip() == "Intervention:":
            bold = True
        if norm(old) != norm(new):
            set_para_red(p, new, bold=bold)
    for idx in nonempty_indices[len(PROPOSAL_INTERVENTION) :]:
        if norm(doc.paragraphs[idx].text):
            set_para_red(doc.paragraphs[idx], "")

    for p in doc.paragraphs:
        new_text = apply_global_subs(p.text)
        if norm(new_text) != norm(p.text):
            bold = p.runs[0].bold if p.runs else None
            set_para_red(p, new_text, bold=bold)

    doc.save(str(PROP))


def update_manuscript() -> None:
    src = MAN_ALT if MAN_ALT.exists() else MAN
    doc = Document(str(src))

    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Methods:") and "single-blind" in p.text.lower():
            replace_para_text(p, p.text, ABSTRACT_METHODS_NEW, red=False)
            break

    # PETTLEP bullet list in section 1.5
    bullet_start = None
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith("Physical:") and bullet_start is None:
            bullet_start = i
            break
    if bullet_start:
        for j, bullet in enumerate(MANUSCRIPT_PETTLEP_BULLETS):
            idx = bullet_start + j
            if idx < len(doc.paragraphs):
                replace_para_text(doc.paragraphs[idx], doc.paragraphs[idx].text, bullet, red=False)

    for old, new in MANUSCRIPT_REPLACEMENTS:
        for p in doc.paragraphs:
            if norm(old) in norm(p.text) or norm(p.text) == norm(old):
                replace_para_text(p, old, new, red=False)
                break

    start = end = None
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("2.7 Intervention Protocol") and start is None:
            start = i
        if start is not None and t.startswith("2.8 Statistical"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Manuscript section 2.7 not found")

    new_paras = [ln.strip() for ln in MANUSCRIPT_SECTION_27.split("\n") if ln.strip()]
    replace_section_before(doc, start, end, new_paras, mark_red=False)

    if doc.tables:
        table = doc.tables[0]
        for ri, (a, b) in enumerate(TABLE_ROWS):
            if ri < len(table.rows):
                table.rows[ri].cells[0].text = a
                table.rows[ri].cells[1].text = b

    for p in doc.paragraphs:
        new_text = apply_global_subs(p.text)
        if norm(new_text) != norm(p.text):
            replace_para_text(p, p.text, new_text, red=False)

    try:
        doc.save(str(MAN_OUT))
    except PermissionError:
        alt = MAN_OUT.with_name(MAN_OUT.stem + "_PROTOCOL_UPDATED.docx")
        doc.save(str(alt))
        print(f"WARNING: saved manuscript to {alt} (original file locked)")


def replace_section_before(doc: Document, start: int, end: int, new_paras: list[str], mark_red: bool) -> None:
    """Replace paragraphs [start, end) with new_paras; paragraph at end is preserved."""
    from docx.oxml import OxmlElement
    from docx.text.paragraph import Paragraph

    end_el = doc.paragraphs[end]._element
    body = end_el.getparent()
    for i in range(end - 1, start - 1, -1):
        body.remove(doc.paragraphs[i]._element)

    heading_prefixes = (
        "2.7", "2.8", "Step ", "Phase ", "Goal:", "Exact ", "Total ",
        "Intervention Fidelity", "The control", "The seven", "Standardized script",
        "Procedure:", "Script:", "• ",
    )
    for text in reversed(new_paras):
        new_p = OxmlElement("w:p")
        end_el.addprevious(new_p)
        para = Paragraph(new_p, body)
        run = para.add_run(text)
        if any(text.strip().startswith(h) for h in heading_prefixes):
            run.bold = True
        if mark_red:
            run.font.color.rgb = RED


def main() -> None:
    update_proposal()
    update_manuscript()
    report = f"Updated protocol in:\n  {PROP}\n  {MAN_OUT}\n"
    Path(r"D:\Thesis app\NeuroLab\forms_extract\protocol_apply_report.txt").write_text(
        report, encoding="utf-8"
    )
    print(report)


if __name__ == "__main__":
    main()
