from docx import Document
from pathlib import Path

doc = Document(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal .docx")
lines = []
for i, p in enumerate(doc.paragraphs[:30]):
    t = p.text.strip()
    lines.append(f"{i}|{t[:95] if t else 'EMPTY'}")
Path(r"D:\Thesis app\NeuroLab\forms_extract\_fixed_head.txt").write_text("\n".join(lines), encoding="utf-8")
print("paras", len(doc.paragraphs))
