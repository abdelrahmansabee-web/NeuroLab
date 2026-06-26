# -*- coding: utf-8 -*-
"""Rebuild proposal cleanly from backup — no padding in form fields."""
from __future__ import annotations

import re
import shutil
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

BACKUP = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal _BACKUP.docx")
OUT = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal .docx")
RED = RGBColor(0xFF, 0x00, 0x00)

# Do not insert text into student-info form area (indices 0–18 in original backup)
FORM_END_IDX = 18

PROPOSAL_INTERVENTION = [
    "Intervention:",
    "The intervention follows a standardized PETTLEP-based AOMI session for the Reach & Wipe task. Total session duration is approximately 17 minutes (within the recommended 15–20 minute range): a 2-minute calibration phase followed by five 3-minute training blocks. Overt physical practice is excluded during the session to isolate cognitive-motor effects of observation and imagery.",
    "1. Baseline Assessment: After system calibration, three Reach & Wipe trials are performed at each time point (pre and post), with the mean used for analysis. Kinematics are recorded via the MediaPipe pipeline (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm).",
    "2. Intervention Phase: Participants are randomized to Experimental (PETTLEP-AOMI) or Control groups. Both sessions are matched in duration (~17 min), room setup, headphones, tablet height, and pre-recorded standardized audio delivery at Istinye and Biruni sites.",
    "Experimental Group — PETTLEP-Based AOMI:",
    "Session timing structure: Calibration (120 s) + five blocks × 180 s = 17 minutes. Each 3-minute block comprises Watch 45 s, Imagine 75 s, and Rest 60 s. Clinical practice allows 4–6 blocks; this study uses five blocks for dose standardization.",
    "Phase 1 — Calibration and Priming (120 seconds): Goal — deep mental focus and sensorimotor transfer from the unaffected to the affected limb before imagery.",
    "Calibration script (delivered in a calm voice via headphones): “Close your eyes. Feel the chair supporting your body weight. Take a deep breath in… and release it slowly.” (15 s). “Now focus on your unaffected hand. Feel the towel beneath it. Imagine wiping the table smoothly with that hand — see through your own eyes (internal perspective). Feel natural shoulder and arm contraction and the comfort of smooth movement. We will call this sensation ‘comfort’.” (45 s). “Now clearly imagine this strong ‘comfort’ signal travelling from your brain to your affected arm, awakening muscle fibres there. You are ready.” (60 s).",
    "Phase 2 — Training Blocks (repeated five times): Each block uses the same Watch → Imagine → Rest sequence applied to Reach & Wipe on the affected side.",
    "Step 1 — Watch / Action Observation (45 s; Perspective & Task): Participant views mirror-reversed first-person video of their own unaffected limb performing Reach & Wipe. External (third-person) monitoring is encouraged: “Open your eyes and watch the video. Observe as an external monitor — smooth arm extension, relaxed shoulders not rising toward the neck, and natural reach–wipe–return timing.”",
    "Step 2 — Imagine / Motor Imagery (75 s; Physical, Timing & Emotion): Eyes closed; internal (first-person) perspective; real-time imagery is critical. Script: “Return inside your body. Look through your own eyes at the table and towel under your affected hand. We will imagine the movement in real time.” The therapist paces with counting: Reach forward 1–2–3 (feel the towel slide); Wipe sideways 1–2–3 (feel arm weight moving smoothly with confidence); Return 1–2–3 (feel muscle relaxation). Three to five complete cycles are performed within 75 seconds according to each participant’s target pace.",
    "Step 3 — Neural Rest (60 s): “Stop imagining. Allow mind and muscles to relax completely. Do not think about movement for one full minute. Breathe calmly.” This prevents mental fatigue and preserves imagery quality in subsequent blocks.",
    "Implementation of the PETTLEP Framework:",
    "P – Physical: Participants adopt the actual task posture — seated on a real chair, normal clothing, real towel under the hand — matching assessment configuration.",
    "E – Environment: Imagery occurs in the same functional environment as assessment (real table setup with everyday objects), not an abstract empty clinical space.",
    "T – Task: Reach & Wipe is tailored to current upper-limb capacity and daily-life relevance.",
    "T – Timing (critical): Imagery duration matches real movement time; paced verbal counting during Imagine enforces 3-second reach, wipe, and return phases.",
    "L – Learning: Imagery content evolves across blocks as vividness improves; corrective guidance is given when imagery becomes jerky or compensatory.",
    "E – Emotion: Instructions integrate comfort, confidence, and awareness of natural muscular effort without strain.",
    "P – Perspective: External (third-person) during Watch for posture monitoring; internal (first-person, kinesthetic) during Imagine for muscle sensation.",
    "Total dose: 2 min calibration + 5 × 45 s observation + 5 × 75 s imagery + 5 × 60 s rest = 17 min. Mental chronometry is verified after each block. The physiotherapist completes a fidelity checklist (engagement, eyes closed during imagery, no overt movement, adverse events).",
    "Control Group — Cognitive and Somatic Control (matched ~17 min):",
    "Two-minute introductory relaxation followed by five 3-minute blocks alternating Body Scanning (sequential somatic attention from feet to head — warmth, pressure, chair contact — without movement or movement imagery) and Spatial Navigation (mental tour of a familiar home/route without human limb movement). Headphones deliver standardized audio; setup matches the experimental group while avoiding sensorimotor activation.",
    "Functional Relevance of the Reach & Wipe Task:",
    "Reach & Wipe provides ecological validity for upper-limb control, suits single-camera markerless tracking, and is sensitive to compensatory trunk lean and shoulder elevation. Kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) quantify movement quality via the NeuroLab pipeline at both Istinye Liv Bahçeşehir and Biruni University Hospital. Doç. Dr. Çiğdem Çınar supervises recruitment and data collection at the Biruni site.",
    "RCT note: Physical Reach & Wipe execution occurs only during pre/post assessment (three trials, mean analysed). The 17-minute intervention contains observation and imagery only — Watch (45 s), Imagine (75 s), and Rest (60 s) per block — without interleaved physical practice.",
]

SUBS = [
    (
        "Immediate Effects of PETTLEP-Based AOMI on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
        "Immediate Effects of a Single Session of PETTLEP-Based Action Observation and Motor Imagery (AOMI) on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
    ),
    ("OpenCap", "MediaPipe"),
    ("Box and Block Test (BBT)", "Wolf Motor Function Test – 4-item short form (WMFT-4)"),
    ("Box and Block Test", "WMFT-4"),
    ("BBT", "WMFT-4"),
    ("Number of Velocity Peaks - NVP", "smoothness pause percentage (smoothness_pause_pct)"),
    ("Reach-to-Grasp", "Reach & Wipe"),
    ("Reach-to-Grasp (RTG)", "Reach & Wipe"),
    ("MIQ-3", "KVIQ-10"),
    ("Movement Imagery Questionnaire-3", "Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10)"),
    ("MAS ≤ 3", "MAS ≤ 2"),
    ("60 frames per second", "30 frames per second"),
    ("iPhone 14 Pro max & iPhone XS max", "a single smartphone or webcam"),
    ("45° angles (one anterior-lateral and one posterior-lateral)", "a standardized frontal view"),
    ("OpenCap cloud platform", "NeuroLab MediaPipe pipeline"),
    ("OpenSim musculoskeletal models", "MediaPipe landmark trajectories with optional OpenSim IK"),
    (
        "Keywords: Motor imagery, action observation, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
        "Keywords: Motor imagery, action observation, AOMI, PETTLEP, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
    ),
    (
        "The study population consists of stroke patients who were admitted or followed up at the Neurorehabilitation Clinic of Istinye University liv Bahçeşehir Hospital, Istanbul, Türkiye.",
        "The study population consists of stroke patients admitted or followed up at Istinye University Liv Bahçeşehir Hospital Neurorehabilitation Clinic and Biruni University Hospital Physical Medicine and Rehabilitation Clinic, Istanbul, Türkiye (multisite).",
    ),
    (
        "prospective randomized controlled trial (RCT) design will be employed, involving two conditions",
        "prospective multisite randomized controlled trial (RCT) design will be employed, involving two conditions",
    ),
    (
        "Uhlrich SD, Falisse A, Kidziński Ł, Muccini J, Ko M, Chaudhari AS, et al. OpenCap: 3D human movement dynamics from smartphone videos. PLOS Comput Biol. 2023;19(10):e1011462.",
        "Lugaresi C, Tang J, Nash H, et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172. 2019.",
    ),
    (
        "Mathiowetz V, Volland G, Kashman N, Weber K. Adult norms for the Box and Block Test of manual dexterity. Am J Occup Ther. 1985;39(6):318-326.",
        "Wolf SL, Catlin PA, Ellis M, et al. Assessing Wolf Motor Function Test scores against normative values. Arch Phys Med Rehabil. 2001;82(5):609-614.",
    ),
    (
        "To minimize inter-rater variability, one physiotherapist at each centre will conduct the evaluations.",
        "To minimize inter-rater variability, one trained physiotherapist at each centre will conduct clinical evaluations using standardized scripts, while kinematic extraction remains fully automated and blinded via the NeuroLab pipeline.",
    ),
    (
        "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature.",
        "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature by providing objective, single-camera kinematic evidence for immediate effects of a PETTLEP-structured AOMI session without interleaved physical practice.",
    ),
    (
        "Gross Manual Dexterity — Box and Block Test (BBT):",
        "Upper Limb Function — Wolf Motor Function Test, 4-item short form (WMFT-4):",
    ),
    (
        "Movement Imagery Ability — Movement Imagery Questionnaire-3 (MIQ-3)",
        "Movement Imagery Ability — Kinesthetic and Visual Imagery Questionnaire (KVIQ-10)",
    ),
    (
        "Correlations Subjective vividness of imagery (MIQ-3 scores)",
        "Correlations Subjective vividness of imagery (KVIQ-10 scores)",
    ),
]

PRIMARY_OUTCOME = (
    "· The primary outcome is change in movement smoothness (smoothness_pause_pct) — percentage of active "
    "movement time below a velocity threshold (lower = smoother). Secondary kinematics include "
    "total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, and total_peak_velocity via the NeuroLab pipeline."
)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u2002", " ").replace("\u00a0", " ")).strip()


def clear_para(p) -> None:
    el = p._element
    for child in list(el):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            el.remove(child)


def apply_red_diff(p, old_text: str, new_text: str, bold: bool | None = None) -> None:
    clear_para(p)
    if not new_text:
        return
    if norm(old_text) == norm(new_text):
        run = p.add_run(new_text)
        if bold:
            run.bold = True
        return
    old_w = norm(old_text).split()
    new_w = norm(new_text).split()
    sm = SequenceMatcher(None, old_w, new_w)
    for tag, _i1, _i2, j1, j2 in sm.get_opcodes():
        chunk = " ".join(new_w[j1:j2])
        if not chunk:
            continue
        if j2 < len(new_w):
            chunk += " "
        run = p.add_run(chunk)
        if bold:
            run.bold = True
        if tag != "equal":
            run.font.color.rgb = RED


def apply_subs(text: str) -> str:
    t = text
    for old, new in SUBS:
        if norm(old) in norm(t):
            t = norm(t).replace(norm(old), norm(new), 1)
    return t


def main() -> None:
    shutil.copy2(BACKUP, OUT)
    doc = Document(str(BACKUP))
    orig = [p.text for p in doc.paragraphs]

    # Build new texts without touching empty form spacers
    new_texts = list(orig)
    for i, t in enumerate(new_texts):
        if not norm(t):
            continue
        if i <= FORM_END_IDX:
            # Only allow title/substitution in form area — never inject new paragraphs
            new_texts[i] = apply_subs(t)
        else:
            new_texts[i] = apply_subs(t)

    for i, t in enumerate(new_texts):
        if "primary outcome of this study is the change in movement smoothness" in norm(t).lower():
            new_texts[i] = PRIMARY_OUTCOME

    # Intervention section
    start = end = None
    for i, t in enumerate(new_texts):
        if start is None and norm(t) == "Intervention:":
            start = i
        if start is not None and norm(t).startswith("Statistical Analysis"):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("Intervention section not found")

    nonempty = [i for i in range(start, end) if norm(new_texts[i])]
    for j, idx in enumerate(nonempty):
        new_texts[idx] = PROPOSAL_INTERVENTION[j] if j < len(PROPOSAL_INTERVENTION) else ""
    for idx in nonempty[len(PROPOSAL_INTERVENTION) :]:
        new_texts[idx] = ""

    # Write back — preserve empty paragraphs exactly
    out_doc = Document(str(BACKUP))
    changed = 0
    for i, p in enumerate(out_doc.paragraphs):
        old = orig[i]
        new = new_texts[i]
        if not norm(old) and not norm(new):
            continue  # leave empty spacers untouched
        bold = p.runs[0].bold if p.runs else None
        if norm(old) != norm(new):
            changed += 1
            apply_red_diff(p, old, new, bold=bold)

    out_doc.save(str(OUT))
    wc = len(re.findall(r"\w+", "\n".join(p.text for p in out_doc.paragraphs), re.UNICODE))
    report = f"Rebuilt: {OUT}\nFrom: {BACKUP}\nParagraphs changed: {changed}\nWords: {wc}\nForm fields preserved.\n"
    Path(r"D:\Thesis app\NeuroLab\forms_extract\proposal_rebuild_report.txt").write_text(report, encoding="utf-8")
    print(report)


if __name__ == "__main__":
    main()
