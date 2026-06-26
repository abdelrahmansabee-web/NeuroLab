# -*- coding: utf-8 -*-
"""Extract signature (transparent BG), enhance quality, insert into kurum izni docx."""
from __future__ import annotations

import shutil
from pathlib import Path

import numpy as np
from docx import Document
from docx.shared import Inches
from PIL import Image, ImageEnhance, ImageFilter, ImageOps

SRC = Path(
    r"C:\Users\acer\.cursor\projects\d-Thesis-app-NeuroLab\assets"
    r"\c__Users_acer_AppData_Roaming_Cursor_User_workspaceStorage_empty-window_images_"
    r"WhatsApp_Image_2026-06-15_at_1.54.52_PM-2b3ad6b3-9e2f-4e0d-be56-744da61505f8.png"
)
PROCESSED = Path(r"D:\Thesis app\NeuroLab\assets\abdelrahman_signature_hq.png")
TARGET = Path(r"D:\Thesis app\manuscript f\REVIZYON_PAKETI\kurum_izni_biruni_DOLDURULMUS.docx")


def extract_signature_hq(src: Path, out: Path, scale: float = 4.0) -> Path:
    img = Image.open(src).convert("RGBA")
    arr = np.array(img, dtype=np.float32)
    r, g, b, _ = arr[..., 0], arr[..., 1], arr[..., 2], arr[..., 3]

    # Blue pen ink: strong B, B > R and B > G
    blue_score = b - np.maximum(r, g)
    max_rgb = np.maximum(np.maximum(r, g), b)
    saturation = b / (max_rgb + 1e-6)
    ink = (blue_score > 18) & (b > 70) & (saturation > 0.45)

    # Soft alpha from ink strength for anti-aliased edges
    strength = np.clip((blue_score - 10) / 45.0, 0, 1) * ink.astype(np.float32)
    strength = np.power(strength, 0.85)

    alpha = (strength * 255).astype(np.uint8)
    # Deep navy ink (natural ballpoint on white paper)
    ink_r = np.clip(15 + (255 - b) * 0.04, 0, 45).astype(np.uint8)
    ink_g = np.clip(25 + (255 - b) * 0.08, 0, 75).astype(np.uint8)
    ink_b = np.clip(120 + b * 0.35, 0, 210).astype(np.uint8)
    ink_rgb = np.stack([ink_r, ink_g, ink_b, alpha], axis=-1)
    sig = Image.fromarray(ink_rgb, "RGBA")

    # Crop to ink bounding box with padding
    ys, xs = np.where(alpha > 12)
    if len(xs) == 0:
        raise ValueError("No signature ink detected")
    pad = 8
    x0, x1 = max(0, xs.min() - pad), min(sig.width, xs.max() + pad)
    y0, y1 = max(0, ys.min() - pad), min(sig.height, ys.max() + pad)
    sig = sig.crop((x0, y0, x1, y1))

    # Upscale for crisp print
    new_w = max(1, int(sig.width * scale))
    new_h = max(1, int(sig.height * scale))
    sig = sig.resize((new_w, new_h), Image.Resampling.LANCZOS)

    # Sharpen ink edges slightly
    sig = sig.filter(ImageFilter.UnsharpMask(radius=1.2, percent=180, threshold=2))

    # Boost contrast on alpha channel only
    r_, g_, b_, a_ = sig.split()
    a_ = ImageEnhance.Contrast(a_).enhance(1.15)
    sig = Image.merge("RGBA", (r_, g_, b_, a_))

    out.parent.mkdir(parents=True, exist_ok=True)
    sig.save(out, "PNG", optimize=True)
    return out


def insert_signature(doc_path: Path, sig_path: Path, width_in: float = 2.0) -> Path:
    doc = Document(str(doc_path))
    target = None
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("Tarih:") and "İmza:" in t and "Onay Tarihi" not in t:
            target = p
            break
    if target is None:
        raise RuntimeError("Signature paragraph not found")

    for run in list(target.runs):
        run._element.getparent().remove(run._element)

    # Remove old inline images in this paragraph if any
    for drawing in list(target._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing")):
        drawing.getparent().remove(drawing)

    target.add_run("Tarih: 15/06/2026  İmza: ")
    pic_run = target.add_run()
    pic_run.add_picture(str(sig_path), width=Inches(width_in))

    doc.save(str(doc_path))
    return doc_path


def main() -> None:
    if not SRC.exists():
        raise FileNotFoundError(SRC)
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)

    sig = extract_signature_hq(SRC, PROCESSED)
    png_copy = TARGET.parent / "abdelrahman_signature_hq.png"
    try:
        out = insert_signature(TARGET, sig, width_in=2.15)
    except PermissionError:
        alt = TARGET.with_name("kurum_izni_biruni_DOLDURULMUS_IMZALI.docx")
        shutil.copy2(TARGET, alt)
        out = insert_signature(alt, sig, width_in=2.15)
        print(f"NOTE: Original was locked; saved as: {alt.name}")
    shutil.copy2(sig, png_copy)
    print(f"Processed signature: {sig} ({sig.stat().st_size} bytes)")
    print(f"Updated document: {out}")


if __name__ == "__main__":
    main()
