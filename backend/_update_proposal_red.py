# -*- coding: utf-8 -*-
"""Update Abdelrahman Sabee Proposal .docx from manuscript facts; mark edits in red."""
from __future__ import annotations

import re
import shutil
from difflib import SequenceMatcher
from pathlib import Path

from docx import Document
from docx.shared import RGBColor

SRC = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal .docx")
BACKUP = Path(r"D:\Thesis app\manuscript f\Abdelrahman Sabee Proposal _BACKUP.docx")
OUT = SRC
TARGET_WORDS = 4631
RED = RGBColor(0xFF, 0x00, 0x00)

# paragraph index -> bold heading (from structure extract)
BOLD_IDX = {
    0, 12, 14, 16, 19, 20, 37, 47, 59, 67, 92, 111, 118, 127, 137, 151, 161, 171, 179, 196,
    200, 211, 235, 246, 281, 312, 321, 324, 325, 327, 329, 330, 333,
}

NEW_INTERVENTION = [
    "Intervention:",
    "The intervention protocol follows the NeuroLab PETTLEP-based AOMI session application (22 minutes total): a 2-minute sensorimotor transfer calibration phase followed by five 4-minute blocks (20 minutes), administered immediately after pre-assessment in the same environmental setting. Unlike earlier scaffolded protocols with interleaved physical practice, this version excludes overt execution during imagery to isolate the specific cognitive-motor effects of observation and imagery and to reduce burden on participants with severe motor impairment.",
    "1. Baseline Assessment: After system calibration, three Reach & Wipe trials are performed at each time point (pre and post), with the mean used for analysis. Kinematics are recorded via the MediaPipe pipeline (pre-intervention smoothness_pause_pct, total_trunk_palm_ratio, and shoulder_vert_norm).",
    "2. Intervention Phase: Participants are randomized to Experimental (PETTLEP-AOMI) or Control (cognitive/somatic) groups. Both sessions last 22 minutes in the same room with headphones, tablet at standardized height, and pre-recorded standardized digital audio files to ensure replication across participants and sites.",
    "Experimental Group — PETTLEP-Based AOMI (NeuroLab protocol):",
    "Calibration (2 min): Participants learn the notice–name–transfer strategy on the unaffected limb. They attend to a specific kinesthetic sensation (e.g., smoothness, stability, relaxation), label it with a single-word cue, and mentally transfer that sensation to the affected limb. Supported limb awareness and surface-contact anchoring enhance sensorimotor engagement without resistance or active movement.",
    "Standardized therapist scripts (Reach & Wipe — adapted from NeuroLab PETTLEP guide): (1) Setup — sit upright, feet flat, back supported; towel under affected hand on table; slow deep breath. (2) External observation (block minute 1) — watch mirror-reversed video; note relaxed shoulder and stable trunk, no compensatory elevation. (3) Internal imagery (minutes 2–3) — eyes closed; first-person view; focus on upper arm and shoulder muscles; imagine real-time reach (~3 s) and return (~3 s). (4) Emotion — feel towel slide smoothly; comfort and confidence without strain. (5) Learning — if imagery becomes jerky, return to last smooth point; adjust script per block (e.g., finger opening). After each block: vividness rating (1–10) and kinesthetic muscle sensation check.",
    "Each of five 4-minute blocks comprises: Minute 1 — Action Observation of mirror-reversed first-person video of the participant’s own unaffected upper limb performing Reach & Wipe (horizontally flipped to appear as the affected side); Minutes 2–3 — Corrective motor imagery with eyes closed, guided by block-specific audio cues delivered via headphones; Minute 4 — Rest with eyes open, without movement or mental imagery.",
    "If imagined movement becomes jerky, effortful, or associated with trunk or shoulder compensation, the participant mentally returns to the last smooth point and continues with the corrected motor plan.",
    "Block 1 — Movement Initiation Transfer: attention to calm, clear start; imagery transfers an easy “start now” signal without hesitation or rush.",
    "Block 2 — Trunk Stability Transfer: attention to quiet, stable trunk; imagery reinforces feedforward postural control — arm reaches forward while trunk remains still and supported.",
    "Block 3 — Shoulder Girdle Control Transfer: attention to low, relaxed shoulder; imagery transfers a calm, heavy shoulder staying away from the ear throughout reach.",
    "Block 4 — Elbow Opening and Muscle Quieting: attention to easy elbow opening; imagery targets smooth extension and reduced co-contraction / antagonist quieting.",
    "Block 5 — Integrated Smooth Movement: all constraints integrated — clear start, stable trunk, relaxed shoulder, smooth elbow, continuous hand trajectory — imagining complete Reach & Wipe at natural timing.",
    "Total dose: 2 min calibration + 5 min observation + 10 min corrective imagery + 5 min rest = 22 min. Mental chronometry is verified after each block (“Did the imagined movement feel at the same pace as normal?”). The supervising physiotherapist completes a structured fidelity checklist after each block (engagement, eyes closed during imagery, no overt movement, adverse events).",
    "Control Group — Cognitive and Somatic Control (matched 22 min):",
    "Five 4-minute blocks preceded by a 2-minute introductory relaxation phase, matching AOMI in duration, room setup, tablet presence, headphones, and overall cognitive engagement while avoiding sensorimotor activation. Blocks alternate Body Scanning (sequential somatic attention from feet to head — warmth, pressure, chair contact — without movement or movement imagery) and Spatial Navigation (mental tour of a familiar home/route, visualizing spatial layout without human limb movement). Both tasks use standardized headphone audio.",
    "Implementation of the PETTLEP Framework:",
    "P – Physical: identical seated posture and Reach & Wipe starting configuration as assessment.",
    "E – Environment: same room, chair, lighting, table distance, and tablet placement as assessment.",
    "T – Task: Reach & Wipe — in-air reaching and wiping; imagery constraints target initiation, smoothness, trunk–limb dissociation, and shoulder control.",
    "T – Timing: real-time imagery; mental chronometry checked after each block with corrective guidance if temporal distortion is reported.",
    "L – Learning: observation-then-imagery cycles with notice–name–transfer calibration; individualized mirror-reversed video provides calibrated visual reference before each imagery bout.",
    "E – Emotion: audio cues emphasize ease, fluidity, heaviness of shoulder, stillness of trunk; counteract anxiety and effort-related hypertonicity.",
    "P – Perspective: internal first-person kinesthetic imagery; mirror-reversed video reinforces anatomical congruence with the affected side.",
    "Functional Relevance of the Reach & Wipe Task:",
    "Reach & Wipe was selected for ecological validity in upper limb control, suitability for single-camera markerless tracking, and sensitivity to compensatory trunk lean and shoulder elevation. The proximal-to-distal block order mirrors normal motor control (initiation and trunk stability before distal reach). Kinematic outcomes (smoothness_pause_pct, total_trunk_palm_ratio, shoulder_vert_norm, total_peak_velocity) quantify movement quality beyond ordinal clinical scales and align with the NeuroLab analysis pipeline used for automated, blinded extraction at both Istinye and Biruni sites.",
    "RCT note: Physical Reach & Wipe execution occurs only during pre/post assessment (three trials, mean analysed). The 22-minute intervention contains observation and imagery only — mapping clinic PETTLEP structure (video observation + mental practice + feedback) to five 4-minute blocks without interleaved physical practice during the session.",
]

TEAM_LINES = [
    "",
    "Sorumlu Araştırmacı (Principal Investigator): Abdelrahman Walid Hamza Mohamed Elsayed Sabee — Istinye University",
    "",
    "Yardımcı Araştırmacı (Co-Investigator): Doç. Dr. Çiğdem Çınar — Biruni Üniversitesi Hastanesi, Fiziksel Tıp ve Rehabilitasyon",
    "",
    "Veri Toplama Merkezleri (Multisite): Istinye University Liv Bahçeşehir Hospital; Biruni University Hospital",
]

SUBS = [
    (
        "Immediate Effects of PETTLEP-Based AOMI on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
        "Immediate Effects of a Single Session of PETTLEP-Based Action Observation and Motor Imagery (AOMI) on Upper Limb Kinematics in Stroke Survivors: A Randomized Controlled Trial.",
    ),
    ("OpenCap", "MediaPipe"),
    ("Box and Block Test (BBT)", "Wolf Motor Function Test – 4-item short form (WMFT-4)"),
    ("Box and Block Test", "WMFT-4"),
    ("BBT", "WMFT-4"),
    ("Number of Velocity Peaks - NVP", "smoothness pause percentage (smoothness_pause_pct)"),
    ("NVP", "smoothness_pause_pct"),
    ("Reach-to-Grasp", "Reach & Wipe"),
    ("Reach-to-Grasp (RTG)", "Reach & Wipe"),
    ("MIQ-3", "KVIQ-10"),
    ("Movement Imagery Questionnaire-3", "Kinesthetic and Visual Imagery Questionnaire – 10 items (KVIQ-10)"),
    ("MAS ≤ 3", "MAS ≤ 2"),
    ("MAS 2–3", "MAS 2"),
    ("60 frames per second", "30 frames per second"),
    ("iPhone 14 Pro max & iPhone XS max", "a single smartphone or webcam"),
    ("45° angles (one anterior-lateral and one posterior-lateral)", "a standardized frontal view"),
    ("OpenCap cloud platform", "NeuroLab MediaPipe pipeline"),
    ("OpenSim musculoskeletal models", "MediaPipe landmark trajectories with optional OpenSim IK"),
    (
        "Keywords: Motor imagery, action observation, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
        "Keywords: Motor imagery, action observation, AOMI, PETTLEP, stroke, MediaPipe, kinematics, upper limb rehabilitation.",
    ),
    (
        "The study population consists of stroke patients who were admitted or followed up at the Neurorehabilitation Clinic of Istinye University liv Bahçeşehir Hospital, Istanbul, Türkiye.",
        "The study population consists of stroke patients admitted or followed up at Istinye University Liv Bahçeşehir Hospital Neurorehabilitation Clinic and Biruni University Hospital Physical Medicine and Rehabilitation Clinic, Istanbul, Türkiye (multisite).",
    ),
    (
        "prospective randomized controlled trial (RCT) design will be employed, involving two conditions",
        "prospective multisite randomized controlled trial (RCT) design will be employed, involving two conditions",
    ),
    (
        "Uhlrich SD, Falisse A, Kidziński Ł, Muccini J, Ko M, Chaudhari AS, et al. OpenCap: 3D human movement dynamics from smartphone videos. PLOS Comput Biol. 2023;19(10):e1011462.",
        "Lugaresi C, Tang J, Nash H, et al. MediaPipe: A Framework for Building Perception Pipelines. arXiv:1906.08172. 2019.",
    ),
    (
        "Mathiowetz V, Volland G, Kashman N, Weber K. Adult norms for the Box and Block Test of manual dexterity. Am J Occup Ther. 1985;39(6):318-326.",
        "Wolf SL, Catlin PA, Ellis M, et al. Assessing Wolf Motor Function Test scores against normative values. Arch Phys Med Rehabil. 2001;82(5):609-614.",
    ),
    (
        "markerless motion capture pipeline, on the kinematics",
        "markerless motion capture pipeline (NeuroLab Stroke Rehabilitation Research Platform), on the kinematics",
    ),
    (
        "To minimize inter-rater variability, one physiotherapist at each centre will conduct the evaluations.",
        "To minimize inter-rater variability, one trained physiotherapist at each centre will conduct clinical evaluations using standardized scripts, while kinematic extraction remains fully automated and blinded via the NeuroLab pipeline.",
    ),
    (
        "Visual Analog Scale for pain (VAS)",
        "Visual Analog Scale for pain (VAS); Visual Analog Mood Scale – 4 dimensions (VAMS-4) pre/post; Motor Difference Rating Scale (MDRS) post-only; International Physical Activity Questionnaire (IPAQ) at baseline",
    ),
    (
        "Handling Missing Data: In order to preserve randomization and reduce bias in the estimate of effect, an ITT analysis will be rigorously performed.",
        "Handling Missing Data: In order to preserve randomization and reduce bias in the estimate of effect, an intention-to-treat (ITT) analysis will be rigorously performed.",
    ),
    (
        "instead they will be treated by statistical imputation (e.g. Last Observation Carried Forward [LOCF] or Linear Mixed Models)",
        "instead missing values will be imputed using Last Observation Carried Forward (LOCF) as the primary method, with Linear Mixed Models as a sensitivity analysis",
    ),
    (
        "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature.",
        "Therefore, the expectation of this study is to fill a critical gap in neuro-rehabilitation literature by providing the first objective, single-camera kinematic evidence for immediate effects of a PETTLEP-structured AOMI session without interleaved physical practice.",
    ),
    (
        "gross manual dexterity of the affected upper limb",
        "upper limb functional performance (WMFT-4) of the affected upper limb",
    ),
    (
        "Correlations Subjective vividness of imagery (MIQ-3 scores)",
        "Correlations Subjective vividness of imagery (KVIQ-10 scores)",
    ),
    (
        "The necessary written permission for use has been obtained from the author who conducted the Turkish validity and reliability study of the MIQ-3 used",
        "The KVIQ-10 is a validated imagery ability instrument; Turkish linguistic adaptation will follow published psychometric guidelines for imagery questionnaires",
    ),
    (
        "Gross Manual Dexterity — Box and Block Test (BBT):",
        "Upper Limb Function — Wolf Motor Function Test, 4-item short form (WMFT-4):",
    ),
    (
        "Movement Imagery Ability — Movement Imagery Questionnaire-3 (MIQ-3)",
        "Movement Imagery Ability — Kinesthetic and Visual Imagery Questionnaire (KVIQ-10)",
    ),
    (
        "Williams SE, Cumming J, Ntoumanis N, Nordin-Bates SM, Ramsey R, Hall C. Further validation and development of the Movement Imagery Questionnaire-3",
        "Malouin F, Richards CL, Jackson PL, et al. The Kinesthetic and Visual Imagery Questionnaire (KVIQ) for assessing motor imagery in persons with physical disabilities",
    ),
    (
        "Uğur Y, Coşkun H, Şenyurt AY. The Movement Imagery Questionnaire-3: Reliability and Validity Study on Turkish Sample. Spormetre Beden Eğitimi ve S",
        "Malouin F, Richards CL, Jackson PL, et al. The Kinesthetic and Visual Imagery Questionnaire (KVIQ) for assessing motor imagery in persons with physical disabilities: a reliability and construct validity study. J Neurol Phys Ther. 2007;31(1):20-29.",
    ),
]

PADS = [
    "Secondary clinical outcomes include WMFT-4 (upper limb function), KVIQ-10 (kinesthetic and visual imagery ability), VAMS-4 (happy, calm, sad, tense mood dimensions), VAS pain, IPAQ physical activity, and MDRS post-intervention motor control perception. Together these capture functional transfer, imagery capacity, affective state, and safety alongside objective kinematics.",
    "Kinematic videos are processed offline through the NeuroLab Stroke Rehabilitation Research Platform using MediaPipe Pose Landmarker (33 landmarks, 30 fps) with ZoeDepth metric scaling. Optional OpenSim inverse kinematics (14 DOF) supports exploratory joint-angle analysis. All participants receive identical segmentation algorithms to preserve blinding.",
    "Moderator analyses will test whether higher baseline KVIQ-10 scores predict greater kinematic improvement in the AOMI group, and whether VAMS-4 positive affect post-session associates with larger reductions in smoothness_pause_pct. Holm–Bonferroni correction will be applied across the secondary kinematic family (k = 8).",
    "Multisite procedures are harmonized: both Istinye Liv Bahçeşehir and Biruni University Hospital use the same assessment script, intervention audio files, Reach & Wipe setup, and NeuroLab export format. Doç. Dr. Çiğdem Çınar supervises recruitment and data collection at the Biruni site as co-investigator.",
    "Session fidelity is documented in NeuroLab for each participant: group assignment, pre/post clinical scales, three Reach & Wipe video trials per time point, and automated kinematic export. Standardized pre-recorded audio files and mirror-reversed participant video ensure identical delivery at Istinye and Biruni sites.",
    "Safety monitoring includes VAS pain at each time point, adverse event recording on the fidelity checklist, and immediate cessation of imagery if distress, dizziness, or pain escalation occurs; participants may withdraw without penalty at any stage.",
]

PRIMARY_OUTCOME = (
    "· The primary outcome is change in movement smoothness (smoothness_pause_pct) — percentage of active "
    "movement time below a velocity threshold (lower = smoother). Secondary kinematics include "
    "total_duration_s, total_trunk_palm_ratio, shoulder_vert_norm, and total_peak_velocity via the NeuroLab pipeline."
)


def norm(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").replace("\u2002", " ").replace("\u00a0", " ")).strip()


def word_count(texts: list[str]) -> int:
    return len(re.findall(r"\w+", "\n".join(texts), re.UNICODE))


def clear_para(p) -> None:
    el = p._element
    for child in list(el):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            el.remove(child)


def apply_red_diff(p, old_text: str, new_text: str, bold: bool = False, all_red: bool = False) -> None:
    clear_para(p)
    if not new_text:
        return
    if all_red or norm(old_text) == norm(new_text):
        run = p.add_run(new_text)
        if bold:
            run.bold = True
        if all_red and norm(old_text) != norm(new_text):
            run.font.color.rgb = RED
        return
    old_w = norm(old_text).split()
    new_w = norm(new_text).split()
    sm = SequenceMatcher(None, old_w, new_w)
    for tag, _i1, _i2, j1, j2 in sm.get_opcodes():
        chunk = " ".join(new_w[j1:j2])
        if not chunk:
            continue
        if j2 < len(new_w):
            chunk += " "
        run = p.add_run(chunk)
        if bold:
            run.bold = True
        if tag != "equal":
            run.font.color.rgb = RED


def apply_subs(text: str) -> str:
    t = text
    for old, new in SUBS:
        if norm(old) in norm(t):
            t = norm(t).replace(norm(old), norm(new), 1)
    return t


def rewrite_intervention(texts: list[str]) -> list[str]:
    start = end = None
    for i, t in enumerate(texts):
        if start is None and norm(t) == "Intervention:":
            start = i
        if start is not None and norm(t).startswith("Statistical Analysis"):
            end = i
            break
    if start is None or end is None:
        return texts
    out = list(texts)
    block_len = end - start
    for j in range(block_len):
        idx = start + j
        if j < len(NEW_INTERVENTION):
            out[idx] = NEW_INTERVENTION[j]
        else:
            out[idx] = ""
    return out


def insert_team_lines(texts: list[str]) -> tuple[list[str], int | None]:
    out: list[str] = []
    advisor_idx: int | None = None
    for i, t in enumerate(texts):
        out.append(t)
        if advisor_idx is None and "Danışmanı (Advisor):" in t:
            advisor_idx = i
            out.extend(TEAM_LINES)
    return out, advisor_idx


def pad_texts(texts: list[str], target: int) -> list[str]:
    out = list(texts)
    pi = 0
    while word_count(out) < target - 5 and pi < len(PADS):
        placed = False
        for i, t in enumerate(out):
            if not norm(t):
                out[i] = PADS[pi]
                pi += 1
                placed = True
                break
        if not placed:
            break
    return out


def aligned_orig(orig: list[str], advisor_idx: int | None) -> list[str]:
    if advisor_idx is None:
        return list(orig)
    out = list(orig)
    insert_at = advisor_idx + 1
    for _ in TEAM_LINES:
        out.insert(insert_at, "")
    return out


def shift_bold_idx(advisor_idx: int | None) -> set[int]:
    if advisor_idx is None:
        return set(BOLD_IDX)
    shift = len(TEAM_LINES)
    return {i if i <= advisor_idx else i + shift for i in BOLD_IDX}


def main() -> None:
    if not BACKUP.exists():
        shutil.copy2(SRC, BACKUP)

    doc = Document(str(BACKUP))
    orig = [p.text for p in doc.paragraphs]
    wc0 = word_count(orig)

    texts = [apply_subs(t) for t in orig]

    for i, t in enumerate(texts):
        if "primary outcome of this study is the change in movement smoothness" in norm(t).lower():
            texts[i] = PRIMARY_OUTCOME
            break

    texts = rewrite_intervention(texts)
    texts, advisor_idx = insert_team_lines(texts)
    orig_aligned = aligned_orig(orig, advisor_idx)
    bold_idx = shift_bold_idx(advisor_idx)
    texts = pad_texts(texts, TARGET_WORDS)

    new_doc = Document(str(BACKUP))
    paras = new_doc.paragraphs

    from docx.text.paragraph import Paragraph
    from docx.oxml import OxmlElement

    if advisor_idx is not None:
        anchor = paras[advisor_idx]
        cur = anchor
        for _ in TEAM_LINES:
            new_p = OxmlElement("w:p")
            cur._p.addnext(new_p)
            para = Paragraph(new_p, cur._parent)
            if anchor.style:
                para.style = anchor.style
            cur = para

    paras = new_doc.paragraphs
    n = min(len(paras), len(texts), len(orig_aligned))
    changed = 0
    for i in range(n):
        p = paras[i]
        old = orig_aligned[i]
        new = texts[i]
        team_line = advisor_idx is not None and advisor_idx < i <= advisor_idx + len(TEAM_LINES)
        bold = i in bold_idx
        if norm(old) != norm(new):
            changed += 1
        apply_red_diff(
            p,
            old,
            new,
            bold=bold,
            all_red=team_line or norm(old) != norm(new),
        )

    new_doc.save(str(OUT))
    wc1 = word_count([p.text for p in Document(str(OUT)).paragraphs])
    report = "\n".join(
        [
            f"Source backup: {BACKUP}",
            f"Updated: {OUT}",
            f"Paragraphs changed: {changed}",
            f"Words before: {wc0}",
            f"Words after: {wc1}",
            f"Target: {TARGET_WORDS} (delta {wc1 - TARGET_WORDS:+d})",
        ]
    )
    Path(r"D:\Thesis app\NeuroLab\forms_extract\proposal_red_update_report.txt").write_text(
        report, encoding="utf-8"
    )
    print(report)


if __name__ == "__main__":
    main()
